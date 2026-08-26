# -*- coding: utf-8 -*-
"""
Generate Official Word Document for Puzzle Data & Quantum AI Fullstack Use Cases
File: offering/generate_puzzle_quantum_usecases_docx.py
Output: offering/docx/2026-08-26_Puzzle_Data_Quantum_AI_Fullstack_Usecases.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Colors
COLOR_MZC_BLUE = RGBColor(0x00, 0xAB, 0xF0)
COLOR_MZC_NAVY = RGBColor(0x00, 0x2B, 0x49)
COLOR_NVIDIA_GREEN = RGBColor(0x76, 0xB9, 0x00)
COLOR_DELL_BLUE = RGBColor(0x00, 0x71, 0xC5)
COLOR_DARK_TEXT = RGBColor(0x1F, 0x29, 0x37)
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)
COLOR_PURPLE = RGBColor(0x9D, 0x4E, 0xDD)
COLOR_GOLD = RGBColor(0xD4, 0xAF, 0x37)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_DARK_NAVY = "002B49"
HEX_MZC_BLUE = "00ABF0"
HEX_LIGHT_GRAY = "F8FAFC"
HEX_BORDER = "E2E8F0"
HEX_PURPLE_BG = "F3E8FF"

FONT_MAIN = "맑은 고딕"
FONT_ENG = "Calibri"

def set_cell_background(cell, hex_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, border_color="D1D5DB"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
            f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="{border_color}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_header_block(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = FONT_MAIN
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_MZC_NAVY

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(8)
        run_sub = p_sub.add_run(subtitle)
        run_sub.font.name = FONT_MAIN
        run_sub.font.size = Pt(9.5)
        run_sub.font.color.rgb = COLOR_MUTED

def add_bullet(doc, strong_text, body_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.4)
    
    r1 = p.add_run(f"• {strong_text}: ")
    r1.font.name = FONT_MAIN
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_MZC_NAVY

    r2 = p.add_run(body_text)
    r2.font.name = FONT_MAIN
    r2.font.size = Pt(9)
    r2.font.color.rgb = COLOR_DARK_TEXT

def build_document():
    doc = Document()

    # Margin Setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Document Header & Title
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(0)
    p_top.paragraph_format.space_after = Pt(2)
    r_tag = p_top.add_run("MEGAZONECLOUD AI FULL STACK OFFERING REPORT")
    r_tag.font.name = FONT_ENG
    r_tag.font.size = Pt(8.5)
    r_tag.font.bold = True
    r_tag.font.color.rgb = COLOR_MZC_BLUE

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("퍼즐데이터 & 퀀텀AI AI Full Stack 접목 Use Case 및 기술 분석 보고서")
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_MZC_NAVY

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(16)
    r_meta = p_meta.add_run("문서 번호: 2026-08-26_DOC_AI_FULLSTACK_PARTNERSHIP | 작성 주체: MegazoneCloud AI Full Stack TF / ISSU 팀")
    r_meta.font.name = FONT_MAIN
    r_meta.font.size = Pt(8.5)
    r_meta.font.color.rgb = COLOR_MUTED

    # Section 1. 개요
    add_header_block(doc, "1. 개요 및 전략적 배경", "엔터프라이즈 소버린 플랫폼(Articul8) 기반 파트너 기술 융합")
    
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_before = Pt(2)
    p_intro.paragraph_format.space_after = Pt(8)
    r_intro = p_intro.add_run(
        "메가존클라우드의 AI Full Stack 전략은 Dell PowerEdge/NVIDIA H100 인프라(L1)와 Red Hat OpenShift AI(L2)를 바탕으로, "
        "핵심 오케스트레이션 소버린 플랫폼인 Articul8(A8) 위에 퍼즐데이터의 프로세스 인텔리전스 및 퀀텀AI의 올인원 AICC·멀티모달 AI를 "
        "유기적으로 결합하여 차별화된 산업별 자율 실행(Autonomous) Use Case를 완성합니다."
    )
    r_intro.font.name = FONT_MAIN
    r_intro.font.size = Pt(9.5)
    r_intro.font.color.rgb = COLOR_DARK_TEXT

    # 4-Layer 배치 표
    t_arch = doc.add_table(rows=5, cols=3)
    t_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_arch.autofit = False
    set_table_borders(t_arch)

    col_widths = [Cm(3.5), Cm(7.0), Cm(6.0)]
    headers = ["계층 (4-Layer)", "핵심 구성 솔루션", "주요 역할 및 특징"]
    for i, h in enumerate(headers):
        cell = t_arch.cell(0, i)
        cell.width = col_widths[i]
        set_cell_background(cell, HEX_DARK_NAVY)
        set_cell_margins(cell, top=140, bottom=140)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = FONT_MAIN
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = COLOR_WHITE

    arch_data = [
        ("L4: Sovereign Apps", "• Articul8 (A8) 자율 비즈니스 엔진\n• [Quantum AI] SOONi 올인원 AICC / 보이스봇\n• [Puzzle Data] ProDiscovery 프로세스 인텔리전스\n• 사내 특화 경영/운영 포털 UI", "엔터프라이즈 소버린 도메인 앱, 통합 AICC, 프로세스 대시보드"),
        ("L3: AI Middleware", "• MCP Server 표준 프로토콜 브로커\n• A8 Ingestion & Reasoning 엔진\n• [Quantum AI] Data2Vec / GraphRAG 멀티모달\n• [Puzzle Data] PiDi & ClickHouse 파이프라인\n• vLLM, LiteLLM, Redis 시맨틱 캐시", "지능형 에이전트 오케스트레이션, 비정형 멀티모달 처리, 초저지연 서빙"),
        ("L2: Platform/MLOps", "• Red Hat OpenShift AI (RHOAI)\n• Single Node OpenShift (SNO) 에어갭\n• Nutanix NKP / Kubernetes 런타임", "사내 폐쇄망 컨테이너 런타임, MLOps, 단일 테넌트 보안 격리"),
        ("L1: Infrastructure", "• Dell PowerEdge XE9680 / R760\n• NVIDIA H100 SXM / RTX 6000 Ada\n• Dell PowerScale F900 고속 스토리지\n• 국산 NPU (리벨리온, 퓨리오사AI)", "고성능 GPU 가속 연산, 초고속 AI 스토리지, 국산 NPU 복합 구성")
    ]

    for row_idx, data in enumerate(arch_data, start=1):
        for col_idx, text in enumerate(data):
            cell = t_arch.cell(row_idx, col_idx)
            cell.width = col_widths[col_idx]
            set_cell_margins(cell, top=100, bottom=100)
            if row_idx in [1, 2]:
                set_cell_background(cell, "F0F9FF" if row_idx == 1 else "FDF4FF")
            elif row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_DARK_TEXT
            if col_idx == 0:
                r.font.bold = True
                if row_idx in [1, 2]:
                    r.font.color.rgb = COLOR_MZC_BLUE if row_idx == 1 else COLOR_PURPLE

    # Section 2. 퍼즐데이터 x Articul8 Use Cases
    add_header_block(doc, "2. 퍼즐데이터 × Articul8 (A8) 연계 및 대체 3대 Use Cases", "프로세스 탐지(Puzzle Data)와 자율 실행(Articul8)의 결합")

    add_bullet(doc, "Use Case 1-1. A8 Autonomous Supply Chain & ERP Process Optimization (제조 SCM 자율 최적화)",
               "ERP/MES 전 공정 로그에서 원부자재 지연 위험을 감지하고, A8 Multi-Agent가 BOM 및 협력사 재고를 크로스체크하여 대체 생산 스케줄 수립 및 구매 발주 수정안을 자율 등록. (조치 리드타임 3일 ➔ 5분 단축, 공장 라인 정지 손실 제로화)")

    add_bullet(doc, "Use Case 1-2. A8 Process-Aware Financial Compliance & Forensic AML (금융 이상거래·AML 자율 감사)",
               "코어뱅킹 계좌 간 다단계 쪼개기 송금 및 비정상 승인 우회 경로를 실시간 시퀀스 그래프로 탐지하고, A8 온프레미스 에어갭 LLM이 금융감독원 제출용 의심거래보고서(STR) 초안 및 감사 브리프를 100% 자동 생성. (감사 소명 공수 85% 절감)")

    add_bullet(doc, "Use Case 1-3. A8 Self-Healing Enterprise ITOM & Incident Remediation (엔터프라이즈 IT 자율 치유)",
               "ServiceNow/Jira 및 APM 로그의 프로세스 흐름 병목을 감지하고, 최근 Git 배포 커밋과 매핑하여 근본 원인(RCA)을 규명한 후 사전 승인된 CI/CD 롤백 및 우회 라우팅을 A8이 자율 실행. (MTTR 70% 단축)")

    # Section 3. 퀀텀AI Use Cases
    add_header_block(doc, "3. 퀀텀AI (SOONi & Data2Vec) 3대 엔터프라이즈 Use Cases", "초저지연 음성인식(98%+) 및 올인원 Infra-Free AICC")

    add_bullet(doc, "Use Case 2-1. 금융권 에어갭(Air-gapped) 풀스택 AICC & 심사 자동화",
               "MZC Dell/NVIDIA 에어갭 인프라 위에 SOONi를 탑재하여 망분리 규제를 준수하고, 목소리 인증(5초) 기반 단순 업무 60% 무인 처리 및 IDOP(GraphRAG)를 통한 비정형 보험금/대출 심사 서류 자동화 구현. (AICC 도입비용 50% 절감, 구축기간 1~3개월 단축)")

    add_bullet(doc, "Use Case 2-2. 국방/방산/특수제조: 온프레미스 음성/VLM 기반 드론 및 로봇 PLC 자율 제어",
               "현장 작전 서버(RTX 6000 Ada / H100)에 퀀텀AI의 초저지연 음성인식과 Llama/sLLM, VLM을 탑재하여 1초 미만 자연어 음성 명령을 드론/로봇 PLC 제어 코드(JSON)로 실시간 생성. (통신 재밍 시 에어갭 독립 작동, 현대로템 프로젝트 연계)")

    add_bullet(doc, "Use Case 2-3. 스마트 헬스케어: 병의원 행정 자동화 및 AI 환자 모니터링 해피콜",
               "24/7 진료/검사 예약을 보이스봇이 무인 처리하고, 퇴원 환자 및 만성질환자 대상 정기 해피콜을 자율 수행하여 복약 상태와 통증 수치를 EMR에 자동 기록. (상담사 만족도 45% 상승, No-Show율 획기적 감축)")

    # Section 4. 부록: 기술적 장단점 평가 매트릭스
    add_header_block(doc, "4. [부록] 각 솔루션별 심층 기술적 장단점 평가 매트릭스", "아키텍처, 성능, 데이터 연동, 온프레미스 에어갭 관점 평가")

    # Table 1: Puzzle Data vs A8
    p_t1 = doc.add_paragraph()
    p_t1.paragraph_format.space_before = Pt(4)
    p_t1.paragraph_format.space_after = Pt(4)
    r_t1 = p_t1.add_run("■ 퍼즐데이터 (ProDiscovery) vs Articul8 (A8) 비교")
    r_t1.font.name = FONT_MAIN
    r_t1.font.size = Pt(10)
    r_t1.font.bold = True
    r_t1.font.color.rgb = COLOR_MZC_NAVY

    t_p1 = doc.add_table(rows=4, cols=3)
    t_p1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_p1.autofit = False
    set_table_borders(t_p1)

    t1_widths = [Cm(3.5), Cm(6.5), Cm(6.5)]
    t1_headers = ["구분 영역", "주요 기술적 강점 (Pros)", "한계 및 고려사항 (Cons)"]
    for i, h in enumerate(t1_headers):
        cell = t_p1.cell(0, i)
        cell.width = t1_widths[i]
        set_cell_background(cell, HEX_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_WHITE

    p1_data = [
        ("아키텍처 & 성능", "• ClickHouse 기반 대용량 로그 초고속 집계\n• 금융/공공 완전 폐쇄망 에어갭 동일 지원", "• 대용량 트래픽 시 ClickHouse 분산 샤딩 및 스토리지 튜닝 인력 필요"),
        ("AI & 분석 기능", "• PiDi 생성형 AI 도우미 (자연어 질의 대시보드 생성)\n• ML 기반 프로세스 지연 시뮬레이터", "• 실행 Worker 부재 (모니터링 중심, 시스템 자동 변경은 A8 연동 필수)"),
        ("데이터 연동", "• 유연한 모델링 (Case ID/Activity/Timestamp 매핑)\n• 국내 주요 ERP(SAP, 영림원) 스키마 노하우", "• 초기 수작업 모델링 3~4개월 소요 (A8 자동 인제스천 파이프라인으로 대체 권장)")
    ]

    for row_idx, data in enumerate(p1_data, start=1):
        for col_idx, text in enumerate(data):
            cell = t_p1.cell(row_idx, col_idx)
            cell.width = t1_widths[col_idx]
            set_cell_margins(cell, top=80, bottom=80)
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_DARK_TEXT
            if col_idx == 0:
                r.font.bold = True

    # Table 2: Quantum AI
    p_t2 = doc.add_paragraph()
    p_t2.paragraph_format.space_before = Pt(8)
    p_t2.paragraph_format.space_after = Pt(4)
    r_t2 = p_t2.add_run("■ 퀀텀AI (SOONi & Data2Vec) 기술적 장단점")
    r_t2.font.name = FONT_MAIN
    r_t2.font.size = Pt(10)
    r_t2.font.bold = True
    r_t2.font.color.rgb = COLOR_MZC_NAVY

    t_p2 = doc.add_table(rows=4, cols=3)
    t_p2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_p2.autofit = False
    set_table_borders(t_p2)

    for i, h in enumerate(t1_headers):
        cell = t_p2.cell(0, i)
        cell.width = t1_widths[i]
        set_cell_background(cell, HEX_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_WHITE

    p2_data = [
        ("시스템 구조 & 도입", "• Infra-Free 올인원 (CTI/PBX/IVR/STT/KMS 내장)\n• 착신 전환 1~3개월 도입 (구축비 50%↓)", "• 대형 하드웨어 PBX 기구축 고객은 소프트웨어 PBX 전환 시 설득 필요"),
        ("AI 원천 기술력", "• 초저지연 음성인식 (STT 98%+, 지연 1초 미만)\n• Data2Vec 및 GraphRAG 결합 환각 제로화", "• 오픈 도메인 복합 추론 시 파인튜닝 데이터셋 품질 의존도 존재"),
        ("온프레미스 & 운영", "• 완전 폐쇄망 지원 (국방/금융 규제 완벽 대응)\n• 1:N 상담 효율화 (상담사 1인이 동시 5명 응대)", "• 글로벌 플랫폼(Articul8) 대비 자동화된 배포 거버넌스 관리 도구 보완 필요")
    ]

    for row_idx, data in enumerate(p2_data, start=1):
        for col_idx, text in enumerate(data):
            cell = t_p2.cell(row_idx, col_idx)
            cell.width = t1_widths[col_idx]
            set_cell_margins(cell, top=80, bottom=80)
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_DARK_TEXT
            if col_idx == 0:
                r.font.bold = True

    # Output Save
    out_dir = r"c:\dev\antigravity-workspace\aifullstack\offering\docx"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "2026-08-26_Puzzle_Data_Quantum_AI_Fullstack_Usecases.docx")
    doc.save(out_path)
    print(f"✅ Word 문서 생성 완료: {out_path} ({os.path.getsize(out_path)/1024:.1f} KB)")

    docs_dir = r"c:\dev\antigravity-workspace\aifullstack\docs"
    docs_path = os.path.join(docs_dir, "2026-08-26_Puzzle_Data_Quantum_AI_Fullstack_Usecases.docx")
    doc.save(docs_path)
    print(f"✅ Word 문서 복사본 저장 완료: {docs_path} ({os.path.getsize(docs_path)/1024:.1f} KB)")

if __name__ == "__main__":
    build_document()
