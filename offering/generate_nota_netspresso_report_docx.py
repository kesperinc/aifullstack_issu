"""
Nota.ai NetsPresso 기술 심층 분석 및 MZC AI Full Stack 연계 전략 보고서 생성기 (DOCX)
- 한국어 및 표준 엔터프라이즈 스타일 가이드라인 준수
- Plan-Code-Doc 트라이어드 준수
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# 폰트 & 색상 상수
FONT_MAIN = "Malgun Gothic"
FONT_TITLE = "Malgun Gothic"

COLOR_MZC_BLUE = RGBColor(0x00, 0xAB, 0xF0)
COLOR_NAVY_DARK = RGBColor(0x0F, 0x17, 0x2A)
COLOR_TEXT_MAIN = RGBColor(0x22, 0x22, 0x22)
COLOR_TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_GOLD = RGBColor(0xD4, 0xAF, 0x37)
COLOR_NVIDIA_GREEN = RGBColor(0x76, 0xB9, 0x00)
COLOR_PURPLE = RGBColor(0x9D, 0x4E, 0xDD)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_BG_DARK_NAVY = "0F172A"
HEX_BG_LIGHT_GRAY = "F8FAFC"
HEX_BORDER = "E2E8F0"


def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="none"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="none"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))


def add_section_header(doc, sec_num, title_text, subtitle_text=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r_num = p.add_run(f"[{sec_num}] ")
    r_num.font.name = FONT_TITLE
    r_num.font.size = Pt(14)
    r_num.bold = True
    r_num.font.color.rgb = COLOR_PURPLE

    r_title = p.add_run(title_text)
    r_title.font.name = FONT_TITLE
    r_title.font.size = Pt(14)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    if subtitle_text:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(10)
        p_sub.paragraph_format.keep_with_next = True
        r_sub = p_sub.add_run(subtitle_text)
        r_sub.font.name = FONT_MAIN
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = COLOR_TEXT_MUTED
        r_sub.italic = True


def add_sub_header(doc, title_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title_text)
    r.font.name = FONT_TITLE
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = COLOR_NAVY_DARK


def add_narrative_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    r.font.name = FONT_MAIN
    r.font.size = Pt(9.2)
    r.font.color.rgb = COLOR_TEXT_MAIN
    return p


def add_callout_box(doc, title, bullets, border_color_hex="9D4EDD", bg_color_hex="FDF4FF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Cm(17.0)
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    set_cell_shading(cell, bg_color_hex)
    
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color_hex}"/>
        <w:top w:val="none"/>
        <w:right w:val="none"/>
        <w:bottom w:val="none"/>
    </w:tcBorders>
    '''
    cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f"💡 {title}")
    r_t.font.name = FONT_TITLE
    r_t.font.size = Pt(9.8)
    r_t.bold = True
    r_t.font.color.rgb = COLOR_NAVY_DARK

    for b in bullets:
        p_b = cell.add_paragraph()
        p_b.paragraph_format.space_before = Pt(2)
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Cm(0.3)
        p_b.paragraph_format.line_spacing = 1.25
        r_b = p_b.add_run(b)
        r_b.font.name = FONT_MAIN
        r_b.font.size = Pt(8.6)
        r_b.font.color.rgb = COLOR_TEXT_MAIN

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)


def build_nota_report(output_path):
    doc = Document()

    # 페이지 여백 설정 (A4)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
        # 헤더 / 푸터
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Nota.ai NetsPresso 기술 분석 & MZC AI Full Stack 연계 전략 | Confidential")
        hrun.font.name = FONT_MAIN
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = COLOR_TEXT_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("© 2026 MEGAZONECLOUD Corp. All Rights Reserved. | Executive Technical Report")
        frun.font.name = FONT_MAIN
        frun.font.size = Pt(8)
        frun.font.color.rgb = COLOR_TEXT_MUTED

    # ── [표지 메타 영역] ──────────────────────────────────────────────
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(4)
    p_badge.paragraph_format.space_after = Pt(4)
    r_b = p_badge.add_run("⚡ ISV TECHNOLOGY DEEP-DIVE & INTEGRATION STRATEGY REPORT")
    r_b.font.name = FONT_MAIN
    r_b.font.size = Pt(9)
    r_b.bold = True
    r_b.font.color.rgb = COLOR_PURPLE

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("Nota.ai NetsPresso 플랫폼 기술 심층 분석\n& MZC AI Full Stack 연계 오퍼링 전략")
    r_title.font.name = FONT_TITLE
    r_title.font.size = Pt(17)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run(
        "하드웨어 인지 경량화·양자화(AutoRound/Mixed Precision) 기술 분석, Red Hat OpenShift AI(RHOAI)와의 기능 비교 및 상호 보완 파이프라인, "
        "그리고 MZC AI Full Stack 3대 단독/복합 비즈니스 제안 시나리오 검토"
    )
    r_sub.font.name = FONT_MAIN
    r_sub.font.size = Pt(9.5)
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    # ── [목차 요약] ───────────────────────────────────────────────────
    toc_items = [
        "1. Nota.ai NetsPresso 플랫폼 핵심 기술 요약 (경량화, 양자화, 디바이스 팜)",
        "2. MZC AI Full Stack 편입 시 핵심 검토 및 사전 점검 과제 (Gap Analysis)",
        "3. Red Hat OpenShift AI (RHOAI)와의 기능 비교 및 상호 보완 공존 방안",
        "4. MZC AI Full Stack 내 단독/복합 솔루션 제안 시나리오 (3대 Use Case)",
        "5. 종합 결론 및 MZC 파트너십 추진 제언"
    ]
    for item in toc_items:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(1)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.left_indent = Cm(0.4)
        r = p_t.add_run(f"• {item}")
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.8)
        r.font.color.rgb = COLOR_TEXT_MAIN
        r.bold = True

    # ── [1. NetsPresso 핵심 기술 요약] ─────────────────────────────────
    add_section_header(doc, "1", "Nota.ai NetsPresso 플랫폼 핵심 기술 요약",
                       "하드웨어 인지 AI 모델 최적화, 고급 양자화 및 자동 컴파일 플랫폼")

    add_narrative_paragraph(
        doc,
        "NetsPresso(넷스프레소)는 AI 기술 기업 노타(Nota.ai)가 개발한 '하드웨어 인지 AI 모델 최적화 및 경량화 자동화 플랫폼'입니다. "
        "AI 모델을 다양한 타깃 디바이스(NVIDIA GPU, Intel CPU, ARM, NPU, 엣지 칩셋)에 배포할 때 발생하는 극심한 수작업 튜닝 병목(Manual Tuning Fatigue)을 제거하고, "
        "모델 압축·양자화부터 타깃 칩셋 컴파일 및 실측 벤치마킹까지 전 과정을 자율화합니다."
    )

    t_tech = doc.add_table(rows=4, cols=2)
    t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tech.autofit = False
    set_table_borders(t_tech)

    tech_headers = [("핵심 기술 영역", Cm(4.5)), ("상세 기술 내용 및 엔터프라이즈 차별점", Cm(12.5))]
    for col_idx, (h_text, w_val) in enumerate(tech_headers):
        cell = t_tech.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    tech_data = [
        ("1. 하드웨어 인지 경량화\n(Hardware-Aware Pruning)",
         "• 150+ 최적화 패턴 라이브러리: 대상 칩셋(GPU/NPU/CPU)의 연산 유닛 및 메모리 대역폭 특성에 맞춰 최적의 압축 기법 자동 매칭\n"
         "• 구조적 가지치기 (Structured Pruning): 실제 가속기에서 하드웨어 가속이 가능한 형태로 채널/필터 단위 압축\n"
         "• 필터 분해 & 레이어 융합 (Layer Fusion): 트랜스포머 어텐션 블록 및 합성곱 연산 경량화"),
        ("2. 지능형 양자화 & LLM\n(Advanced Quantization)",
         "• 혼합 정밀도 양자화 (Mixed Precision): 레이어 민감도를 자동 평가하여 FP16, INT8, INT4, FP8 복합 배정\n"
         "• LLM 특화 양자화 (AutoRound): Llama 3, Mistral, Qwen 등 거대 모델의 정확도 손실을 1% 미만으로 억제하며 가중치 압축\n"
         "• HAQA (Hardware-Aware Quantization Agent): LLM 기반 에이전트가 타깃 칩셋 제약에 맞춰 하이퍼파라미터 자율 탐색"),
        ("3. Device Farm & 자동 컴파일\n(Benchmarking & Toolchain)",
         "• 물리 칩셋 실측 (Device Farm): 단순 시뮬레이션이 아닌 실제 물리 칩셋에서 Latency, Memory Footprint, 전력 소모 실측\n"
         "• 원클릭 타깃 툴체인 변환: ONNX, TensorRT, OpenVINO, TFLite, SNPE, QNN 등으로 자동 컴파일\n"
         "• 노코드/로우코드 파이프라인: GUI 웹 포털 및 파이썬 SDK/API를 통한 CI/CD MLOps 파이프라인 자동화")
    ]

    for row_idx, (t_name, t_desc) in enumerate(tech_data, start=1):
        row_cells = t_tech.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([t_name, t_desc], [Cm(4.5), Cm(12.5)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.2)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_PURPLE

    # ── [2. AI Full Stack 편입 시 사전 검토 사항] ──────────────────────
    add_section_header(doc, "2", "MZC AI Full Stack 편입 시 핵심 검토 및 사전 점검 과제",
                       "성공적인 엔터프라이즈 오퍼링 편입을 위한 4대 Gap 및 Risk 분석")

    add_narrative_paragraph(
        doc,
        "메가존클라우드의 엔터프라이즈 AI Full Stack 프레임워크에 NetsPresso를 공식 솔루션 컴포넌트로 결합하기 위해 다음 4대 핵심 항목에 대한 심층 기술 및 비즈니스 검토가 선행되어야 합니다."
    )

    gaps = [
        ("• 1. 대형 파운데이션 모델(LLM / MoE / VLM) 지원 성숙도 검증: ",
         "NetsPresso의 뿌리는 컴퓨터 비전(CNN, YOLO) 및 엣지 AI에 있으므로, 엔터프라이즈 고객이 주로 요구하는 32B~70B 대형 언어 모델(Qwen2.5-Coder, Llama 3.3), MoE 구조(Mixtral), 멀티모달(VLM)에 대한 압축·양자화 지원 안정성과 실측 처리 성능(TPS) 검증이 필수적입니다."),
        ("• 2. 온프레미스 완전 폐쇄망(Air-Gap) 설치 배포 지원 여부: ",
         "금융·공공·제조 대기업 고객은 외부 인터넷과 차단된 소버린 폐쇄망 환경을 요구합니다. 현재 NetsPresso의 주요 서비스 형태가 클라우드 SaaS인지, 또는 사내 쿠버네티스(OpenShift/Nutanix) 상에 온프레미스 프라이빗 패키지(컨테이너 Appliance)로 배포 가능한지 및 오프라인 라이선스 인증 방식을 확인해야 합니다."),
        ("• 3. 엔터프라이즈 MLOps 및 서빙 엔진 연동성 (Interoperability): ",
         "MZC 스택의 핵심 서빙 프레임워크인 vLLM, Triton Inference Server, KServe, DataRobot 및 LangChain/Articul8과의 표준 API(REST/gRPC/CLI/SDK) 연계 파이프라인 호환성을 사전 검증해야 합니다."),
        ("• 4. 상용 라이선스 체계 및 SI 비즈니스 마진 구조: ",
         "디바이스당 과금(Per-device), 모델당 과금(Per-model), 또는 엔터프라이즈 연간 서브스크립션(Annual License) 등 노타의 가격 정책을 확인하고, MZC의 H/W(Dell) + 가상화(Nutanix/Red Hat) 패키지 번들링 시 경쟁력 있는 고객 TCO와 당사 SI 마진율을 산정해야 합니다.")
    ]

    for g_title, g_desc in gaps:
        p_g = doc.add_paragraph()
        p_g.paragraph_format.space_before = Pt(2)
        p_g.paragraph_format.space_after = Pt(4)
        p_g.paragraph_format.left_indent = Cm(0.4)
        r_gt = p_g.add_run(g_title)
        r_gt.font.name = FONT_MAIN
        r_gt.font.size = Pt(8.8)
        r_gt.bold = True
        r_gt.font.color.rgb = COLOR_PURPLE

        r_gd = p_g.add_run(g_desc)
        r_gd.font.name = FONT_MAIN
        r_gd.font.size = Pt(8.8)
        r_gd.font.color.rgb = COLOR_TEXT_MAIN

    # ── [3. RHOAI와의 기능 비교 및 공존 방안] ──────────────────────────
    add_section_header(doc, "3", "Red Hat OpenShift AI (RHOAI)와의 기능 비교 및 상호 보완 공존 방안",
                       "엔터프라이즈 인프라 플랫폼(RHOAI)과 전문 모델 경량화 엔진(NetsPresso)의 시너지")

    add_narrative_paragraph(
        doc,
        "Red Hat OpenShift AI(RHOAI)와 NetsPresso는 직접적인 경쟁 관계가 아니며, 엔터프라이즈 MLOps 플랫폼(RHOAI)과 전문 모델 최적화 엔진(NetsPresso)으로서 완벽한 상호 보완(Complementary) 관계를 형성합니다."
    )

    t_rhoai = doc.add_table(rows=5, cols=3)
    t_rhoai.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_rhoai.autofit = False
    set_table_borders(t_rhoai)

    rhoai_headers = [("비교 항목", Cm(3.2)), ("Red Hat OpenShift AI (RHOAI)", Cm(6.8)), ("Nota.ai NetsPresso", Cm(7.0))]
    for col_idx, (h_text, w_val) in enumerate(rhoai_headers):
        cell = t_rhoai.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    rhoai_data = [
        ("핵심 역할 & 계층", "Layer 02 / 03 코어 MLOps & 서빙 플랫폼\n(엔터프라이즈 K8s 기반 전주기 AI 인프라 오케스트레이션)", "Layer 03 특화 모델 최적화 & 컴파일러 엔진\n(타깃 하드웨어 인지 모델 경량화/양자화/컴파일 전문 도구)"),
        ("주요 핵심 기능", "• 분산 학습 (Ray, PyTorch)\n• 파이프라인 자동화 (Kubeflow, Elyra)\n• 엔터프라이즈 서빙 (KServe, vLLM, OpenVINO)\n• vGPU/MIG 스케줄링 & 보안 거버넌스", "• 하드웨어 인지 구조적 가지치기 (Pruning)\n• LLM AutoRound / Mixed Precision 양자화\n• 150+ 타깃 칩셋 맞춤 IR 변환 & 컴파일\n• Device Farm 실측 벤치마킹"),
        ("모델 경량화 깊이", "표준 프레임워크 수준 (일반 Quantization 지원)", "하드웨어 맞춤형 초정밀 경량화 & 최적화 전문"),
        ("인프라 제어 범위", "전사 클러스터, 멀티 테넌시, 스토리지, 보안 제어", "단일 모델 아티팩트 압축 및 컴파일에 집중")
    ]

    for row_idx, r_data in enumerate(rhoai_data, start=1):
        row_cells = t_rhoai.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(6.8), Cm(7.0)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.0)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY_DARK

    add_callout_box(
        doc,
        "RHOAI + NetsPresso 통합 시너지 파이프라인 아키텍처",
        [
            "• [연동 흐름]: RHOAI Workbench(학습) ➔ NetsPresso Plugin(자동 양자화/컴파일) ➔ RHOAI KServe/vLLM(엔터프라이즈 서빙 배포)",
            "• [도입 효과]: 개발자는 RHOAI에서 단일 워크플로우로 작업하며, 백엔드에서 NetsPresso가 모델을 1/2~1/4로 압축하여 단일 GPU 서빙 TPS를 3배 이상 향상시킵니다."
        ],
        border_color_hex="9D4EDD",
        bg_color_hex="FDF4FF"
    )

    # ── [4. AI Full Stack 3대 제안 시나리오] ────────────────────────────
    add_section_header(doc, "4", "MZC AI Full Stack 내 단독/복합 솔루션 제안 시나리오",
                       "고객 페인포인트 해결을 위한 3대 대표 비즈니스 오퍼링 패키지")

    add_narrative_paragraph(
        doc,
        "MZC AI Full Stack의 포트폴리오 관점에서 NetsPresso를 단독 또는 복합 패키지로 제안할 수 있는 3가지 대표적인 고객 적용 시나리오입니다."
    )

    scenarios = [
        ("시나리오 1. 스마트팩토리 엣지 AI 패키지 (Vision AI & IoT)",
         "• 적용 분야: 제조 공장 실시간 불량 검출(Vision AI), 지능형 CCTV 관제, 물류 로봇\n"
         "• 고객 페인포인트: 고성능 비전 AI 모델이 현장 엣지 칩셋(NVIDIA Jetson, Intel NUC)에서 메모리 초과 및 발열로 가동 불가\n"
         "• MZC 오퍼링: Dell 엣지 서버 + NetsPresso 비전 경량화(크기 1/5, 속도 3~5배 향상) + 온프레미스 통합 관제 시스템 구축"),
        ("시나리오 2. 소버린 sLLM 경량 서빙 패키지 (단일 GPU 고성능 RAG)",
         "• 적용 분야: 금융·공공 사내 Q&A RAG, 온프레미스 사내 코딩 에이전트\n"
         "• 고객 페인포인트: 8-GPU 대형 서버(10억+) 도입 예산 부담, 1~2 GPU 서버로 고속 추론을 구현해야 하는 예산 제약\n"
         "• MZC 오퍼링: Dell PowerEdge R760(2x RTX 6000 Ada / L40S) + NetsPresso AutoRound INT4 양자화 + vLLM 결합으로 단일 가속기에서 60+ TPS 서빙 달성 (인프라 도입 TCO 60% 절감)"),
        ("시나리오 3. 소버린 NPU 멀티 벤더 전환 패키지 (이기종 가속기)",
         "• 적용 분야: 공공·국방 국산 AI 반도체 실증 사업, 엔터프라이즈 엔비디아 GPU 종속 탈피\n"
         "• 고객 페인포인트: 엔비디아 GPU 수급난 및 과도한 비용, 국산 NPU로의 포팅 시 모델 호환성 및 성능 저하 우려\n"
         "• MZC 오퍼링: 리벨리온, 퓨리오사AI NPU + NetsPresso Device Farm 자동 컴파일로 이기종 칩셋 원클릭 마이그레이션 실현")
    ]

    for s_title, s_desc in scenarios:
        add_sub_header(doc, f"■ {s_title}")
        p_sc = doc.add_paragraph()
        p_sc.paragraph_format.space_before = Pt(1)
        p_sc.paragraph_format.space_after = Pt(6)
        p_sc.paragraph_format.left_indent = Cm(0.4)
        r_sc = p_sc.add_run(s_desc)
        r_sc.font.name = FONT_MAIN
        r_sc.font.size = Pt(8.5)
        r_sc.font.color.rgb = COLOR_TEXT_MAIN

    # ── [5. 종합 결론 및 향후 기술 검증(PoC) 제언] ───────────────────────
    add_section_header(doc, "5", "종합 결론 및 향후 기술 검증(PoC) 제언",
                       "Layer 03 구성요소 편입 검토 가능성 및 사전 기술 검증 선행 과제")

    add_narrative_paragraph(
        doc,
        "Nota.ai의 NetsPresso는 동일 인프라 환경에서 모델 경량화 및 추론 서빙 처리량(TPS)을 최대 3배 이상 극대화할 수 있는 잠재력을 보유하고 있어, "
        "향후 MZC AI Full Stack의 'Layer 03 (LLM Model + Tool Layer)' 내 압축/최적화 구성요소 중 하나로 편입을 검토할 수 있는 솔루션입니다."
    )

    add_narrative_paragraph(
        doc,
        "그러나 즉각적인 상용 도입이나 필수 탑재를 결정하기에 앞서, 엔터프라이즈 거대 LLM(32B~70B)에 대한 실증 성능, 온프레미스 폐쇄망(Air-Gap) 배포 안정성, "
        "그리고 상용 라이선스 비즈니스 모델에 대한 사전 기술 검증(PoC)이 반드시 선행되어야 하는 '향후 필요 시 검토 대상(Candidate)' 솔루션으로 포지셔닝합니다."
    )

    conclusions = [
        ("• 1. 사내 랩 기반 사전 기술 검증(PoC) 우선 추진: ", "당사 사내 AI 데모 랩(Dell PowerEdge R760)에서 엔터프라이즈 코딩 모델(Qwen2.5-Coder) 및 사내 sLLM(Solar)을 대상으로 실제 정확도 보존율과 TPS 향상률을 정밀 실측 검증"),
        ("• 2. 온프레미스 완전 폐쇄망(Air-Gap) 패키징 검토: ", "금융·공공 고객 환경에 필수적인 독립 컨테이너 Appliance 배포 및 오프라인 라이선스 인증 체계의 기술적 성숙도 점검"),
        ("• 3. 점진적 파트너십 및 옵션형 솔루션 검토: ", "필수 기본 탑재가 아닌, 엣지 AI나 단일 GPU sLLM 최적화 등 특정 고객 요구 시 선택적으로 제안할 수 있는 '선택형 확장 옵션(Optional Add-on)'으로 파트너십 조건 협의")
    ]

    for c_title, c_desc in conclusions:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_before = Pt(2)
        p_c.paragraph_format.space_after = Pt(4)
        p_c.paragraph_format.left_indent = Cm(0.4)
        r_ct = p_c.add_run(c_title)
        r_ct.font.name = FONT_MAIN
        r_ct.font.size = Pt(8.8)
        r_ct.bold = True
        r_ct.font.color.rgb = COLOR_PURPLE

        r_cd = p_c.add_run(c_desc)
        r_cd.font.name = FONT_MAIN
        r_cd.font.size = Pt(8.8)
        r_cd.font.color.rgb = COLOR_TEXT_MAIN

    # 저장
    doc.save(str(output_path))
    print(f"✅ Nota.ai NetsPresso 보고서 DOCX 생성 완료: {output_path}")
    print(f"   파일 크기: {output_path.stat().st_size / 1024:.1f} KB")


def main():
    base_dir = Path(__file__).resolve().parent
    docx_dir = base_dir / "docx"
    docx_dir.mkdir(exist_ok=True)
    
    out_dated = docx_dir / "2026-08-20_Nota_AI_NetsPresso_Fullstack_Analysis_Report.docx"
    out_docs = base_dir.parent / "docs" / "2026-08-20_Nota_AI_NetsPresso_Fullstack_Analysis_Report.docx"

    build_nota_report(out_dated)
    build_nota_report(out_docs)


if __name__ == "__main__":
    main()
