"""
Liqid Composable GPU & Memory Pooling 기술 분석 및 MZC AI Full Stack / Private Cloud 총판 타당성 보고서 생성기 (DOCX)
- ISSU 조직 배경 (Sales/Presales, 파트너 H/W 판매, 테스트 인프라 부재 제약) 의무 명시
- 표준 스타일 가이드라인 준수
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

FONT_MAIN = "Malgun Gothic"
FONT_TITLE = "Malgun Gothic"

COLOR_CYAN = RGBColor(0x00, 0x97, 0xA7)
COLOR_NAVY_DARK = RGBColor(0x0F, 0x17, 0x2A)
COLOR_TEXT_MAIN = RGBColor(0x22, 0x22, 0x22)
COLOR_TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_GOLD = RGBColor(0xD4, 0xAF, 0x37)
COLOR_PURPLE = RGBColor(0x9D, 0x4E, 0xDD)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEX_BG_DARK_NAVY = "0F172A"
HEX_BG_LIGHT_GRAY = "F8FAFC"


def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:left w:val="none"/>
        <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:right w:val="none"/>
        <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))


def add_section_header(doc, sec_num, title_text, subtitle_text=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r_num = p.add_run(f"[{sec_num}] ")
    r_num.font.name = FONT_TITLE
    r_num.font.size = Pt(13)
    r_num.bold = True
    r_num.font.color.rgb = COLOR_CYAN

    r_title = p.add_run(title_text)
    r_title.font.name = FONT_TITLE
    r_title.font.size = Pt(13)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    if subtitle_text:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(8)
        p_sub.paragraph_format.keep_with_next = True
        r_sub = p_sub.add_run(subtitle_text)
        r_sub.font.name = FONT_MAIN
        r_sub.font.size = Pt(9.2)
        r_sub.font.color.rgb = COLOR_TEXT_MUTED
        r_sub.italic = True


def add_sub_header(doc, title_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title_text)
    r.font.name = FONT_TITLE
    r.font.size = Pt(10.5)
    r.bold = True
    r.font.color.rgb = COLOR_NAVY_DARK


def add_narrative_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    r.font.name = FONT_MAIN
    r.font.size = Pt(9.2)
    r.font.color.rgb = COLOR_TEXT_MAIN
    return p


def add_callout_box(doc, title, bullets, border_color_hex="0097A7", bg_color_hex="F0FDFA"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Cm(17.0)
    set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
    set_cell_shading(cell, bg_color_hex)
    
    borders_xml = f'''
    <w:tcBorders {nsdecls("w")}>
        <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color_hex}"/>
        <w:top w:val="none"/>
        <w:right w:val="none"/>
        <w:bottom w:val="none"/>
    </w:tcBorders>
    '''
    cell._tc.get_or_add_tcPr().append(parse_xml(borders_xml))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f"🏛️ {title}")
    r_t.font.name = FONT_TITLE
    r_t.font.size = Pt(9.8)
    r_t.bold = True
    r_t.font.color.rgb = COLOR_NAVY_DARK

    for b in bullets:
        p_b = cell.add_paragraph()
        p_b.paragraph_format.space_before = Pt(2)
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Cm(0.3)
        p_b.paragraph_format.line_spacing = 1.25
        r_b = p_b.add_run(b)
        r_b.font.name = FONT_MAIN
        r_b.font.size = Pt(8.6)
        r_b.font.color.rgb = COLOR_TEXT_MAIN

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)


def build_liqid_report(output_path):
    doc = Document()

    # 여백 설정 (A4)
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Liqid CDI 기술 분석 & MZC AI Full Stack / Private Cloud 총판 검토 | Confidential")
        hrun.font.name = FONT_MAIN
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = COLOR_TEXT_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("© 2026 MEGAZONECLOUD Corp. All Rights Reserved. | ISSU Strategic Review")
        frun.font.name = FONT_MAIN
        frun.font.size = Pt(8)
        frun.font.color.rgb = COLOR_TEXT_MUTED

    # ── [표지 메타] ───────────────────────────────────────────────────
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(4)
    p_badge.paragraph_format.space_after = Pt(4)
    r_b = p_badge.add_run("⚡ HARDWARE-LEVEL CDI & PRIVATE CLOUD FEASIBILITY REPORT")
    r_b.font.name = FONT_MAIN
    r_b.font.size = Pt(9)
    r_b.bold = True
    r_b.font.color.rgb = COLOR_CYAN

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("Liqid Composable GPU & Memory Pooling 기술 분석\n& MZC AI Full Stack 연계 / 총판 사업 타당성 검토")
    r_title.font.name = FONT_TITLE
    r_title.font.size = Pt(16)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run(
        "소프트웨어 정의 GPU/DRAM 풀링(CDI) 기술 심층 분석, NVIDIA NVL-72 대비 7배 코어 비용 절감 및 90% KV Cache 가속 효과 검토, "
        "그리고 ISSU(Solution Sales & Presales) 조직 특성 및 인프라 제약 기반의 Private Cloud 파트너 세일즈 오퍼링 전략"
    )
    r_sub.font.name = FONT_MAIN
    r_sub.font.size = Pt(9.5)
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    # ── [1. ISSU 조직 배경 및 분석 전제] ──────────────────────────────
    add_section_header(doc, "1", "ISSU 조직 배경 및 분석 기본 전제",
                       "Solution Sales / Presales 조직 특성, 파트너 판매 구조 및 인프라 제약 조건")

    add_callout_box(
        doc,
        "ISSU(Integrated Solution Sales Unit) 조직 정의 및 비즈니스 모델",
        [
            "• [조직의 본질]: ISSU는 메가존클라우드의 솔루션 세일즈(Solution Sales), 기술 영업 및 프리세일즈(Presales) 중심 조직입니다.",
            "• [하드웨어 솔루션의 성격]: H/W를 직접 제조하거나 사내에 장비를 보유하지 않으며, 당사는 '세일즈 오퍼링(Offering)' 및 아키텍처 컨설팅 패키지 형태로 포트폴리오를 보유합니다 (현재 AIDC 사업에 대하여 검토 중이나, 구체화된 내용은 없음).",
            "• [판매 및 공급 체계]: 실제 하드웨어의 조달, 납품, 유지보수 및 구축은 델(Dell)과 엔비디아(NVIDIA) 파트너를 통하여 이루어짐.",
            "• [현재의 인프라 제약]: 현재 ISSU 내부에는 하드웨어를 직접 설치하여 테스트할 수 있는 자체 실증 Infra가 부재한 상태입니다."
        ],
        border_color_hex="D4AF37",
        bg_color_hex="FFFDF5"
    )

    # ── [2. Liqid 핵심 기술 요약] ─────────────────────────────────────
    add_section_header(doc, "2", "Liqid(리키드) 핵심 기술 요약",
                       "PCIe Gen5 기반 GPU & DRAM 분리 및 소프트웨어 정의 동적 풀링 (CDI)")

    add_narrative_paragraph(
        doc,
        "Liqid는 서버 내부에 고정되어 있던 GPU와 메모리(DRAM)를 섀시 단위로 분리(Disaggregation)하고, "
        "초고속 PCIe Gen5 x16 패브릭과 소프트웨어(Liqid Matrix)를 통해 필요한 만큼 동적으로 묶어주는 Composable Disaggregated Infrastructure (CDI) 플랫폼입니다."
    )

    t_tech = doc.add_table(rows=4, cols=2)
    t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tech.autofit = False
    set_table_borders(t_tech)

    tech_headers = [("핵심 기술 영역", Cm(4.5)), ("상세 기술 내용 및 엔터프라이즈 가치", Cm(12.5))]
    for col_idx, (h_text, w_val) in enumerate(tech_headers):
        cell = t_tech.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    tech_data = [
        ("1. PCIe Gen5 패브릭 기반\nGPU/DRAM 풀링",
         "• Liqid Expansion Chassis: 섀시당 10x GPU(최대 30x GPU 풀링), 섀시당 40TB DRAM(최대 200TB 풀링)\n"
         "• 기존 서버 자산 100% 보호: Dell, HPE, Cisco, Supermicro 등 기존 1U~4U 서버에 PCIe HIC 카드만 장착하여 연결\n"
         "• 초저지연 버스 전송: 네트워크 스위치 오버헤드 없이 네이티브 PCIe Gen5 속도로 자원 공유"),
        ("2. Liqid Matrix SW\n동적 자원 오케스트레이션",
         "• 소프트웨어 정의 자원 할당: 물리적 케이블 재연결 없이 GUI/API로 서버에 GPU/메모리를 1초 만에 증설/회수\n"
         "• Kubernetes & Slurm 플러그인: 컨테이너 워크로드 요구량에 맞춰 GPU 개수 및 DRAM 크기를 자율 스케줄링\n"
         "• 자원 과다 프로비저닝(Overprovisioning) 75% 절감 및 서버 수 75% 감소"),
        ("3. 대용량 DRAM 풀링 &\n90% KV Cache 가속",
         "• 최대 200TB 초대형 DRAM 풀: 단일 GPU 메모리 한계를 극복하는 거대 메모리 공유 풀 제공\n"
         "• 90% KV Cache Hit Rate: LLM 긴 문맥 추론 시 토큰 재연산(Recompute)을 제거하여 추론 속도 7배 향상\n"
         "• Tokens per Watt 2배 극대화 및 코어당 전력 소모 5배 절감")
    ]

    for row_idx, (t_name, t_desc) in enumerate(tech_data, start=1):
        row_cells = t_tech.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([t_name, t_desc], [Cm(4.5), Cm(12.5)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.2)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_CYAN

    # ── [3. NVL-72 vs Liqid PCIe-72 비교] ─────────────────────────────
    add_section_header(doc, "3", "엔비디아 NVL-72 vs Liqid PCIe-72 심층 비교 분석",
                       "독점 수랭 턴키 팩토리 대비 공랭 데이터센터 TCO 최적화 비교")

    t_comp = doc.add_table(rows=7, cols=3)
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_comp.autofit = False
    set_table_borders(t_comp)

    comp_headers = [("비교 지표", Cm(3.2)), ("NVIDIA NVL-72 (독점 턴키 팩토리)", Cm(6.8)), ("Liqid PCIe-72 (컴포저블 풀링)", Cm(7.0))]
    for col_idx, (h_text, w_val) in enumerate(comp_headers):
        cell = t_comp.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    comp_data = [
        ("GPU 및 아키텍처", "72x B200 NVLink 일체형 (고정형 랙)", "72x RTX 6000 Ada / L40S PCIe (동적 풀링)"),
        ("코어당 단가 / 비용", "~$30,000/GPU (Core당 $2.21)", "~$7,500/GPU (Core당 $0.31, 약 7배 우수)"),
        ("전력 소모 (Power)", "~150 kW (특수 수랭식 랙 필수)", "~60 kW (일반 공랭식 랙 지원, 5배 우수)"),
        ("데이터센터 요구조건", "수랭 설비 완비된 신축 데이터센터 전용", "기존 일반 기업 전산실/데이터센터 즉시 설치"),
        ("인프라 유연성", "정적 구성 (Static, 변경 불가)", "S/W 기반 동적 자원 재할당 (Composable)"),
        ("LLM 추론 경제성", "대규모 학습 중심, 추론 시 과도한 TCO", "대용량 DRAM KV Cache 연동 시 최상위 Tokens/$")
    ]

    for row_idx, r_data in enumerate(comp_data, start=1):
        row_cells = t_comp.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(6.8), Cm(7.0)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.0)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY_DARK

    # ── [4. 엔터프라이즈 HA 안정성 및 PCIe/NTB 구조적 제약 · 성능 비교] ─────
    add_section_header(doc, "4", "엔터프라이즈 고가용성(HA) 및 PCIe/NTB 구조적 제약 · 성능 비교",
                       "고가용성 리스크, 랙 스케일 제약 및 NVLink / InfiniBand 대비 대역폭·지연속도 비교")

    add_sub_header(doc, "■ 1. 고가용성(HA) 및 단일 장애점(SPOF) 리스크 평가")
    add_narrative_paragraph(
        doc,
        "• Expansion Chassis 레벨 SPOF: Liqid 섀시 1대에 최대 30개 GPU 및 다수 호스트가 연결되므로, 섀시 전원/PCIe 스위치 고장 시 연결된 모든 서버의 AI 서비스가 동시 중단되는 장애 확산 위험이 존재합니다.\n"
        "• PCIe 링크 페일오버의 한계: 이더넷/InfiniBand(LACP/BGP)와 달리 PCIe 하드웨어 버스 단절 시 OS 커널 패닉(PCIe AER) 및 CUDA 컨텍스트 강제 종료가 발생할 수 있어 무중단 하드웨어 페일오버에 태생적 한계가 있습니다.\n"
        "• 컨트롤 플레인(Liqid Matrix) 이중화: Matrix 오케스트레이터의 Active-Standby 이중화 및 메타데이터 동기화 아키텍처가 필수 검증 항목입니다."
    )

    add_sub_header(doc, "■ 2. PCIe & NTB(Non-Transparent Bridge) 확장성 제약")
    add_narrative_paragraph(
        doc,
        "• 초기 세팅 후 동적 토폴로지 변경의 제약: PCIe는 기본적으로 트리(Tree) 구조의 정적 버스 열거(Bus Enumeration)에 의존합니다. 런타임 동적 재할당(Hot-plug) 시 OS 커널 드라이버 및 CUDA 런타임 리로드 간헐적 불안정성이 발생할 수 있습니다.\n"
        "• 랙 스케일(Rack-Scale) 한계: 단일 섀시(최대 30x GPU) 및 단일 Pod(최대 16대 서버) 범위 내에서는 효과적이나, 멀티 랙(Multi-Rack) 이상의 대규모 클러스터로 확장 시에는 결국 InfiniBand/RoCEv2 네트워크 브릿징이 강제됩니다."
    )

    add_sub_header(doc, "■ 3. 패브릭 성능 비교 매트릭스: NVLink vs InfiniBand vs Liqid PCIe (NTB)")
    t_fab = doc.add_table(rows=4, cols=4)
    t_fab.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_fab.autofit = False
    set_table_borders(t_fab)

    fab_headers = [("패브릭 규격", Cm(3.2)), ("양방향 대역폭", Cm(3.5)), ("전송 지연 (Latency)", Cm(3.3)), ("워크로드 적합성 평가", Cm(7.0))]
    for col_idx, (h_text, w_val) in enumerate(fab_headers):
        cell = t_fab.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    fab_data = [
        ("NVIDIA NVLink 5\n(B200 / NVL-72)", "1,800 GB/s (1.8 TB/s)\n(Liqid 대비 약 28배)", "~100 ns\n(초저지연)", "초거대 LLM 분산 사전학습(Pre-training) 및 텐서 병렬화(TP) 필수"),
        ("NVIDIA InfiniBand\n(NDR 400G / XDR 800G)", "50 ~ 100 GB/s per Port\n(400G ~ 800G)", "~600 ns ~ 1 µs\n(RDMA 가속)", "멀티 노드/멀티 랙 대규모 분산 클러스터링 글로벌 표준"),
        ("Liqid PCIe Gen5 x16\n(NTB Fabric Pooling)", "~64 GB/s\n(단방향 32 GB/s)", "~200 ~ 400 ns\n(NTB 맵핑 오버헤드)", "❌ 대규모 분산 학습 부적합\n✅ 독립 추론 서빙 & 대용량 KV Cache 공유 최적")
    ]

    for row_idx, r_data in enumerate(fab_data, start=1):
        row_cells = t_fab.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(3.5), Cm(3.3), Cm(7.0)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.0)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_CYAN

    # ── [5. 엔터프라이즈 하드웨어 토탈 패키지 제안 모델] ───────────────
    add_section_header(doc, "5", "엔터프라이즈 프라이빗 AI 하드웨어 토탈 패키지 제안 모델",
                       "서버 + GPU + 풀링 섀시 + 스토리지 + RHOAI 결합 턴키 토탈 어플라이언스")

    add_narrative_paragraph(
        doc,
        "엔터프라이즈 고객에게 Liqid를 개별 부품이나 단품 카드로 제안할 경우 도입 및 검증 장벽이 높아지므로, "
        "'서버(Dell) + 가속기(NVIDIA) + 풀링 섀시(Liqid) + 플랫폼(RHOAI) + 올인원 랙'이 결합된 '엔터프라이즈 프라이빗 AI 하드웨어 토탈 패키지' 형태의 턴키 오퍼링으로 제안합니다."
    )

    t_pkg = doc.add_table(rows=6, cols=3)
    t_pkg.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_pkg.autofit = False
    set_table_borders(t_pkg)

    pkg_headers = [("구성 영역", Cm(3.5)), ("토탈 패키지 표준 하드웨어 & S/W 규격", Cm(8.0)), ("조달 및 수행 주체", Cm(5.5))]
    for col_idx, (h_text, w_val) in enumerate(pkg_headers):
        cell = t_pkg.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    pkg_data = [
        ("호스트 컴퓨팅\n(Compute Host)", "Dell PowerEdge R760 / R660 (Dual Intel Xeon, 4~8 노드)\n• PCIe Gen5 HIC 기본 탑재", "델(Dell) 총판 파트너사\n(물리 서버 조달 및 H/W 보증)"),
        ("컴포저블 풀링\n(CDI Chassis)", "Liqid UltraStack PCIe Gen5 섀시 (1~2대)\n• 섀시당 10~30x GPU 수용, 최대 200TB DRAM 풀링", "MZC ISSU (총판 라이선스)\n+ H/W 파트너 랙 마운팅"),
        ("AI 가속기\n(GPU Accelerator)", "NVIDIA RTX 6000 Ada (48GB) 또는 L40S PCIe 16~32대\n• 공랭식 데이터센터 최적화 가속기 팩", "엔비디아(NVIDIA) 파트너사\n(GPU 공급 및 엔터프라이즈 지원)"),
        ("엔터프라이즈 스토리지\n(Storage)", "Dell PowerScale / All-Flash NVMe 스토리지 (100TB~)\n• RAG 문서 파이프라인 및 고속 벡터 저장소", "델(Dell) 파트너사"),
        ("플랫폼 & 솔루션\n(AI Software Stack)", "Liqid Matrix SW + Red Hat OpenShift AI (RHOAI)\n• MZC 엔터프라이즈 문서 파이프라인/코딩 에이전트 탑재", "MZC ISSU\n(AI 아키텍처 SI & 통합 플랫폼 구축)")
    ]

    for row_idx, r_data in enumerate(pkg_data, start=1):
        row_cells = t_pkg.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.5), Cm(8.0), Cm(5.5)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.0)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY_DARK

    # ── [6. 엔터프라이즈 비용 분석 및 TCO 경제성 검토] ───────────────
    add_section_header(doc, "6", "엔터프라이즈 비용 분석 및 TCO 경제성 검토",
                       "초기 투자(CapEx), S/W 라이선스, 전력/상면(OpEx) 및 3개년 TCO 비교")

    add_narrative_paragraph(
        doc,
        "Liqid CDI 기반 솔루션의 사업성 평가에서 핵심은 초기 하드웨어 투자(CapEx), 소프트웨어 라이선스, 그리고 데이터센터 운영비(OpEx)를 종합한 3개년 TCO 경제성입니다."
    )

    t_tco = doc.add_table(rows=6, cols=4)
    t_tco.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tco.autofit = False
    set_table_borders(t_tco)

    tco_headers = [("비용 비교 항목", Cm(3.2)), ("[시나리오 A] NVIDIA NVL-72\n(B200 수랭 턴키)", Cm(4.5)), ("[시나리오 B] 전통적 고정형 서버\n(4x 8-GPU R760, 32x L40S)", Cm(4.5)), ("[시나리오 C] MZC Composable AI\n(1x Liqid + 32x L40S 풀링)", Cm(4.8))]
    for col_idx, (h_text, w_val) in enumerate(tco_headers):
        cell = t_tco.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.2)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    tco_data = [
        ("초기 H/W 도입비 (CapEx)", "약 35억 ~ 45억 원", "약 12억 ~ 14억 원 (서버 4대)", "약 7억 ~ 9억 원 (서버 1~2대)"),
        ("특수 설비 공사비", "수십억 원 (수랭 배관 필수)", "없음 (일반 공랭 전산실)", "없음 (일반 공랭 전산실 즉시 구축)"),
        ("3개년 S/W 라이선스", "NVAIE 기본 포함", "RHOAI 4노드분 (상대적 고비용)", "RHOAI 1~2노드분 + Matrix (약 40% 절감)"),
        ("3개년 전력 및 상면비", "약 4.5억 ~ 6억 원 (150kW)", "약 2.5억 ~ 3.5억 원 (서버 4대)", "약 1.5억 ~ 2억 원 (60kW, 40% 절감)"),
        ("3개년 총 TCO 합계", "50억 원 이상 (특수 연구소)", "약 16억 ~ 19억 원", "약 10억 ~ 12억 원 (약 35~40% 절감)")
    ]

    for row_idx, r_data in enumerate(tco_data, start=1):
        row_cells = t_tco.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(4.5), Cm(4.5), Cm(4.8)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY_DARK
            elif col_idx == 3:
                r.bold = True
                r.font.color.rgb = COLOR_CYAN

    add_sub_header(doc, "■ 3. 신규 비즈니스 인력/조직 셋업 비용 및 BEP(손익분기점) 시뮬레이션")
    
    t_bep = doc.add_table(rows=5, cols=3)
    t_bep.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_bep.autofit = False
    set_table_borders(t_bep)

    bep_headers = [("비즈니스 셋업 & 수익 항목", Cm(4.5)), ("비용 및 시장 현실 산정 근거", Cm(8.0)), ("연간 재무 영향", Cm(4.5))]
    for col_idx, (h_text, w_val) in enumerate(bep_headers):
        cell = t_bep.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.2)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    bep_data = [
        ("전담 조직 셋업 비용\n(5인 전담팀 구성)", "• 세일즈/프리세일즈 2명 + 전문 CDI 엔지니어 2명 + PM 1명\n• 교육 및 데모 랩 운영비 포함", "연간 고정비 약 7억 ~ 9억 원"),
        ("건당 사업 순마진\n(Unit Margin)", "• Liqid S/W 라이선스 마진 + 아키텍처 SI 용역비\n• 건당 평균 마진 약 3,000만 ~ 5,000만 원", "건당 순마진 약 4,000만 원"),
        ("국내 시장 수주 현실\n(Annual Pipeline)", "• 국내 엔터프라이즈 CDI 시장은 초기 니치 단계\n• 1차년도 예상 수주: 연간 1 ~ 2건", "1차년도 이익 약 6,000만 ~ 1억 원"),
        ("예상 BEP 도달 기간", "연간 고정비(8억) 회수를 위해 연간 20건 이상 수주 필요\n➔ 단기간 내 달성 불가능", "최소 2.5년 ~ 3년 소요\n(1년 내 달성 불가 ➔ 부정적)")
    ]

    for row_idx, r_data in enumerate(bep_data, start=1):
        row_cells = t_bep.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(4.5), Cm(8.0), Cm(4.5)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY_DARK
            elif col_idx == 2 and row_idx == 4:
                r.bold = True
                r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    add_narrative_paragraph(
        doc,
        "• 200여 가지 ISV 파트너 재평가 기준 적용: 경쟁력, 시장 성장성, 매출 기여도를 기준으로 200여 개 파트너십을 엄격히 정예화했듯이, "
        "기술적으로 문제가 없더라도 BEP가 1년 단위 이상 소요(2.5~3년)되는 신규 솔루션은 비즈니스 편입에 대해 부정적으로 평가합니다.\n"
        "• 현재 상태 규정: 현재 Liqid는 '단순 기술적 검토(Technical Exploration)' 단계로 동결하며, 신규 전담 조직 셋업이나 공식 솔루션 추가는 전면 보류합니다."
    )

    # ── [7. AI Full Stack 합류 및 Private Cloud 총판 사업 타당성] ──────
    add_section_header(doc, "7", "MZC AI Full Stack 합류 및 Private Cloud 총판 사업 타당성",
                       "프레임워크 파편화 방지, 분기 100억 총판 허들 및 H/W 직접 핸들링 검토")

    add_sub_header(doc, "■ 1. 솔루션 프레임워크 파편화 방지 및 공식 편입 부정적")
    add_narrative_paragraph(
        doc,
        "• 프레임워크 파편화 방지: 기술적으로 가능하다고 무분별하게 신규 H/W를 추가할 경우, 솔루션 프레임워크가 파편화되고 사후 지원 부담이 가중되므로 공식 편입을 지양합니다.\n"
        "• 분기 100억 원 총판 허들 미달: MZC 총판 비즈니스 편입 기준(분기 100억 원 이상)에 크게 미달하며, BEP 도달 기간이 1년 이상 소요되므로 총판 사업 편입은 불가합니다.\n"
        "• 하드웨어 직접 핸들링 배제: MZC가 H/W 수입·재고·AS를 직접 담당하는 위험은 원천 차단합니다."
    )

    add_sub_header(doc, "■ 2. Dell / NVIDIA 공인 총판 지위 활용에 집중")
    add_narrative_paragraph(
        doc,
        "• 기존 핵심 총판 역량 레버리지: MZC가 이미 공인 총판 지위를 확보하고 있는 Dell PowerEdge 서버 및 NVIDIA GPU 표준 라인업에 모든 영업 및 엔지니어링 역량을 집중합니다.\n"
        "• 특수 고객 제한적 연계: 사내 기존 서버 재활용을 강력히 요구하는 특수 사이트에 한하여, 물리 섀시는 H/W 파트너가 납품하고 MZC는 S/W SI로만 선별 대응합니다."
    )

    # ── [8. 결론 및 실행 제언] ─────────────────────────────────────────
    add_section_header(doc, "8", "종합 결론 및 전략적 의사결정 가이드라인",
                       "Dell/NVIDIA 총판 비즈니스 집중 및 H/W 직접 핸들링 배제")

    conclusions = [
        ("• 1. 현재 상태 규정: '단순 기술적 검토' 단계로 유지: ",
         "기술적 가능성은 확인하였으나, BEP 1년 초과(2.5~3년) 및 전담 조직 셋업 부담으로 인해 MZC 공식 솔루션 편입은 부정적(보류)으로 결론"),
        ("• 2. Dell / NVIDIA 공인 총판 사업에 집중 (Core Business First): ",
         "솔루션 프레임워크 파편화를 방지하고, 메가존클라우드의 핵심 강점인 Dell PowerEdge + NVIDIA 표준 엔터프라이즈 인프라 비즈니스에 전사 역량 집중"),
        ("• 3. H/W 직접 핸들링 배제 원칙 준수: ",
         "분기 100억 원 미달 및 재고·AS 리스크에 따라 Liqid 하드웨어 직접 총판/유통은 추진하지 않음")
    ]

    for c_title, c_desc in conclusions:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_before = Pt(2)
        p_c.paragraph_format.space_after = Pt(4)
        p_c.paragraph_format.left_indent = Cm(0.4)
        r_ct = p_c.add_run(c_title)
        r_ct.font.name = FONT_MAIN
        r_ct.font.size = Pt(8.8)
        r_ct.bold = True
        r_ct.font.color.rgb = COLOR_CYAN

        r_cd = p_c.add_run(c_desc)
        r_cd.font.name = FONT_MAIN
        r_cd.font.size = Pt(8.8)
        r_cd.font.color.rgb = COLOR_TEXT_MAIN

    doc.save(str(output_path))
    print(f"✅ Liqid 분석 보고서 DOCX 생성 완료: {output_path}")
    print(f"   파일 크기: {output_path.stat().st_size / 1024:.1f} KB")


def main():
    base_dir = Path(__file__).resolve().parent
    docx_dir = base_dir / "docx"
    docx_dir.mkdir(exist_ok=True)
    
    out_dated = docx_dir / "2026-08-20_Liqid_Composable_GPU_Memory_Pooling_Analysis_Report.docx"
    out_docs = base_dir.parent / "docs" / "2026-08-20_Liqid_Composable_GPU_Memory_Pooling_Analysis_Report.docx"

    build_liqid_report(out_dated)
    build_liqid_report(out_docs)


if __name__ == "__main__":
    main()
