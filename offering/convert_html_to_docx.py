#!/usr/bin/env python3
"""
proposal 폴더 내 HTML 파일들을 DOCX(Word) 문서로 변환하는 스크립트.

HTML의 구조(제목, 표, 목록, 코드 블록, 단락 등)를 파싱하여
서식이 보존된 Word 문서를 생성합니다.
"""

import os
import sys
import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── 스타일 설정 상수 ──────────────────────────────────────────────
FONT_NAME = '맑은 고딕'
FONT_SIZE_TITLE = Pt(22)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(15)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(12)
FONT_SIZE_BODY = Pt(10)
FONT_SIZE_CODE = Pt(9)
FONT_SIZE_SMALL = Pt(8)

COLOR_PRIMARY = RGBColor(0x00, 0xAB, 0xF0)
COLOR_CYAN = RGBColor(0x00, 0xE5, 0xFF)
COLOR_GOLD = RGBColor(0xD4, 0xAF, 0x37)
COLOR_DARK_BG = RGBColor(0x12, 0x15, 0x22)
COLOR_MUTED = RGBColor(0x8B, 0x95, 0xB5)
COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_HEADER_BG = RGBColor(0x1B, 0x1F, 0x33)
COLOR_TABLE_BORDER = RGBColor(0x26, 0x2C, 0x45)


def clean_text(text):
    """HTML 텍스트에서 불필요한 공백과 줄바꿈을 정리"""
    if not text:
        return ''
    # 연속 공백/줄바꿈을 단일 공백으로
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def get_all_text(element):
    """요소 내 모든 텍스트를 재귀적으로 수집"""
    if isinstance(element, NavigableString):
        return str(element)
    texts = []
    for child in element.children:
        texts.append(get_all_text(child))
    return ''.join(texts)


def set_cell_shading(cell, color_hex):
    """테이블 셀 배경색 설정"""
    shading_elm = cell._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading_elm.append(shading)


def add_formatted_text(paragraph, element, font_size=FONT_SIZE_BODY, color=COLOR_BLACK):
    """HTML 인라인 요소를 Word 문단에 서식 적용하여 추가"""
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            run = paragraph.add_run(text)
            run.font.size = font_size
            run.font.color.rgb = color
            run.font.name = FONT_NAME
        return

    if not isinstance(element, Tag):
        return

    tag_name = element.name.lower() if element.name else ''

    if tag_name in ('strong', 'b'):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    run = paragraph.add_run(text)
                    run.bold = True
                    run.font.size = font_size
                    run.font.color.rgb = color
                    run.font.name = FONT_NAME
            else:
                add_formatted_text(paragraph, child, font_size, color)
    elif tag_name in ('em', 'i'):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip():
                    run = paragraph.add_run(text)
                    run.italic = True
                    run.font.size = font_size
                    run.font.color.rgb = color
                    run.font.name = FONT_NAME
            else:
                add_formatted_text(paragraph, child, font_size, color)
    elif tag_name == 'code':
        text = get_all_text(element)
        if text.strip():
            run = paragraph.add_run(text)
            run.font.size = FONT_SIZE_CODE
            run.font.name = 'Consolas'
            run.font.color.rgb = COLOR_PRIMARY
    elif tag_name == 'a':
        text = get_all_text(element)
        href = element.get('href', '')
        if text.strip():
            run = paragraph.add_run(text)
            run.font.size = font_size
            run.font.color.rgb = COLOR_PRIMARY
            run.font.underline = True
            run.font.name = FONT_NAME
    elif tag_name == 'br':
        paragraph.add_run('\n')
    elif tag_name == 'span':
        # span은 클래스에 따라 스타일 결정
        css_class = ' '.join(element.get('class', []))
        if 'price-tag' in css_class or 'accent-gold' in css_class:
            for child in element.children:
                add_formatted_text(paragraph, child, font_size, COLOR_GOLD)
        else:
            for child in element.children:
                add_formatted_text(paragraph, child, font_size, color)
    else:
        for child in element.children:
            add_formatted_text(paragraph, child, font_size, color)


def process_list(doc, ul_element, level=0):
    """HTML 목록(ul/ol)을 Word 문서에 변환"""
    is_ordered = ul_element.name.lower() == 'ol'
    items = ul_element.find_all('li', recursive=False)
    for idx, li in enumerate(items):
        # 리스트 항목 텍스트 수집 (하위 ul/ol 제외)
        text_parts = []
        sub_lists = []
        for child in li.children:
            if isinstance(child, Tag) and child.name in ('ul', 'ol'):
                sub_lists.append(child)
            elif isinstance(child, Tag):
                text_parts.append(get_all_text(child))
            elif isinstance(child, NavigableString):
                text_parts.append(str(child))

        text = clean_text(' '.join(text_parts))
        if text:
            prefix = '  ' * level
            if is_ordered:
                bullet = f'{prefix}{idx + 1}. '
            else:
                bullets = ['•', '◦', '▪']
                bullet = f'{prefix}{bullets[min(level, 2)]} '

            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.left_indent = Cm(1.5 * (level + 1))

            # 인라인 서식 보존하여 추가
            run = para.add_run(bullet)
            run.font.size = FONT_SIZE_BODY
            run.font.name = FONT_NAME
            for child in li.children:
                if isinstance(child, Tag) and child.name in ('ul', 'ol'):
                    continue
                add_formatted_text(para, child, FONT_SIZE_BODY, COLOR_BLACK)

        # 하위 리스트 재귀 처리
        for sub_list in sub_lists:
            process_list(doc, sub_list, level + 1)


def process_table(doc, table_element):
    """HTML 테이블을 Word 테이블로 변환"""
    rows = table_element.find_all('tr')
    if not rows:
        return

    # 최대 열 수 계산
    max_cols = 0
    for row in rows:
        cells = row.find_all(['th', 'td'])
        col_count = sum(int(cell.get('colspan', 1)) for cell in cells)
        max_cols = max(max_cols, col_count)

    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for row_idx, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        col_idx = 0
        for cell in cells:
            if col_idx >= max_cols:
                break
            cell_text = clean_text(get_all_text(cell))
            is_header = cell.name == 'th'

            doc_cell = table.cell(row_idx, col_idx)
            doc_cell.text = ''

            para = doc_cell.paragraphs[0] if doc_cell.paragraphs else doc_cell.add_paragraph()
            run = para.add_run(cell_text)
            run.font.size = FONT_SIZE_BODY
            run.font.name = FONT_NAME

            if is_header:
                run.bold = True
                run.font.color.rgb = COLOR_PRIMARY
                set_cell_shading(doc_cell, 'F0F4FA')
            else:
                run.font.color.rgb = COLOR_BLACK

            colspan = int(cell.get('colspan', 1))
            if colspan > 1 and col_idx + colspan <= max_cols:
                doc_cell.merge(table.cell(row_idx, col_idx + colspan - 1))

            col_idx += colspan

    # 테이블 이후 빈 줄 추가
    doc.add_paragraph()


def process_code_block(doc, pre_element):
    """코드 블록을 Word 문서에 변환"""
    code = pre_element.find('code')
    if code:
        text = code.get_text()
    else:
        text = pre_element.get_text()

    if not text.strip():
        return

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Cm(1)

    run = para.add_run(text.strip())
    run.font.name = 'Consolas'
    run.font.size = FONT_SIZE_CODE
    run.font.color.rgb = COLOR_PRIMARY

    # 배경색 효과 (단락 수준 음영)
    pPr = para._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F4FA',
    })
    pPr.append(shd)


def process_element(doc, element, depth=0):
    """HTML 요소를 재귀적으로 Word 문서 요소로 변환"""
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text and depth == 0:
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = FONT_SIZE_BODY
            run.font.name = FONT_NAME
        return

    if not isinstance(element, Tag):
        return

    tag = element.name.lower() if element.name else ''

    # 스크립트, 스타일, 메타 태그 무시
    if tag in ('script', 'style', 'link', 'meta', 'head', 'noscript', 'svg'):
        return

    # 제목 태그
    if tag == 'h1':
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_heading(text, level=1)
            for run in para.runs:
                run.font.color.rgb = COLOR_PRIMARY
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_H1
        return

    if tag == 'h2':
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_heading(text, level=2)
            for run in para.runs:
                run.font.color.rgb = COLOR_PRIMARY
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_H2
        return

    if tag == 'h3':
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_heading(text, level=3)
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_H3
        return

    if tag == 'h4':
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_heading(text, level=4)
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_H4
        return

    # 테이블
    if tag == 'table':
        process_table(doc, element)
        return

    # 코드 블록
    if tag == 'pre':
        process_code_block(doc, element)
        return

    # 목록
    if tag in ('ul', 'ol'):
        process_list(doc, element)
        doc.add_paragraph()  # 목록 후 빈 줄
        return

    # 단락
    if tag == 'p':
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            # 인라인 서식 보존
            for child in element.children:
                add_formatted_text(para, child, FONT_SIZE_BODY, COLOR_BLACK)
        return

    # 수평선
    if tag == 'hr':
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run('─' * 80)
        run.font.size = FONT_SIZE_SMALL
        run.font.color.rgb = COLOR_MUTED
        return

    # div 계열 — 클래스에 따라 처리
    css_classes = ' '.join(element.get('class', []))

    # badge 요소
    if 'badge' in css_classes:
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f'[ {text} ]')
            run.bold = True
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = FONT_NAME
        return

    # subtitle 요소
    if 'subtitle' in css_classes:
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_MUTED
            run.font.name = FONT_NAME
            run.italic = True
        return

    # card-title 요소
    if 'card-title' in css_classes:
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_heading(text, level=3)
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_H3
                run.font.color.rgb = COLOR_PRIMARY
        return

    # alert-box 요소
    if 'alert-box' in css_classes or 'alert' in css_classes:
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(f'💡 {text}')
            run.font.size = FONT_SIZE_BODY
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = FONT_NAME
            run.italic = True

            # 배경 음영
            pPr = para._element.get_or_add_pPr()
            shd = pPr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): 'E8F4FD',
            })
            pPr.append(shd)
        return

    # cite-box 요소
    if 'cite-box' in css_classes or 'cite' in css_classes:
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(f'📎 {text}')
            run.font.size = FONT_SIZE_SMALL
            run.font.color.rgb = COLOR_MUTED
            run.font.name = FONT_NAME
        return

    # mermaid 다이어그램 — 텍스트로 표시
    if 'mermaid' in css_classes:
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(f'[다이어그램]\n{text}')
            run.font.size = FONT_SIZE_CODE
            run.font.name = 'Consolas'
            run.font.color.rgb = COLOR_MUTED
        return

    # footer
    if tag == 'footer':
        text = clean_text(get_all_text(element))
        if text:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(20)
            run = para.add_run('─' * 60)
            run.font.size = FONT_SIZE_SMALL
            run.font.color.rgb = COLOR_MUTED

            para2 = doc.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = para2.add_run(text)
            run2.font.size = FONT_SIZE_SMALL
            run2.font.color.rgb = COLOR_MUTED
            run2.font.name = FONT_NAME
        return

    # header 태그 (HTML semantic)
    if tag == 'header':
        for child in element.children:
            process_element(doc, child, depth + 1)
        return

    # 기본: 하위 요소 재귀 처리
    for child in element.children:
        process_element(doc, child, depth + 1)


def convert_html_to_docx(html_path, output_path):
    """단일 HTML 파일을 DOCX로 변환"""
    print(f'  📄 변환 중: {html_path.name} → {output_path.name}')

    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'lxml')

    # 문서 제목 추출
    title_tag = soup.find('title')
    doc_title = clean_text(title_tag.get_text()) if title_tag else html_path.stem

    # Word 문서 생성
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    style.font.color.rgb = COLOR_BLACK

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 문서 제목 추가
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(20)
    title_run = title_para.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = FONT_SIZE_TITLE
    title_run.font.color.rgb = COLOR_PRIMARY
    title_run.font.name = FONT_NAME

    # body 내용 변환
    body = soup.find('body')
    if body:
        for child in body.children:
            process_element(doc, child)
    else:
        # body가 없으면 전체 문서 처리
        for child in soup.children:
            process_element(doc, child)

    # 저장
    doc.save(str(output_path))
    print(f'  ✅ 완료: {output_path.name} ({output_path.stat().st_size / 1024:.1f} KB)')


def main():
    """proposal 폴더 내 모든 HTML 파일을 DOCX로 변환"""
    proposal_dir = Path(__file__).parent

    # 출력 디렉토리 생성
    output_dir = proposal_dir / 'docx'
    output_dir.mkdir(exist_ok=True)

    # HTML 파일 목록
    html_files = sorted(proposal_dir.glob('*.html'))

    if not html_files:
        print('❌ proposal 폴더에 HTML 파일이 없습니다.')
        sys.exit(1)

    print(f'\n🔄 HTML → DOCX 변환 시작 (총 {len(html_files)}개 파일)')
    print(f'   출력 경로: {output_dir}\n')

    success_count = 0
    error_count = 0

    for html_path in html_files:
        output_path = output_dir / f'{html_path.stem}.docx'
        try:
            convert_html_to_docx(html_path, output_path)
            success_count += 1
        except Exception as e:
            print(f'  ❌ 오류 ({html_path.name}): {e}')
            error_count += 1

    print(f'\n📊 변환 결과: 성공 {success_count}개 / 실패 {error_count}개')
    print(f'   출력 위치: {output_dir}')


if __name__ == '__main__':
    main()
