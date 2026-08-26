#!/usr/bin/env python3
"""
MZC AI Fullstack COE 보고용 종합 전략 문서 생성기 (DOCX)
- 서술형 문장(Narrative Style) 중심의 완성도 높은 Executive Report
- AI Full Stack 본질적 정의, 글로벌 빅테크 비교 (NVIDIA, Dell, Oracle OCI, IBM), ISSU 4-Layer 파트너 배치
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ── 디자인 색상 ───────────────────────────────────────────────────
COLOR_MZC_BLUE      = RGBColor(0x00, 0xAB, 0xF0)   # #00ABF0
COLOR_NAVY_DARK     = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A
COLOR_PRIMARY_DARK  = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B
COLOR_NVIDIA_GREEN  = RGBColor(0x76, 0xB9, 0x00)   # #76B900
COLOR_DELL_BLUE     = RGBColor(0x00, 0x71, 0xC5)   # #0071C5
COLOR_OCI_RED       = RGBColor(0xC7, 0x46, 0x34)   # #C74634
COLOR_IBM_BLUE      = RGBColor(0x0F, 0x62, 0xFE)   # #0F62FE
COLOR_GOLD          = RGBColor(0xD4, 0xAF, 0x37)   # #D4AF37
COLOR_TEXT_MAIN     = RGBColor(0x1E, 0x24, 0x32)   # #1E2432
COLOR_TEXT_MUTED    = RGBColor(0x64, 0x74, 0x8B)   # #64748B
COLOR_WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

HEX_BG_LIGHT_BLUE   = "F0F7FF"
HEX_BG_LIGHT_GREEN  = "F4FBF0"
HEX_BG_LIGHT_DELL   = "F0F6FC"
HEX_BG_LIGHT_GOLD   = "FFFDF0"
HEX_BG_LIGHT_GRAY   = "F8FAFC"
HEX_BG_DARK_NAVY    = "0F172A"
HEX_BORDER_LIGHT    = "CBD5E1"

FONT_MAIN = "맑은 고딕"


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=140, bottom=140, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=HEX_BORDER_LIGHT, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)


def add_section_header(doc, number_str, title_str, subtitle_str=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r_num = p.add_run(f"제 {number_str} 장.  ")
    r_num.font.name = FONT_MAIN
    r_num.font.size = Pt(14.5)
    r_num.bold = True
    r_num.font.color.rgb = COLOR_MZC_BLUE
    
    r_title = p.add_run(title_str)
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(14.5)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    if subtitle_str:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(10)
        p_sub.paragraph_format.keep_with_next = True
        r_sub = p_sub.add_run(subtitle_str)
        r_sub.font.name = FONT_MAIN
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = COLOR_TEXT_MUTED
        r_sub.italic = True


def add_sub_header(doc, title_str, color=COLOR_PRIMARY_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title_str)
    r.font.name = FONT_MAIN
    r.font.size = Pt(11.5)
    r.bold = True
    r.font.color.rgb = color


def add_narrative_paragraph(doc, text_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text_str)
    r.font.name = FONT_MAIN
    r.font.size = Pt(9.8)
    r.font.color.rgb = COLOR_TEXT_MAIN


def add_callout_box(doc, title, text_lines, border_color_hex="00ABF0", bg_color_hex="F0F7FF"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_shading(cell, bg_color_hex)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:color="{border_color_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    
    r_title = p.add_run(f"💡 {title}\n")
    r_title.bold = True
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(10)
    if border_color_hex == "00ABF0":
        r_title.font.color.rgb = COLOR_MZC_BLUE
    elif border_color_hex == "D4AF37":
        r_title.font.color.rgb = COLOR_GOLD
    else:
        r_title.font.color.rgb = COLOR_NAVY_DARK

    for line in text_lines:
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(2)
        r_line = p_line.add_run(line)
        r_line.font.name = FONT_MAIN
        r_line.font.size = Pt(9.2)
        r_line.font.color.rgb = COLOR_TEXT_MAIN

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)


def build_coe_report(output_path):
    doc = Document()
    
    # ── 기본 페이지 설정 (A4 여백 2.2cm) ──────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        
        # 헤더 / 푸터 설정
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("MZC AI Full Stack Master Report | Confidential")
        hrun.font.name = FONT_MAIN
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = COLOR_TEXT_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("© 2026 MEGAZONECLOUD Corp. All Rights Reserved. | Executive Strategic Review")
        frun.font.name = FONT_MAIN
        frun.font.size = Pt(8)
        frun.font.color.rgb = COLOR_TEXT_MUTED

    # ── [표지 / 메타 정보 영역] ──────────────────────────────────────────
    p_badge = doc.add_paragraph()
    r_b = p_badge.add_run("🏛️ MEGAZONECLOUD AI FULL STACK STRATEGIC REPORT · 2026.08")
    r_b.font.name = FONT_MAIN
    r_b.font.size = Pt(9)
    r_b.bold = True
    r_b.font.color.rgb = COLOR_MZC_BLUE

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("MZC AI Full Stack 아키텍처 서비스 전략 보고서")
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(18)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("단일 벤더 종속(Lock-in)을 탈피한 개방형·모듈형 AI 풀스택 정의, 글로벌 빅테크 심층 비교 및 ISSU ISV 4-Layer 오케스트레이션")
    r_sub.font.name = FONT_MAIN
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    # 메타 정보 테이블
    t_meta = doc.add_table(rows=2, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_meta.autofit = False
    set_table_borders(t_meta)
    
    meta_info = [
        ("• 보고 대상: MZC", "• 발행 부서: Integrated Solution Sales Unit (ISSU)"),
        ("• 작성 일자: 2026년 8월 26일 (v3.2 Final)", "• 보안 등급: 사내 대외비 (Confidential - Internal Only)")
    ]
    for r_idx, (c1, c2) in enumerate(meta_info):
        for c_idx, text in enumerate([c1, c2]):
            cell = t_meta.cell(r_idx, c_idx)
            cell.width = Cm(8.2)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_TEXT_MUTED

    # 구분선
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(10)
    p_hr.paragraph_format.space_after = Pt(12)
    r_hr = p_hr.add_run("━" * 58)
    r_hr.font.color.rgb = COLOR_MZC_BLUE
    r_hr.font.size = Pt(8)

    # ── [전체 목차] ────────────────────────────────────────────────────
    add_sub_header(doc, "📋 전체 보고서 목차 (Table of Contents)")
    toc_items = [
        "제 1 장. Executive Summary 및 추진 배경",
        "제 2 장. AI Full Stack의 본질적 정의 및 도입 당위성 (Why AI Full Stack?)",
        "제 3 장. 글로벌 빅테크 AI 스택 심층 비교 분석 (NVIDIA, Dell, Oracle OCI, IBM vs MZC)",
        "제 4 장. MZC ISSU 4-Layer AI Full Stack Ecosystem (H/W, OS/가상화, LLM/Tool, Solution)",
        "제 5 장. 기업 및 사용자 규모별 표준 H/W & 솔루션 사이징 오퍼링 가이드 (Sizing Tier Guide)",
        "제 6 장. MZC AI Full stack 점진적 구축을 위한 4단계 Process",
        "제 7 장. AI Full Stack 솔루션 패키지 및 산업군별 적용 시나리오 (예)",
        "제 8 장. 향후 실행 계획 (Offering 가이드 및 차기 보고 준비 과제)",
        "[부록 1] Layer 3 (LLM Model + Tool Layer) 세부 분야별 ISV 솔루션 경쟁력 우선순위 분석",
        "[부록 2] 주요 파트너 솔루션 공식 웹사이트 및 참고 문헌 (References & Ecosystem Links)"
    ]
    for item in toc_items:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(1)
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.left_indent = Cm(0.5)
        r = p_t.add_run(f"• {item}")
        r.font.name = FONT_MAIN
        r.font.size = Pt(9.2)
        r.font.color.rgb = COLOR_TEXT_MAIN
        r.bold = True

    # ── [제 1 장. Executive Summary 및 추진 배경] ────────────────────────
    add_section_header(doc, "1", "Executive Summary 및 추진 배경", 
                       "2026 엔터프라이즈 AI 패러다임 전환과 메가존클라우드의 전략적 포지셔닝")

    add_narrative_paragraph(
        doc,
        "2026년 글로벌 및 국내 엔터프라이즈 AI 시장은 단순한 파운데이션 모델의 개념 검증(PoC)이나 단편적인 생성형 AI SaaS 도입 단계를 지나, "
        "기업 내부의 핵심 데이터 자산과 업무 프로세스가 결합되어 24시간 365일 비즈니스 인텔리전스와 자동화 태스크를 연속 생산하는 'AI Factory(AI 공장)' 체제로 전면 재편되고 있습니다. "
        "이러한 전환기 속에서 기업 고객들은 퍼블릭 클라우드 LLM API 호출에 따른 천문학적인 토큰 과금 부담, 금융·공공·제조 분야의 엄격한 망분리 및 데이터 주권(Data Sovereignty) 규제, "
        "그리고 기존 레거시 인프라(서버, 스토리지, ERP, 가상화)와 AI 신기술 간의 극심한 파편화라는 3대 장벽에 직면해 있습니다."
    )
    add_narrative_paragraph(
        doc,
        "메가존클라우드 ISSU(Integrated Solution Sales Unit)는 이러한 시장의 본질적 갈증을 해결하기 위해, 특정 단일 벤더나 솔루션에 종속되지 않는 "
        "'개방형·모듈형(Vendor-Agnostic) 엔터프라이즈 AI Full Stack' 프레임워크를 수립하여 고객들의 AX 전환을 가속화 할 수 있도록 준비하였습니다. "
        "본 보고서는 메가존클라우드가 보유한 국내 최고의 ISV 파트너십 생태계를 바탕으로, H/W부터 인프라 OS, 파운데이션 모델/도구, 그리고 최상위 비즈니스 솔루션에 이르기까지 "
        "4개 계층(4-Layer)으로 체계화된 구축 전략을 정리하여 실행 동력을 확보하고자 작성되었습니다."
    )

    add_callout_box(
        doc,
        "MZC AI Full Stack의 핵심 철학: 'Zero Lock-in, Open Modular Orchestration'",
        [
            "• 특정 솔루션(Nutanix, Articul8, Cohere, Acryl Jonathan 등)을 단일 필수로 강제하지 않고, 고객의 IT 환경과 예산에 맞춘 최적의 Best-of-Breed 조합을 제공합니다.",
            "• 고객의 기존 투자 자산(Dell, VMware, 사내 DB/스토리지)을 100% 보호하면서 필요한 계층만 모듈식으로 증설합니다.",
            "• 메가존클라우드는 독립적인 최상위 시스템 인티그레이터(SI)이자 MSP로서 전 계층의 통합 안정성과 완벽한 폐쇄망(Air-Gap) 소버린 보안을 책임집니다."
        ],
        border_color_hex="00ABF0",
        bg_color_hex="F0F7FF"
    )

    # ── [제 2 장. AI Full Stack의 본질적 정의 및 도입 당위성] ────────────
    add_section_header(doc, "2", "AI Full Stack의 본질적 정의 및 도입 당위성",
                       "과거 기술 매칭을 넘어선 '데이터 집약적 지능 공장'으로서의 필요성 분석")

    add_narrative_paragraph(
        doc,
        "과거 IT 아키텍처에서 논의되던 3-Tier(Web-WAS-DB)나 OSI 7-Layer는 네트워크 트랜잭션 처리를 위한 '기능적·기술적 계층 매칭'에 불과했습니다. "
        "반면, 생성형 AI와 에이전틱(Agentic) 워크로드를 위한 AI Full Stack은 단순한 소프트웨어 레이어의 나열이 아니라, "
        "하드웨어 연산 능력(Compute), 초고속 I/O 패브릭, 가상화/오케스트레이션, 도메인 특화 모델, 데이터 거버넌스, 비즈니스 애플리케이션이 상호 유기적으로 결합되어 "
        "실시간으로 추론과 지능을 생성하는 '데이터 집약적 지능 생산 시스템'을 의미합니다."
    )

    add_sub_header(doc, "■ 엔터프라이즈에 AI Full Stack이 반드시 필요한 3대 당위성")
    add_narrative_paragraph(
        doc,
        "첫째, '데이터 주권(Data Sovereignty)과 완전 폐쇄망(Air-Gap) 보안의 실현'입니다. 금융, 공공, 국방, 첨단 제조 분야는 민감 데이터의 외부 반출이 법적으로 엄격히 통제됩니다. "
        "단편적인 클라우드 API 호출로는 규제를 충족할 수 없으며, 물리적 온프레미스 인프라부터 폐쇄망 내부에서 독립 구동되는 파운데이션 모델, 벡터 DB, PII 비식별화 거버넌스가 "
        "하나의 스택으로 결합되어야만 진정한 데이터 주권이 보장됩니다."
    )
    add_narrative_paragraph(
        doc,
        "둘째, '토큰 비용 폭증 방어 및 장기 TCO 절감'입니다. 사내 수천 명의 임직원이 외부 LLM API를 일상적으로 사용할 경우 연간 발생하는 토큰 과금은 통제 불가능한 수준에 도달합니다. "
        "사내 온프레미스 GPU/NPU 클러스터 상에 최적화된 경량 모델(sLLM)과 고속 추론 서빙 엔진(vLLM, NIM)을 내재화하면 외부 API 대비 60% 이상의 비용 절감 효과를 거둘 수 있습니다."
    )
    add_narrative_paragraph(
        doc,
        "셋째, '인프라-추론-애플리케이션 간 수직 최적화를 통한 극대화된 성능'입니다. AI 추론의 지연 시간(Latency)과 초당 토큰 처리량(TPS)은 하드웨어 가속기, "
        "스토리지 I/O(GPUDirect Storage), 컨테이너 스케줄러, 추론 서빙 엔진이 일체화되어 조율될 때 비로소 극대화됩니다. 분절된 개별 솔루션 도입은 필연적으로 심각한 병목을 유발합니다."
    )

    # ── [제 3 장. 글로벌 빅테크 AI 스택 심층 비교 분석] ─────────────────
    add_section_header(doc, "3", "글로벌 빅테크 AI 스택 심층 비교 분석",
                       "NVIDIA, Dell, Oracle OCI, IBM watsonx 대비 MZC의 차별적 경쟁 우위")

    add_narrative_paragraph(
        doc,
        "현재 글로벌 엔터프라이즈 AI 시장은 각 벤더의 전통적 강점에 따라 상이한 풀스택 전략을 제시하고 있습니다. "
        "엔비디아의 하드웨어 중심 수직 통합부터 델의 완제품 패키지, 오라클의 DB 결합형 클라우드, IBM의 거버넌스 중심 스택에 이르기까지 "
        "각 진영의 특징과 한계를 비교 분석하고, 메가존클라우드의 개방형 모듈형 스택이 갖는 독보적 가치를 규명합니다."
    )

    t_comp = doc.add_table(rows=7, cols=5)
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_comp.autofit = False
    set_table_borders(t_comp)

    headers_comp = [
        ("구분 차원", Cm(2.2)),
        ("⚡ NVIDIA AI Factory", Cm(3.5)),
        ("💻 Dell AI Factory", Cm(3.6)),
        ("🔴 Oracle OCI AI", Cm(3.5)),
        ("🌐 MZC Modular Fullstack", Cm(3.7))
    ]
    for col_idx, (h_text, w_val) in enumerate(headers_comp):
        cell = t_comp.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    comp_rows = [
        ("아키텍처 성격", "단일 벤더 수직 통합\n(Vertical Lock-in)", "엔터프라이즈 H/W 턴키\n(Turnkey Infrastructure)", "DB 중심 퍼블릭 클라우드\n(Cloud DB Centric)", "모듈형 소버린 플랫폼\n(Vendor-Agnostic Modular)"),
        ("컴퓨트 & 가속기", "DGX B200/GB200 독점\n타사 칩셋 지원 불가", "Dell PowerEdge XE9680\n(NVIDIA GPU 탑재)", "OCI Supercluster\n(H100 / AMD MI300X)", "멀티 벤더 지원\n(Dell + NVIDIA, 리벨리온, 퓨리오사 NPU)"),
        ("스토리지 & 패브릭", "InfiniBand / Spectrum-X\n(자체 스토리지 없음)", "Dell PowerScale (DGX 인증)\nPowerStore 일체형", "OCI 고성능 블록/오브젝트\nRoCEv2 패브릭", "고객 기존 스토리지 보호\n+ Dell PowerScale"),
        ("가상화 & 플랫폼", "DGX OS + Run:ai\n+ NVAIE 런타임", "Dell Hub on Hugging Face\n+ NVAIE 탑재", "OCI OKE (Kubernetes)\n+ Managed AI Services", "Red Hat RHOAI / Nutanix AHV\n유연 선택"),
        ("모델 & 솔루션", "NIM, NeMo, Blueprints\n(자사 레퍼런스 중심)", "NVAIE NIM 탑재\n제휴 오픈소스 배포", "OCI GenAI, Cohere,\nLlama, HeatWave DB", "A8, Cohere,\n오픈소스 LLM"),
        ("주요 한계 / 특성", "고비용 구독료 부담,\n스토리지 서드파티 의존", "H/W 중심 패키지,\n사내 기간계 커스텀 SI 한계", "퍼블릭 클라우드 종속,\n완전 온프레미스 폐쇄망 한계", "Zero Lock-in, 기존 인프라 보호,\n100% 폐쇄망 온프레미스 소버린")
    ]

    for row_idx, r_data in enumerate(comp_rows, start=1):
        row_cells = t_comp.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(2.2), Cm(3.5), Cm(3.6), Cm(3.5), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
            elif col_idx == 4:
                r.bold = True
                r.font.color.rgb = COLOR_MZC_BLUE

    add_callout_box(
        doc,
        "IBM watsonx 및 Bob 대비 MZC의 차별화 우위",
        [
            "• IBM은 watsonx.ai, .data, .governance와 자체 파운데이션 모델(Granite/Bob), Red Hat OpenShift를 결합한 스택을 제안하지만, IBM 전용 제품군 중심의 폐쇄성과 고가의 전문 컨설팅 비용이 수반됩니다.",
            "• 메가존클라우드는 특정 솔루션에 얽매이지 않고 국내외 시장의 다양한 Best-of-Breed ISV 솔루션을 고객 맞춤형으로 유연하게 결합하여 TCO를 최적화합니다."
        ],
        border_color_hex="D4AF37",
        bg_color_hex="FFFDF0"
    )

    # ── [제 4 장. MZC ISSU 4-Layer AI Full Stack Ecosystem] ──────────────
    add_section_header(doc, "4", "MZC ISSU 4-Layer AI Full Stack Ecosystem",
                       "메가존클라우드 ISSU 엔터프라이즈 모듈형 AI 풀스택 4개 계층 정의 및 파트너 에코시스템")

    add_narrative_paragraph(
        doc,
        "메가존클라우드 ISSU의 공식 ISV 파트너십 데이터베이스를 분석하여, "
        "엔터프라이즈 온프레미스 AI Full Stack을 구성하는 4개 핵심 계층(H/W ➔ OS/가상화 ➔ LLM 모델/도구 ➔ 비즈니스 솔루션)에 최적의 파트너 솔루션들을 매핑 배치하였습니다."
    )

    # 4-Layer 정의 및 역할 서술형 요약
    add_sub_header(doc, "■ MZC AI Full Stack 4개 계층(4-Layer) 정의 및 핵심 역할 요약")

    add_narrative_paragraph(
        doc,
        "메가존클라우드의 모듈형 AI Full Stack 프레임워크는 엔터프라이즈 AI 서비스 전 과정을 유기적으로 결합하는 4개 핵심 계층(4-Layer)으로 체계화되어 있으며, 각 계층의 명확한 정의와 핵심 역할은 다음과 같습니다."
    )

    layer_descriptions = [
        ("• Layer 04 (Solution Layer - 엔터프라이즈 AI 솔루션 & 비즈니스 애플리케이션): ",
         "현업 비즈니스 부서와 실무 개발자가 직접 상호작용하는 '최상위 인텔리전스 및 데이터 주권 계층'입니다. 다중 에이전트(Agentic AI) 오케스트레이션과 도메인 특화 지식 그래프(Knowledge Graph) 자동 구축을 통해 환각을 원천 방지하며, 사내 구축형 소버린 파운데이션 모델(Articul8, Cohere), AI 코딩 에이전트 완결 배포, 개인정보 비식별화 및 망분리 규제를 완벽하게 충족하는 온프레미스 솔루션을 제공합니다."),
        ("• Layer 03 (LLM Model + Tool Layer - 모델 서빙·데이터 플랫폼·MLOps): ",
         "하부 인프라와 상부 비즈니스 솔루션을 유기적으로 중계하는 '핵심 AI 미들웨어 및 데이터 플랫폼 계층'입니다. 온프레미스 대규모 비정형 데이터 레이크하우스 구축 및 시맨틱 벡터 검색/캐싱 파이프라인을 지원하고, vLLM·TensorRT-LLM·NVIDIA NIM 기반의 초저지연·고처리량 GPU 추론 서빙과 엔드투엔드 MLOps 수명주기 관리 및 실시간 LLM Observability(AIOps 관제) 체계를 운영합니다."),
        ("• Layer 02 (OS + Virtualization Layer - 운영체제 및 가상화 오케스트레이션): ",
         "물리 하드웨어 자원을 논리적 GPU 풀과 컨테이너 클러스터로 유연하게 추상화하는 '코어 가상화 계층'입니다. Nutanix 무상 AHV 및 Red Hat OpenShift (RHOAI)를 통해 인프라 TCO를 50% 절감하고, 엔터프라이즈 K8s 및 GPU 슬라이싱(vGPU / MIG)을 통해 자원 활용률을 극대화하며 온프레미스 워크로드를 무중단으로 연계합니다."),
        ("• Layer 01 (H/W Layer - 가속 컴퓨팅 서버 & 고성능 스토리지·네트워크): ",
         "24x7 무장애 AI 대규모 연산과 페타바이트급 I/O를 물리적으로 지탱하는 '최하위 가속 하드웨어 인프라 계층'입니다. 최신 엔터프라이즈 가속 서버(Dell PowerEdge XE9680 / XE9640 / R760xa / R760)에 고성능 가속기(NVIDIA GPU, 리벨리온, 퓨리오사AI)를 장착하고, DGX SuperPOD 인증 고성능 올플래시 병렬 파일 스토리지(Dell PowerScale F900/F710)와 무손실 400GbE RoCEv2 / InfiniBand 네트워크 패브릭을 무결성 있게 공급합니다.")
    ]

    for title_text, desc_text in layer_descriptions:
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(2)
        p_l.paragraph_format.space_after = Pt(3)
        p_l.paragraph_format.left_indent = Cm(0.4)
        r_t = p_l.add_run(title_text)
        r_t.font.name = FONT_MAIN
        r_t.font.size = Pt(8.8)
        r_t.bold = True
        if "04" in title_text: r_t.font.color.rgb = RGBColor(0x9D, 0x4E, 0xDD)
        elif "03" in title_text: r_t.font.color.rgb = COLOR_GOLD
        elif "02" in title_text: r_t.font.color.rgb = COLOR_NVIDIA_GREEN
        elif "01" in title_text: r_t.font.color.rgb = COLOR_DELL_BLUE
        
        r_d = p_l.add_run(desc_text)
        r_d.font.name = FONT_MAIN
        r_d.font.size = Pt(8.8)
        r_d.font.color.rgb = COLOR_TEXT_MAIN

    add_sub_header(doc, "■ 계층별 파트너 솔루션 에코시스템 매핑 상세")

    t_layers = doc.add_table(rows=5, cols=3)
    t_layers.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_layers.autofit = False
    set_table_borders(t_layers)

    l_headers = [("계층 (Layer)", Cm(3.5)), ("핵심 역할 및 모듈 정의", Cm(5.2)), ("ISSU 파트너 및 솔루션 배치 (Ecosystem)", Cm(7.8))]
    for col_idx, (h_text, w_val) in enumerate(l_headers):
        cell = t_layers.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.8)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    layer_data = [
        (
            "Layer 04\nSolution Layer\n(비즈니스 솔루션)",
            "【계층 정의 및 역할】\n현업 비즈니스 부서와 최고경영진이 직접 활용하는 최상위 산업 인텔리전스 계층으로, 다중 에이전트(Agent of Agents) 오케스트레이션, 산업 도메인 특화 지식 그래프(Knowledge Graph Substrate), 그리고 사내 소버린 파운데이션 모델(Articul8, Cohere)을 결합하여 미션 크리티컬 비즈니스 의사결정을 자율화합니다.\n\n【주요 적용 비즈니스 분야 & Use Case】\n• 제조·항공우주 MRO: Articul8 기반 설비 예지보전(PdM), MRO 조립·정밀 정비\n• 금융: Articul8 & Cohere 기반 기업 공시/재무제표 자동 분석, 여신 심사 의사결정 지원",
            "• 산업 도메인 특화 GenAI & 에이전틱: Articul8 (A8 Model Mesh & Knowledge Graph)\n"
            "• 사내 구축형 엔터프라이즈 소버린 모델: Cohere (Command R+ / Embed / Rerank)"
        ),
        (
            "Layer 03\nLLM Model + Tool Layer\n(모델·개발·데이터·AIOps)",
            "【계층 정의 및 역할】\n물리 인프라의 가속 연산 파워를 개발 생산성과 엔터프라이즈 지능으로 변환하는 온프레미스 코어 미들웨어 계층으로, 초고속 추론 서빙(vLLM, LiteLLM), 사내 폐쇄망 오픈소스 sLLM 및 멀티모달 vLM, GitLab CI/CD 연동 사내 코딩 에이전트(MCP), 고성능 AI 오브젝트 스토리지(MinIO), 초저지연 시맨틱 캐시(Redis Enterprise), 실시간 데이터 스트리밍(Confluent), 그리고 인과관계 기반 풀스택 AIOps(Dynatrace)를 제공합니다.\n\n【주요 적용 비즈니스 분야 & Use Case】\n• 초고속 추론 & 게이트웨이: vLLM 오픈소스 표준 서빙, LiteLLM 통합 라우팅 게이트웨이 (<80ms)\n• 사내 소버린 언어/비전 모델 & 개발 생산성: sLLM & vLM 기반 비정형 멀티모달 분석, MCP (사내 코딩 에이전트), GitLab (CI/CD)\n• AI 스토리지 & 시맨틱 캐시: MinIO 고성능 AI 스토리지, Redis Enterprise 시맨틱 캐시 (비용 40% 절감), Confluent 실시간 CDC 파이프라인\n• 전사 AIOps 실시간 관제: Dynatrace Managed 기반 LLM 토큰/지연시간/GPU 병목 인과관계 추적",
            "• 초고속 추론 서빙 & 모델 라우터: vLLM (PagedAttention 오픈소스 표준), LiteLLM (통합 모델 라우팅 & API 게이트웨이)\n"
            "• 사내 소버린 모델 & 개발 생산성: sLLM (Qwen2.5-Coder, Llama 3.3), vLM (Qwen2-VL, Llama-Vision), MCP (사내 코딩 에이전트), GitLab (사내 Git/CI/CD)\n"
            "• 온프레미스 AI 스토리지 & 데이터 플랫폼: MinIO (S3 호환 AI 스토리지), Redis Enterprise (시맨틱 캐시), Confluent (Kafka 파이프라인)\n"
            "• 온프레미스 AIOps & 풀스택 옵저버빌리티: Dynatrace (Dynatrace Managed 폐쇄망 지원 & Davis AI RCA)"
        ),
        (
            "Layer 02\nOS + Virtualization\n(가상화·오케스트레이션)",
            "【계층 정의 및 역할】\n물리적 서버와 이기종 가속기 자원을 단일 풀로 추상화하고, GPU 하드웨어 슬라이싱(vGPU/MIG)을 수행하며, 온프레미스 클라우드 네이티브 쿠버네티스 컨테이너 환경을 무중단 통합 운영하는 가상화 인프라 소프트웨어 계층입니다.\n\n【주요 적용 비즈니스 분야 & Use Case】\n• 인프라 비용 방어: Nutanix 무상 AHV / OpenShift 전환으로 가상화 TCO 50% 절감\n• GPU 슬라이싱: 개발/연구/운영팀별 vGPU 및 MIG 분할 할당으로 가동률 극대화\n• 온프레미스 폐쇄망 보안: 완전 망분리 환경에서 엔터프라이즈 K8s 자동화 배포",
            "• 엔터프라이즈 온프레미스 가상화 & 컨테이너: Red Hat OpenShift / RHOAI SNO, Nutanix NCI + Nutanix NKP (K8s Platform)"
        ),
        (
            "Layer 01\nH/W Layer\n(서버·스토리지·패브릭)",
            "【계층 정의 및 역할】\n대규모 파운데이션 모델의 초저지연 실시간 추론 및 분산 연산을 물리적으로 지탱하는 엔터프라이즈 서버, 가속기 실리콘 칩셋, RoCEv2/인피니밴드 초고속 네트워크 패브릭, 그리고 고성능 올플래시 스토리지 인프라 계층입니다.\n\n【주요 적용 비즈니스 분야 & Use Case】\n• 전사 온프레미스 AI Factory: 8-GPU Dell XE9680 기반 24x7 대규모 학습/추론\n• 망분리 소버린 AI DC: 금융/공공/국방/제조 전용 Air-Gap 고성능 컴퓨팅 룸\n• 고속 AI 데이터 I/O: Dell PowerScale F900 올플래시 & 400GbE RoCEv2 패브릭",
            "• 엔터프라이즈 AI 서버: Dell PowerEdge 서버 (XE9680 / XE9640 수랭식 / R760xa / R760)\n"
            "• 가속기 실리콘 (GPU/NPU): NVIDIA GPU (B200/300, H100/200, L40S, RTX 6000 Ada/Blackwell), 리벨리온, 퓨리오사AI\n"
            "• 엔터프라이즈 고성능 스토리지: Dell PowerScale 스토리지 (DGX SuperPOD 인증) / 400GbE RoCEv2, InfiniBand"
        )
    ]

    for row_idx, (l_name, l_def, l_isv) in enumerate(layer_data, start=1):
        row_cells = t_layers.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([l_name, l_def, l_isv], [Cm(3.5), Cm(5.2), Cm(7.8)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                if "04" in text: r.font.color.rgb = RGBColor(0x9D, 0x4E, 0xDD)
                elif "03" in text: r.font.color.rgb = COLOR_GOLD
                elif "02" in text: r.font.color.rgb = COLOR_NVIDIA_GREEN
                elif "01" in text: r.font.color.rgb = COLOR_DELL_BLUE

    # ── [제 5 장. 기업 및 사용자 규모별 표준 H/W & 솔루션 사이징 오퍼링 가이드] ──
    add_section_header(doc, "5", "기업 및 사용자 규모별 표준 H/W & 솔루션 사이징 오퍼링 가이드",
                       "고객 규모, 주요 워크로드 및 예산에 따른 3단계 표준 아키텍처 및 추정 TCO 오퍼링")

    add_narrative_paragraph(
        doc,
        "고객사의 비즈니스 규모(사용자 수, 개발자 규모), 주요 워크로드(사내 코딩 에이전트, Articul8 산업 지식 그래프, 전사 AI Factory), "
        "그리고 예산 규모(S/W 솔루션 TCO 및 H/W 조달 비용)에 따라 3단계 표준 패키지 오퍼링(Small ➔ Medium ➔ Large Enterprise)을 유연하게 제안 및 구성할 수 있습니다."
    )

    t_sizing = doc.add_table(rows=4, cols=6)
    t_sizing.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_sizing.autofit = False
    set_table_borders(t_sizing)

    s_headers = [
        ("오퍼링 티어", Cm(2.2)),
        ("권장 대상 및 규모", Cm(2.8)),
        ("하드웨어 구성 (L1/L2)", Cm(3.6)),
        ("예상 H/W 가격\n(별도 견적)", Cm(2.2)),
        ("솔루션/모델 구성 (L3/L4)", Cm(3.6)),
        ("예상 S/W TCO\n(연간 라이선스)", Cm(2.1))
    ]
    for col_idx, (h_text, w_val) in enumerate(s_headers):
        cell = t_sizing.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=70, right=70)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.0)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    sizing_data = [
        (
            "Option 01\nSmall 티어\n(팀/스타트업)",
            "• 50인 이하 조직 / 개발팀\n• 코딩 에이전트 전용 PoC\n• 단일 부서 파일럿",
            "• Dell PowerEdge R760 (2 GPU)\n• 2x RTX 6000 Ada or 리벨리온 NPU\n• NVMe SSD (2~4TB) + 10/25GbE\n• OpenShift SNO or Nutanix 1-Node",
            "약 4,500만 ~\n6,000만 원\n(서버/GPU)",
            "• GitLab CI/CD + 사내 코딩 에이전트\n• 오픈소스 Qwen2.5-Coder-32B\n• vLLM 초고속 서빙",
            "약 5,000만 ~\n6,000만 원\n(1억 미만 턴키)"
        ),
        (
            "Option 02\nMedium 티어\n(중견/사업부)",
            "• 50인 ~ 200인 중견기업\n• 대기업 제조/금융 사업부\n• Articul8 + 코딩 에이전트",
            "• Dell PowerEdge R760xa / XE9640 (4 GPU)\n• 4x L40S 48GB or 퓨리오사AI NPU\n• 올플래시 NVMe (20~50TB) + 100GbE\n• OpenShift 3-Node / Nutanix NCI",
            "약 1.5억 ~\n2.2억 원\n(H/W 별도)",
            "• Articul8 Model Mesh (지식 그래프)\n• Cohere Command R+\n• vLLM / NVIDIA NIM 이중화\n• MinIO + Redis Enterprise 캐시",
            "약 1.8억 ~\n2.5억 원\n(SW 라이선스)"
        ),
        (
            "Option 03\nLarge 티어\n(대기업/그룹사)",
            "• 200인 ~ 500인+ 전사 규모\n• 엔터프라이즈 AI Factory\n• 제조/MRO/금융 복합 워크로드",
            "• Dell PowerEdge XE9680 (8x H100/B200)\n• Dell PowerScale F900/F710 (100TB~PB)\n• 400GbE Spectrum-X / InfiniBand\n• Multi-Cluster OpenShift / Nutanix NKP",
            "약 6억 ~\n10억+ 원\n(사양별 산정)",
            "• Articul8 엔터프라이즈 (제조/MRO/금융)\n• Cohere Suite / 전사 코딩에이전트 + GitLab\n• MinIO + Confluent + Dynatrace",
            "약 3.5억 ~\n5.5억+ 원\n(전사 라이선스)"
        )
    ]

    for row_idx, r_data in enumerate(sizing_data, start=1):
        row_cells = t_sizing.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(2.2), Cm(2.8), Cm(3.6), Cm(2.2), Cm(3.6), Cm(2.1)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.5)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                if "Small" in text: r.font.color.rgb = COLOR_NVIDIA_GREEN
                elif "Medium" in text: r.font.color.rgb = COLOR_GOLD
                elif "Large" in text: r.font.color.rgb = COLOR_DELL_BLUE
            elif col_idx == 5:
                r.bold = True
                if "Small" in r_data[0]: r.font.color.rgb = COLOR_NVIDIA_GREEN
                elif "Medium" in r_data[0]: r.font.color.rgb = COLOR_GOLD
                elif "Large" in r_data[0]: r.font.color.rgb = COLOR_DELL_BLUE

    p_guide = doc.add_paragraph()
    p_guide.paragraph_format.space_before = Pt(4)
    p_guide.paragraph_format.space_after = Pt(8)
    p_guide.paragraph_format.left_indent = Cm(0.4)
    r_g = p_guide.add_run("💡 비용 산정 가이드라인: 상기 표의 TCO 금액은 S/W 솔루션 및 구축/운영 라이선스 기준이며, H/W 장비 가격(서버, 가속기 GPU/NPU, 스토리지, 네트워크)은 고객사의 요구 사양 및 도입 시점에 맞춰 Dell/NVIDIA 공인 총판 견적으로 별도 제공됩니다.")
    r_g.font.name = FONT_MAIN
    r_g.font.size = Pt(8.2)
    r_g.font.color.rgb = COLOR_TEXT_MUTED
    r_g.italic = True

    # ── [제 6 장. MZC AI Full stack 점진적 구축을 위한 4단계 Process] ─────
    add_section_header(doc, "6", "MZC AI Full stack 점진적 구축을 위한 4단계 Process",
                       "고객 비즈니스 성과 조기 실현 및 확실한 비용 이득을 위한 단계별 구축 프로세스")

    add_narrative_paragraph(
        doc,
        "엔터프라이즈 고객의 초기 투자 리스크를 최소화하고 단계적으로 전사 확장을 유도하기 위해, "
        "메가존클라우드는 'MZC AI Full stack 점진적 구축을 위한 4단계 Process를 제안'합니다. "
        "각 단계는 사전 TCO 진단 및 신속한 파일럿 PoC(Phase 1), VMware 비용 폭증 방어 및 코어 인프라 현대화(Phase 2), "
        "온프레미스 GPU/NPU 기반 소버린 모델 및 RAG/코딩 에이전트 가동(Phase 3), "
        "그리고 24x7 SLA 무장애 관제 및 전사 AI Factory 확장(Phase 4)을 위한 것입니다."
    )

    add_narrative_paragraph(
        doc,
        "고객은 각 단계별로 초기 기술 적합성 검증(Go/No-Go), 무중단 인프라 이관, 라인 자동완성 지연시간 80ms 미만 달성, "
        "개발 생산성 40% 향상 및 PII 100% 자동 마스킹과 같은 실질적인 성과를 직접 확인할 수 있으며, "
        "ISV 파트너 펀딩 바우처를 통한 초기 PoC 비용 최소화, Broadcom VMware 라이선스 50% 절감액 인프라 재투자, "
        "외부 상용 API 대비 토큰 비용 60% 이상 절감이라는 강력한 비용 이득을 확실하게 확보할 수 있습니다."
    )

    t_roadmap = doc.add_table(rows=5, cols=5)
    t_roadmap.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_roadmap.autofit = False
    set_table_borders(t_roadmap)

    r_headers = [
        ("Phase 구분", Cm(2.4)),
        ("소요 기간", Cm(1.8)),
        ("협력 목적 및 핵심 활동 (위한 것)", Cm(4.6)),
        ("확인 가능한 핵심 성과 (성과 지표)", Cm(4.0)),
        ("확보 가능한 비용 이득 (Cost Benefits)", Cm(3.7))
    ]
    for col_idx, (h_text, w_val) in enumerate(r_headers):
        cell = t_roadmap.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=70, right=70)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.2)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    roadmap_data = [
        (
            "PHASE 01\n진단 & 파일럿",
            "2 ~ 4주",
            "• 사전 TCO 진단 및 10~20 VM 신속 실증\n• IT 인벤토리 분석 & 클라우드 송환 타당성 분석\n• 파트너 펀딩 지원금 신청 및 실증 랩 가동",
            "• 파일럿 기술 검증 성공 (Go/No-Go)\n• 시스템 전환 타당성 명세서 확보\n• 기술 적합성 100% 사전 확인",
            "초기 검증 비용 최소화\n(AWS/Intel/Nutanix\n파트너 펀딩 바우처)"
        ),
        (
            "PHASE 02\n코어 인프라 현대화",
            "1 ~ 3개월",
            "• VMware 비용 폭증 방어 및 인프라 가상화\n• Nutanix 무상 AHV / OpenShift 최적화\n• 무중단 워크로드 본 이관 및 NC2 하이브리드",
            "• 무중단 워크로드 이관 완결\n• GPU 가상화(vGPU/MIG) 슬라이싱 가동\n• 하이브리드 망분리 DR 체계 수립",
            "가상화 TCO 50% 절감\n(절감된 VMware 예산으로\nGPU 인프라 증설)"
        ),
        (
            "PHASE 03\nAI 플랫폼 & 모델",
            "1 ~ 2개월",
            "• 사내 GPU/NPU 소버린 모델 & 에이전트 가동\n• Dell PowerEdge + vLLM/NIM 기반 sLLM\n• 사내 RAG 파이프라인 및 코딩 에이전트 연동",
            "• 코딩 자동완성 응답 지연 < 80ms\n• 개발 생산성 40% 이상 향상\n• Spiceware PII 100% 자동 마스킹",
            "토큰 비용 60% 이상 절감\n(외부 상용 API 호출 비용\n대비 영구적 절감)"
        ),
        (
            "PHASE 04\n운영 관리 & 확장",
            "지속 운영",
            "• 24x7 무장애 관제 & 전사 AI Factory 확장\n• Dynatrace Managed 기반 AIOps 관제\n• 전사 부서별 Use Case 발굴 및 Scale-out",
            "• 서비스 SLA 99.9% 무장애 운영\n• 비즈니스 의사결정 속도 5배 향상\n• 전사 엔터프라이즈 AI 자산 내재화",
            "운영 TCO 40% 추가 절감\n(자동화 장애 복구 및\n규모의 경제 달성)"
        )
    ]

    for row_idx, r_data in enumerate(roadmap_data, start=1):
        row_cells = t_roadmap.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(2.4), Cm(1.8), Cm(4.6), Cm(4.0), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=60, right=60)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.6)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                if "01" in text: r.font.color.rgb = COLOR_MZC_BLUE
                elif "02" in text: r.font.color.rgb = COLOR_NVIDIA_GREEN
                elif "03" in text: r.font.color.rgb = COLOR_GOLD
                elif "04" in text: r.font.color.rgb = RGBColor(0x9D, 0x4E, 0xDD)
            elif col_idx == 4:
                r.bold = True
                r.font.color.rgb = COLOR_NVIDIA_GREEN

    # ── [제 7 장. AI Full Stack 솔루션 패키지 및 산업군별 적용 시나리오 (예)] ──
    add_section_header(doc, "7", "AI Full Stack 솔루션 패키지 및 산업군별 적용 시나리오 (예)",
                       "고객 비즈니스 목적 직결 4대 패키지 정의 및 기업 규모별 추정 사양 가이드")

    add_narrative_paragraph(
        doc,
        "메가존클라우드는 복잡한 물리 구성도 대신 고객의 실제 비즈니스 목적과 직결된 "
        "'4대 핵심 솔루션 패키지(코딩 에이전트, Articul8 산업 지식, 데이터 인텔리전스, 소버린 AI 팩토리)'를 정의하고, "
        "각 패키지별 기업 규모(50인/200인/500인+)에 따른 추정 인프라 사양 및 정량적 도입을 제안합니다."
    )

    # 4대 패키지 총괄 비교 요약 (불릿형 서술형)
    add_sub_header(doc, "■ 4대 핵심 비즈니스 솔루션 패키지 총괄 요약")

    add_narrative_paragraph(
        doc,
        "메가존클라우드가 제공하는 4대 핵심 비즈니스 솔루션 패키지의 주요 특성 및 정량적 도입 효과 요약은 다음과 같습니다."
    )

    pkg_summary_bullets = [
        (
            "• Package 01. 엔터프라이즈 사내 코딩 에이전트 & DevOps 패키지 (Coding Agent + GitLab): ",
            "IT/SW, 핀테크, 게임 및 첨단 제조 R&D를 위한 완전 폐쇄망(Air-Gap) 개발 생산성 및 DevSecOps 통합 솔루션입니다. GitLab (사내 CI/CD) + Qwen2.5-Coder-32B + vLLM + Continue.dev/MCP를 결합하여 소스코드 외부 유출을 0% 원천 차단하고, 초고속 라인 자동완성(FIM <80ms)으로 개발 생산성을 40% 이상 향상시킵니다. (권장 규모: Small ~50인 S/W 약 5,000만~6,000만 원 / Medium 100~300인 S/W 약 1.8억~2.5억 원)"
        ),
        (
            "• Package 02. Articul8 AI 산업 도메인 특화 지식 패키지 (Domain Intelligence: 제조·항공MRO / 금융): ",
            "제조·스마트팩토리(PdM), 항공우주 MRO 정밀 정비, 금융 여신/신용 심사 분야를 위한 턴키 지식 그래프 솔루션입니다. Articul8 Model Mesh + Knowledge Graph + Dell PowerEdge/NVIDIA + OpenShift를 결합하여 비정형 도면, 정비 매뉴얼, 재무 공시를 자동 구조화하고, 환각률을 95% 이상 억제하며 업무 분석 시간을 85% 이상 단축합니다. (권장 규모: Medium 50~200인 S/W 약 1.8억~2.5억 원 / Large 200~500인+ S/W 약 3.5억~5.5억 원)"
        ),
        (
            "• Package 03. 온프레미스 문서 & 데이터 인텔리전스 파이프라인 (Data RAG): ",
            "공공·금융 FDS, 규제 감사 보고서 자동화 및 대기업 법무/계약 검토를 위한 전사 온프레미스 데이터 자산화 솔루션입니다. MinIO (AI 오브젝트 스토리지) + Cohere Command R+ + Redis Enterprise를 결합하여 온프레미스 시맨틱 캐시를 통한 토큰 비용 40% 절감, 규제 감사 보고서 작성 시간 10시간 ➔ 15분 단축을 실현합니다. (권장 규모: Medium 50~200인 S/W 약 1.8억~2.5억 원 / Large 500인+ S/W 약 3.5억~5.0억 원)"
        ),
        (
            "• Package 04. 온프레미스 소버린 AI 팩토리 & PoC 랜딩 패키지 (AI Factory): ",
            "전사 AI 인프라를 내재화하려는 대기업 및 가상화 인프라 현대화 고객을 위한 풀스택 턴키 인프라 솔루션입니다. Dell PowerEdge XE9680 (8-GPU) + Dell PowerScale F900 + Nutanix AHV / Red Hat OpenShift + NVIDIA NIM을 결합하여 가상화 TCO 50% 절감 및 외부 API 대비 토큰 비용 60% 이상 절감, 24x7 SLA 무장애 관제를 제공합니다. (권장 규모: 파일럿 PoC 2~4주 펀딩 바우처 지원 / 프로덕션 팩토리 500인+ 사양별 맞춤 구성)"
        )
    ]

    for title_text, desc_text in pkg_summary_bullets:
        p_b = doc.add_paragraph()
        p_b.paragraph_format.space_before = Pt(2)
        p_b.paragraph_format.space_after = Pt(4)
        p_b.paragraph_format.left_indent = Cm(0.4)
        r_t = p_b.add_run(title_text)
        r_t.font.name = FONT_MAIN
        r_t.font.size = Pt(8.8)
        r_t.bold = True
        if "Package 01" in title_text: r_t.font.color.rgb = COLOR_MZC_BLUE
        elif "Package 02" in title_text: r_t.font.color.rgb = COLOR_DELL_BLUE
        elif "Package 03" in title_text: r_t.font.color.rgb = COLOR_GOLD
        elif "Package 04" in title_text: r_t.font.color.rgb = COLOR_NVIDIA_GREEN

        r_d = p_b.add_run(desc_text)
        r_d.font.name = FONT_MAIN
        r_d.font.size = Pt(8.8)
        r_d.font.color.rgb = COLOR_TEXT_MAIN

    add_sub_header(doc, "■ 4대 핵심 패키지별 세부 아키텍처 및 오퍼링 상세")

    package_items = [
        (
            "Package 01. 엔터프라이즈 코딩 에이전트 & DevOps 패키지 (Coding Agent + GitLab)",
            "IT/소프트웨어, 금융 핀테크, 게임, 첨단 전자 제조 R&D",
            "GitLab (사내 Git/CI/CD) + Qwen2.5-Coder-32B + vLLM 초고속 서빙 + Continue.dev / 사내 MCP Server",
            "• Small (개발자 ~50인): Dell PowerEdge R760 (2x RTX 6000 Ada or 리벨리온 NPU) + SNO (S/W 약 5,000만~6,000만 원, H/W 별도)\n"
            "• Medium/Large (개발자 100~300인+): Dell PowerEdge R760xa (4x L40S) + OpenShift 3-Node (S/W 약 1.8억~2.5억 원, H/W 별도)",
            "소스코드 외부 유출 0% 원천 차단(Air-Gap), 외부 SaaS 대비 비용 70% 절감, 라인 자동완성(FIM) 80ms 미만 초고속 응답, 개발 생산성 40% 이상 향상"
        ),
        (
            "Package 02. Articul8 산업 도메인 특화 지식 패키지 (Domain Intelligence: 제조·항공MRO / 금융)",
            "[제조·항공MRO] 설비 예지보전, MRO 항공 정비 | [금융] 여신심사, 재무분석, 사모신용",
            "Articul8 Model Mesh Orchestrator + Knowledge Graph Substrate + Dell PowerEdge + Red Hat OpenShift",
            "• Medium (50~200인 사업부): Dell PowerEdge R760xa (4x L40S) + 올플래시 NVMe (S/W 약 1.8억~2.5억 원, H/W 별도)\n"
            "• Large (200~500인+ 전사): Dell PowerEdge XE9640/XE9680 + Dell PowerScale F900 (S/W 약 3.5억~5.5억 원, H/W 별도)",
            "비정형 CAD 도면/매뉴얼/공시 데이터 지식 그래프 자동 구조화, 환각률(Hallucination) 95% 이상 억제, 도면 분석 시간 85% 단축, 미션 크리티컬 의사결정 정확도 99% 달성"
        ),
        (
            "Package 03. 온프레미스 문서 & 데이터 인텔리전스 파이프라인 (Data RAG Package)",
            "공공·금융(FDS 이상거래 탐지 & 규제 감사 보고서 자동화), 대기업 법무/인사/계약 관리",
            "MinIO AI Storage + Cohere Command R+ / Embed / Rerank + Redis Enterprise",
            "• Medium (50~200인): Dell PowerEdge R760xa (4 GPU) + MinIO 50TB + Redis Enterprise (S/W 약 1.8억~2.5억 원, H/W 별도)\n"
            "• Large (전사 500인+): Dell PowerEdge XE9680 + PowerScale + Cohere / Redis Enterprise 전사 라이선스 (S/W 약 3.5억~5.0억 원, H/W 별도)",
            "초저지연 시맨틱 캐시(토큰 비용 40% 절감), 망분리 금융/공공 규제 100% 충족, 사내 감사 보고서 작성 시간 10시간 ➔ 15분 단축"
        ),
        (
            "Package 04. 온프레미스 소버린 AI 팩토리 & PoC 랜딩 패키지 (AI Factory Package)",
            "전사 AI 전환 추진 대기업/금융지주, 가상화 인프라 현대화 고객",
            "Dell PowerEdge XE9680 (8x H100/B200) + Dell PowerScale F900 + Nutanix NCI/무상 AHV + NVIDIA NIM",
            "• Pilot / PoC (2~4주): 10~20 VM 단위 Landing Zone + 파트너 펀딩 바우처 수혜 (초기 비용 최소화)\n"
            "• Production Factory (500인+): Dell XE9680 8x H100/B200 + 400GbE 패브릭 + Dynatrace (맞춤형 턴키 견적)",
            "인프라 TCO 50% 절감액으로 GPU 인프라 증설 재투자, 외부 상용 API 대비 토큰 비용 60% 이상 절감, 24x7 SLA 무장애 관제 및 TPS 2배 증대"
        )
    ]

    for p_title, p_target, p_stack, p_sizing, p_effect in package_items:
        add_sub_header(doc, f"■ {p_title}")
        
        p_info = doc.add_paragraph()
        p_info.paragraph_format.space_before = Pt(1)
        p_info.paragraph_format.space_after = Pt(2)
        p_info.paragraph_format.left_indent = Cm(0.4)
        
        r_t1 = p_info.add_run("【적용 산업】 ")
        r_t1.font.name = FONT_MAIN
        r_t1.font.size = Pt(8.8)
        r_t1.bold = True
        r_t1.font.color.rgb = COLOR_DELL_BLUE
        
        r_v1 = p_info.add_run(f"{p_target}\n")
        r_v1.font.name = FONT_MAIN
        r_v1.font.size = Pt(8.8)
        r_v1.font.color.rgb = COLOR_TEXT_MAIN

        r_t2 = p_info.add_run("【핵심 기술 스택】 ")
        r_t2.font.name = FONT_MAIN
        r_t2.font.size = Pt(8.8)
        r_t2.bold = True
        r_t2.font.color.rgb = COLOR_GOLD

        r_v2 = p_info.add_run(f"{p_stack}\n")
        r_v2.font.name = FONT_MAIN
        r_v2.font.size = Pt(8.8)
        r_v2.font.color.rgb = COLOR_TEXT_MAIN

        r_t3 = p_info.add_run("【기업 규모별 추정 사양 & S/W TCO】\n")
        r_t3.font.name = FONT_MAIN
        r_t3.font.size = Pt(8.8)
        r_t3.bold = True
        r_t3.font.color.rgb = COLOR_NVIDIA_GREEN

        r_v3 = p_info.add_run(f"{p_sizing}\n")
        r_v3.font.name = FONT_MAIN
        r_v3.font.size = Pt(8.5)
        r_v3.font.color.rgb = COLOR_TEXT_MAIN

        r_t4 = p_info.add_run("【정량적 도입 효과】 ")
        r_t4.font.name = FONT_MAIN
        r_t4.font.size = Pt(8.8)
        r_t4.bold = True
        r_t4.font.color.rgb = RGBColor(0x9D, 0x4E, 0xDD)

        r_v4 = p_info.add_run(f"{p_effect}")
        r_v4.font.name = FONT_MAIN
        r_v4.font.size = Pt(8.8)
        r_v4.font.color.rgb = COLOR_TEXT_MAIN

    # Articul8 산업 도메인 특화 3대 Use Case 패키지 상세
    add_sub_header(doc, "■ Articul8 산업 도메인 특화 3대 Use Case 패키지 상세 (제조·항공MRO / 금융)")

    articul8_usecases = [
        (
            "Use Case 01. 설비 이상 감지 및 자율 예지보전 AI 패키지 (제조 - Predictive Maintenance)",
            "대형 플랜트/제조 시설 (연간 $250M+ 가동 중단 손실 방지)",
            "• Data-Aware Agent: 센서 로그, 정비 이력, 매뉴얼/도면 데이터 자율 통합\n• Manufacturing Domain Agent: 기계/물리 도메인 지식 기반 고장 예측 및 인과관계 자동 재구성\n• Report Generation Agent: 단계별 조치 제안 진단 요약 보고서 자동 발행",
            "비계획 다운타임 40% 이상 감축, 멀티모달 근본 원인 규명 및 고장 전 사전 선제 조치"
        ),
        (
            "Use Case 04. MRO 프로세스 인텔리전스 & 정밀 정비 AI 패키지 (항공우주/방산 - MRO Aerospace)",
            "항공우주/방산 MRO 및 복잡 조립 공정 (100% 규제 준수 & 감사 추적성 요구)",
            "• Data-Aware Agent: PDF, 스캔 매뉴얼, 부품 허용 오차 및 지침 자동 수집\n• Domain-Specific Agent: 도구/구성요소 식별, 작업 순서 및 안전 경고 연결\n• Traceability Assurance Agent: 원본 문서 및 수정 이력과 연결된 100% 감사 가능 증빙 생성",
            "매뉴얼 해석 정확도 92% 달성, 기술자 온보딩 기간 단축, 서비스 복귀 시간 단축"
        ),
        (
            "Use Case 10. 금융 서비스 리서치 & 여신 심사 AI 패키지 (금융/투자 - Financial Services)",
            "투자은행, 자산운용사, 사모펀드, 여신 심사 부서 (방대한 금융 공시/실적 분석)",
            "• Financial Research Assistant: 100만 건 이상 재무 문서/SEC 자료에서 투자 메모 자동 생성\n• Macro Investment Research: GDP/인플레이션/통화정책 등 거시경제 분석 보고서 95% 자동 완성\n• Private Credit Decisioning: CIM·실사 보고서 분석 및 사모 신용/M&A 여신 심사 지원",
            "실사 소요 기간 '수 주 ➔ 수 분' 단축, 100% 출처 기반 감사 추적성 확보"
        )
    ]

    for u_title, u_target, u_pipeline, u_effect in articul8_usecases:
        p_u = doc.add_paragraph()
        p_u.paragraph_format.space_before = Pt(6)
        p_u.paragraph_format.space_after = Pt(2)
        p_u.paragraph_format.left_indent = Cm(0.4)
        
        r_ut = p_u.add_run(f"• {u_title}\n")
        r_ut.font.name = FONT_MAIN
        r_ut.font.size = Pt(9.0)
        r_ut.bold = True
        r_ut.font.color.rgb = COLOR_DELL_BLUE

        r_tg_l = p_u.add_run("  - 대상 도메인: ")
        r_tg_l.font.name = FONT_MAIN
        r_tg_l.font.size = Pt(8.5)
        r_tg_l.bold = True
        r_tg_v = p_u.add_run(f"{u_target}\n")
        r_tg_v.font.name = FONT_MAIN
        r_tg_v.font.size = Pt(8.5)

        r_pl_l = p_u.add_run("  - 3단계 핵심 에이전트/솔루션 파이프라인:\n")
        r_pl_l.font.name = FONT_MAIN
        r_pl_l.font.size = Pt(8.5)
        r_pl_l.bold = True
        r_pl_v = p_u.add_run(f"    {u_pipeline.replace(chr(10), chr(10) + '    ')}\n")
        r_pl_v.font.name = FONT_MAIN
        r_pl_v.font.size = Pt(8.3)

        r_ef_l = p_u.add_run("  - 핵심 기대 효과: ")
        r_ef_l.font.name = FONT_MAIN
        r_ef_l.font.size = Pt(8.5)
        r_ef_l.bold = True
        r_ef_v = p_u.add_run(f"{u_effect}")
        r_ef_v.font.name = FONT_MAIN
        r_ef_v.font.size = Pt(8.5)

    # ── [제 8 장. 향후 실행 계획] ─────────────────────────────────────────
    add_section_header(doc, "8", "향후 실행 계획",
                       "전사 세일즈 오퍼링 가이드 확립, 파트너 펀딩 풀 확보 및 차기 보고 대비 핵심 과제")

    add_narrative_paragraph(
        doc,
        "본 MZC 엔터프라이즈 AI Full Stack 전략은 메가존클라우드가 단순한 CSP 리셀러의 한계를 넘어, "
        "급성장하는 하이브리드 온프레미스 AI 인프라 및 소버린 솔루션 시장을 선도하는 '최상위 AI 시스템 인티그레이터(AI SI)'로 도약하기 위한 핵심 전략 이니셔티브입니다. "
        "전사 세일즈 오퍼링 가이드 확립과 성공적인 사업 런칭, 그리고 차기 CEO 보고를 대비하여 다음 4대 핵심 실행 과제(4-Action)를 즉시 착수하며, 각 과제별 핵심 실행 목표 및 추진 방향은 다음과 같습니다."
    )

    action_summary_bullets = [
        (
            "• Action 01. 전사 세일즈 오퍼링 가이드 & 브리핑 킷 배포: ",
            "4대 핵심 패키지(코딩 에이전트, Articul8, 데이터 인텔리전스, AI Factory)별 고객 표준 제안서, TCO 계산기 및 규모별(Small/Medium/Large) 맞춤 견적 가이드북을 표준화하고, 전사 영업·프리세일즈 대상 격주 정례 교육을 가동합니다."
        ),
        (
            "• Action 02. ISV 파트너십 번들링 & PoC 펀딩 바우처 풀 확보: ",
            "Dell(H/W), Nutanix/Red Hat(OS), MinIO/Redis(스토리지/캐시), Articul8/Cohere(솔루션)와의 온프레미스 공동 번들 협력을 추진하고, 파트너 마케팅 펀드(MDF) 및 PoC 지원금을 사전 확보하여 고객의 초기 도입 진입 장벽을 제거합니다."
        ),
        (
            "• Action 03. 사내 온프레미스 AI PoC 실증 랩 (Demo Lab) 구축: ",
            "Dell PowerEdge R760(2x RTX 6000 Ada / L40S) 기반 사내 실증 서버를 가동하여 vLLM + Qwen2.5-Coder 코딩 에이전트 및 MinIO/Redis RAG 파이프라인의 실시간 라이브 데모 시연과 고객 초청 핸즈온(Hands-on) 워크숍을 운영합니다."
        ),
        (
            "• Action 04. 산업별 등대 고객 (Lighthouse) 1~2개사 조기 확보 및 차기 보고: ",
            "금융권(FDS/망분리 감사) 및 제조 대기업(사내 코딩에이전트/ERP 연동) 1차 고객사를 조기 발굴하여 파일럿 수주 및 정량적 TCO 절감 실증 데이터를 확보하고, 차기 CEO 보고 안건으로 상정하여 전사 사업 확장을 가속화합니다."
        )
    ]

    for title_text, desc_text in action_summary_bullets:
        p_ab = doc.add_paragraph()
        p_ab.paragraph_format.space_before = Pt(2)
        p_ab.paragraph_format.space_after = Pt(4)
        p_ab.paragraph_format.left_indent = Cm(0.4)
        r_at = p_ab.add_run(title_text)
        r_at.font.name = FONT_MAIN
        r_at.font.size = Pt(8.8)
        r_at.bold = True
        if "Action 01" in title_text: r_at.font.color.rgb = COLOR_MZC_BLUE
        elif "Action 02" in title_text: r_at.font.color.rgb = COLOR_NVIDIA_GREEN
        elif "Action 03" in title_text: r_at.font.color.rgb = COLOR_GOLD
        elif "Action 04" in title_text: r_at.font.color.rgb = RGBColor(0x9D, 0x4E, 0xDD)

        r_ad = p_ab.add_run(desc_text)
        r_ad.font.name = FONT_MAIN
        r_ad.font.size = Pt(8.8)
        r_ad.font.color.rgb = COLOR_TEXT_MAIN

    add_sub_header(doc, "■ 4대 핵심 실행 과제별 세부 추진 계획")

    action_plans = [
        (
            "Action 01. 전사 세일즈 오퍼링 가이드 & 브리핑 킷 배포",
            "• 4대 패키지별 고객 표준 제안서 & TCO 계산기 완성: 코딩 에이전트, Articul8, 데이터 인텔리전스, AI Factory 표준 템플릿 완비\n"
            "• 규모별 맞춤 견적 가이드북: Small (50인 이하 1억 미만), Medium (3~4.5억), Large (8~15억+) 견적 프로세스 내재화\n"
            "• 전사 영업/프리세일즈 교육: AI 인프라 및 ISV 솔루션 결합 세일즈 브리핑 세션 격주 정례화"
        ),
        (
            "Action 02. ISV 파트너십 번들링 & PoC 펀딩 바우처 풀 확보",
            "• 온프레미스 벤더 공동 번들 협력: Dell(H/W), Nutanix/Red Hat(OS), MinIO/Redis(스토리지/캐시), Articul8/Cohere(솔루션)\n"
            "• 고객 무상 PoC 펀딩 바우처 확보: 파트너 마케팅 펀드(MDF) 및 PoC 지원금을 사전 확정하여 초기 진입 장벽 제거\n"
            "• 원스톱 유지보수 SLA 체계 수립: 복수 ISV 솔루션의 단일 창구 기술지원(Tier 1/2) 체계 마련"
        ),
        (
            "Action 03. 사내 온프레미스 AI PoC 실증 랩 (Demo Lab) 구축",
            "• 실증 하드웨어 인프라 확보: Dell PowerEdge R760 (2x RTX 6000 Ada / L40S) 기반 온프레미스 실증 서버 가동\n"
            "• 라이브 데모 시연 환경 완비: vLLM + Qwen2.5-Coder 코딩 에이전트 실시간 시연 및 MinIO/Redis RAG 파이프라인 가동\n"
            "• 고객 초청 기술 세미나 및 Hands-on 워크숍: 주 1~2회 고객사 C-Level 및 테크 리드 초청 체험 세션 운영"
        ),
        (
            "Action 04. 산업별 등대 고객 (Lighthouse) 1~2개사 조기 확보 및 차기 보고",
            "• Top Account 타깃팅: 금융권(FDS/망분리 감사) 및 제조 대기업(사내 코딩에이전트/ERP 연동) 1차 고객사 조기 발굴\n"
            "• Phase 1 파일럿 수주 및 실증 데이터 수집: 실제 고객 현장 배포를 통한 벤치마크 및 TCO 절감 정량 데이터 확보\n"
            "• 차기 CEO 보고 안건 상정: 파일럿 실증 성과, 1차 수주 실적, 전사 확장 로드맵 및 사업 손익 현황 보고"
        )
    ]

    for a_title, a_desc in action_plans:
        add_sub_header(doc, f"• {a_title}")
        p_act = doc.add_paragraph()
        p_act.paragraph_format.space_before = Pt(1)
        p_act.paragraph_format.space_after = Pt(6)
        p_act.paragraph_format.left_indent = Cm(0.4)
        r_ad = p_act.add_run(a_desc)
        r_ad.font.name = FONT_MAIN
        r_ad.font.size = Pt(8.8)
        r_ad.font.color.rgb = COLOR_TEXT_MAIN

    # ── [부록 1. 계층별 핵심 솔루션 특징, 적용 분야 및 평가 매트릭스] ──────
    p_app_header = doc.add_paragraph()
    p_app_header.paragraph_format.space_before = Pt(28)
    p_app_header.paragraph_format.space_after = Pt(4)
    p_app_header.paragraph_format.keep_with_next = True
    
    r_app_tag = p_app_header.add_run("[부록] Appendix.  ")
    r_app_tag.font.name = FONT_MAIN
    r_app_tag.font.size = Pt(14.5)
    r_app_tag.bold = True
    r_app_tag.font.color.rgb = COLOR_GOLD
    
    r_app_title = p_app_header.add_run("MZC AI Full Stack 핵심 솔루션별 특징, 적용 분야 및 평가 매트릭스")
    r_app_title.font.name = FONT_MAIN
    r_app_title.font.size = Pt(14.5)
    r_app_title.bold = True
    r_app_title.font.color.rgb = COLOR_NAVY_DARK

    p_app_sub = doc.add_paragraph()
    p_app_sub.paragraph_format.space_before = Pt(0)
    p_app_sub.paragraph_format.space_after = Pt(10)
    p_app_sub.paragraph_format.keep_with_next = True
    r_app_s = p_app_sub.add_run("MZC ISSU ISV 포트폴리오의 시장 검증도, 기술력, 온프레미스/폐쇄망 지원력 및 TCO 분석 기반 핵심 솔루션별 특장점 및 적용 분야 정리")
    r_app_s.font.name = FONT_MAIN
    r_app_s.font.size = Pt(9.5)
    r_app_s.font.color.rgb = COLOR_TEXT_MUTED
    r_app_s.italic = True

    add_narrative_paragraph(
        doc,
        "MZC ISSU의 글로벌 ISV 파트너십 데이터베이스와 엔터프라이즈 시장 검증도(레퍼런스, 온프레미스/폐쇄망 지원력, 기술적 처리 성능, TCO)를 바탕으로 "
        "Layer 04(비즈니스 솔루션) 및 Layer 03(미들웨어·개발·데이터·AIOps) 핵심 솔루션의 세부 특징과 적용 분야를 체계적으로 수록합니다."
    )

    # [부록 1-1] Layer 04 솔루션 특징 및 분야
    add_sub_header(doc, "Layer 04: Solution Layer (엔터프라이즈 AI 비즈니스 솔루션)")
    t_app_l4 = doc.add_table(rows=3, cols=4)
    t_app_l4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_app_l4.autofit = False
    set_table_borders(t_app_l4)

    l4_headers = [("솔루션 (ISV)", Cm(3.2)), ("적용 비즈니스 분야", Cm(3.8)), ("핵심 기술 특징 및 경쟁 우위", Cm(5.8)), ("권장 구축 환경", Cm(3.7))]
    for col_idx, (h_text, w_val) in enumerate(l4_headers):
        cell = t_app_l4.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    l4_data = [
        ("Articul8 AI\n(Intel Spinoff)", "제조(스마트팩토리 PdM)\n항공우주 MRO 정밀정비\n금융(여신/공시 리서치)", "• 산업 도메인 특화 GenAI & 에이전틱 지식 그래프 플랫폼\n• Model Mesh + Knowledge Graph Substrate 결합\n• 비정형 도면·매뉴얼·공시문서 자동 구조화, 환각률 95% 억제 및 분석 시간 85% 단축\n• 100% 사내 폐쇄망 턴키 배포 보장", "사내 온프레미스 GPU/NPU 클러스터, Red Hat OpenShift, Dell XE9680"),
        ("Cohere\n(Command R+ / Embed / Rerank)", "금융·보험·투자\n공공·행정·국방\n사내 RAG 검색 포털", "• 사내 구축형 엔터프라이즈 소버린 파운데이션 모델 사실상 표준\n• Command R+: 다국어 및 복잡한 기업 문서 비즈니스 추론 특화\n• Embed v3 & Rerank 3.5: 글로벌 최고 수준의 RAG 임베딩 및 고정밀 재순위화\n• 사내 폐쇄망 내 데이터 주권(Data Sovereignty) 완벽 보호", "사내 온프레미스 고성능 GPU 서버, Nutanix / OpenShift 컨테이너")
    ]
    for row_idx, r_data in enumerate(l4_data, start=1):
        row_cells = t_app_l4.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(3.8), Cm(5.8), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0x9D, 0x4E, 0xDD)

    # [부록 1-2] Layer 03 분야별 솔루션 매트릭스
    # 1. 고속 추론 서빙 & 라우팅
    add_sub_header(doc, "Layer 03 - 1. 초고속 추론 서빙 엔진 & 모델 라우팅 (Inference Serving & Routing)")
    t_app1 = doc.add_table(rows=3, cols=4)
    t_app1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_app1.autofit = False
    set_table_borders(t_app1)

    app1_headers = [("솔루션 (ISV)", Cm(3.2)), ("적용 분야", Cm(3.8)), ("핵심 기술 특징 및 경쟁 우위", Cm(5.8)), ("권장 워크로드", Cm(3.7))]
    for col_idx, (h_text, w_val) in enumerate(app1_headers):
        cell = t_app1.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    app1_data = [
        ("vLLM\n(PagedAttention)", "초고속 LLM/vLM 추론 서빙", "• 글로벌 오픈소스 고속 추론 사실상 표준 엔진\n• PagedAttention 기반 KV 캐시 메모리 낭비 0% 실현, 고성능 배치 추론\n• 최신 오픈소스 모델(sLLM, vLM) 및 리벨리온, 퓨리오사AI NPU 완벽 호환", "대규모 동시접속 챗봇, 사내 코딩 에이전트 FIM 서빙, 사내 RAG"),
        ("LiteLLM", "모델 라우팅 & API 게이트웨이", "• 100+ LLM 통합 OpenAI 규격 지능형 프록시 게이트웨이\n• 사내 vLLM, 소버린 모델, 외부 API를 단일 인터페이스로 일원화\n• 모델별 스마트 로드밸런싱, Fallback, 부서별 토큰/비용 레이트리밋 통제", "사내 통합 AI 게이트웨이, 멀티 모델 라우팅, API 보안 거버넌스")
    ]
    for row_idx, r_data in enumerate(app1_data, start=1):
        row_cells = t_app1.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(3.8), Cm(5.8), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_GOLD

    # 2. 사내 소버린 모델 & 개발 생산성
    add_sub_header(doc, "Layer 03 - 2. 사내 소버린 모델 & 개발 생산성 / DevSecOps (Model & Developer Productivity)")
    t_app2 = doc.add_table(rows=5, cols=4)
    t_app2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_app2.autofit = False
    set_table_borders(t_app2)

    for col_idx, (h_text, w_val) in enumerate(app1_headers):
        cell = t_app2.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    app2_data = [
        ("sLLM\n(Qwen2.5 / Llama 3.3)", "사내 소버린 경량 언어 모델", "• 고성능 경량 파운데이션 모델 (7B~72B)\n• 사내 폐쇄망 GPU 서버에서 외부 통신 없이 자체 파인튜닝 및 온프레미스 구동\n• 코딩(Qwen2.5-Coder), 보고서 요약, 질의응답 비즈니스 로직 최적화", "사내 코딩 보조, 업무 문서 자동 요약, 보안 민감 텍스트 처리"),
        ("vLM\n(Qwen2-VL / Llama-Vision)", "멀티모달 비전 언어 모델", "• 고해상도 이미지, 도면, 복합 문서 시각 분석\n• 비정형 설계도면, 설비 현장 사진, 복잡한 차트 및 표 데이터를 시각적으로 이해하고 구조화\n• 온프레미스 vLLM 기반 고속 멀티모달 추론", "제조 도면 OCR/구조화, 설비 크랙/결함 시각 검사, 영수증/서식 인식"),
        ("MCP (코딩 에이전트)\n(Model Context Protocol)", "사내 AI 코딩 에이전트 & IDE 연동", "• Model Context Protocol 기반 사내 코드/문서 컨텍스트 실시간 바인딩\n• VS Code, Cursor, JetBrains IDE 연동 및 FIM(<80ms) 초고속 라인 자동완성\n• 사내 개발 생산성 40% 향상 및 코드 보안 누출 방지", "사내 개발자 IDE 자동완성, 코드 리뷰 에이전트, 단위 테스트 생성"),
        ("GitLab", "사내 DevSecOps & CI/CD", "• 엔터프라이즈 사내 소스코드 관리 및 DevSecOps 파이프라인 표준\n• 완전 폐쇄망 설치형 Git 리포지토리, SAST 보안 취약점 사전 스캔, 무중단 자동 배포\n• AI 코딩 에이전트 및 MLOps 파이프라인과의 완벽한 네이티브 연계", "전사 소스코드 단일 저장소, 보안 빌드/배포 자동화, AI 수명주기 관리")
    ]
    for row_idx, r_data in enumerate(app2_data, start=1):
        row_cells = t_app2.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(3.8), Cm(5.8), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_NVIDIA_GREEN

    # 3. 데이터 스토리지 & 스트리밍
    add_sub_header(doc, "Layer 03 - 3. 온프레미스 AI 스토리지 & 실시간 데이터 파이프라인 (AI Storage & Data Pipeline)")
    t_app3 = doc.add_table(rows=4, cols=4)
    t_app3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_app3.autofit = False
    set_table_borders(t_app3)

    for col_idx, (h_text, w_val) in enumerate(app1_headers):
        cell = t_app3.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    app3_data = [
        ("MinIO", "고성능 AI 오브젝트 스토리지", "• 온프레미스 고성능 S3 호환 오브젝트 스토리지 사실상 표준\n• 대규모 비정형 문서, 멀티모달 이미지 데이터셋 초고속 병렬 I/O 처리\n• 완전 폐쇄망(Air-Gap) 배포 및 쿠버네티스 네이티브 아키텍처", "비정형 문서 RAG 저장소, 멀티모달 AI 데이터셋, 모델 아티팩트 백업"),
        ("Redis Enterprise", "초저지연 시맨틱 캐시 & 벡터", "• 인메모리 초저지연 벡터 검색 및 LLM 시맨틱 캐싱(Semantic Cache) 1위\n• 반복 질문 시 GPU 추론을 건너뛰어 토큰 비용 40% 절감 및 밀리초(<5ms) 응답", "실시간 대화 세션 관리, RAG 시맨틱 캐시, 임베딩 초고속 색인"),
        ("Confluent (Kafka)", "실시간 이벤트 데이터 스트리밍", "• 엔터프라이즈 실시간 Kafka 이벤트 스트리밍 파이프라인\n• 기간계 RDBMS 및 ERP/MES 데이터의 실시간 CDC 연계로 AI 데이터 최신성 보장\n• 사내 온프레미스 완전 독립 클러스터 운영", "실시간 데이터 동기화 파이프라인, 금융/제조 엔터프라이즈 기간계 연동")
    ]
    for row_idx, r_data in enumerate(app3_data, start=1):
        row_cells = t_app3.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(3.8), Cm(5.8), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_GOLD

    # 4. AIOps & 옵저버빌리티 표
    add_sub_header(doc, "Layer 03 - 4. AIOps & 풀스택 옵저버빌리티 / 모니터링 (AIOps & Observability)")
    t_app4 = doc.add_table(rows=2, cols=4)
    t_app4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_app4.autofit = False
    set_table_borders(t_app4)

    for col_idx, (h_text, w_val) in enumerate(app1_headers):
        cell = t_app4.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=120, bottom=120, left=90, right=90)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(8.5)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    app4_data = [
        ("Dynatrace\n(Dynatrace Managed)", "온프레미스 풀스택 AIOps & APM", "• 인과관계 분석(Causal AI / Davis AI) 기반 자동 장애 근본원인 분석(RCA) 최강자\n• OneAgent 단일 에이전트로 K8s, GPU, LLM 트레이싱, 인프라 전계층 실시간 관제\n• 완전 폐쇄망(Air-Gap) 사내 온프레미스 전용 클러스터 완벽 지원", "금융·공공·제조 24x7 무장애 SLA 시스템, LLM 토큰/지연시간 정밀 트레이싱")
    ]
    for row_idx, r_data in enumerate(app4_data, start=1):
        row_cells = t_app4.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.2), Cm(3.8), Cm(5.8), Cm(3.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(7.8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                r.font.color.rgb = COLOR_DELL_BLUE

    # ── [부록 2. 주요 파트너 솔루션 공식 웹사이트 및 참고 문헌] ───────────
    doc.add_page_break()
    add_section_header(doc, "부록 2", "주요 파트너 솔루션 공식 웹사이트 및 기술 레퍼런스 (References)",
                       "MZC AI Full Stack 4개 계층(4-Layer) 핵심 파트너 솔루션 공식 홈페이지 및 제품 레퍼런스")

    add_narrative_paragraph(
        doc,
        "본 MZC AI Full Stack 아키텍처 서비스 전략 보고서에 수록된 4개 계층(4-Layer) 핵심 파트너 솔루션들의 공식 제품 홈페이지 및 기술 레퍼런스 목록입니다."
    )

    ref_groups = [
        ("■ Layer 04: Solution Layer (비즈니스 솔루션 & 엔터프라이즈 AI)", RGBColor(0x9D, 0x4E, 0xDD), [
            ("• Articul8 AI: ", "엔터프라이즈 생성형 AI Model Mesh & Knowledge Graph 플랫폼 (Intel 스핀오프) — https://articul8.ai"),
            ("• Cohere: ", "엔터프라이즈 소버린 파운데이션 모델 (Command R+, Embed, Rerank) — https://cohere.com")
        ]),
        ("■ Layer 03: LLM Model + Tool Layer (파운데이션 모델, 코딩 에이전트/DevOps, 데이터, 추론 & AIOps)", COLOR_GOLD, [
            ("• vLLM Project: ", "PagedAttention 기반 오픈소스 사실상 표준 초고속 LLM 추론 엔진 — https://github.com/vllm-project/vllm"),
            ("• LiteLLM: ", "100+ LLM 통합 OpenAI 규격 API 프록시 & 로드밸런싱 게이트웨이 — https://www.litellm.ai"),
            ("• sLLM (소버린 모델): ", "Qwen2.5-Coder, Llama 3.3 등 사내 구축형 고성능 경량 파운데이션 모델 — https://huggingface.co/Qwen"),
            ("• vLM (비전 모델): ", "Qwen2-VL, Llama-Vision 등 비정형 도면·이미지 멀티모달 분석 모델 — https://huggingface.co/models"),
            ("• MCP (코딩 에이전트): ", "Model Context Protocol 기반 사내 코드/문서 바인딩 및 IDE 자동완성 — https://modelcontextprotocol.io"),
            ("• GitLab: ", "엔터프라이즈 사내 Git 소스코드 관리 및 DevSecOps/CI-CD 플랫폼 — https://about.gitlab.com"),
            ("• MinIO: ", "고성능 S3 호환 엔터프라이즈 AI 오브젝트 스토리지 — https://min.io"),
            ("• Redis Enterprise: ", "인메모리 초저지연 벡터 검색 및 LLM 시맨틱 캐시 솔루션 — https://redis.io"),
            ("• Confluent: ", "엔터프라이즈 실시간 Kafka 스트리밍 데이터 파이프라인 — https://www.confluent.io"),
            ("• Dynatrace: ", "Davis AI 기반 인과관계 분석(Causal AI) 및 사내 폐쇄망 지원 풀스택 AIOps — https://www.dynatrace.com")
        ]),
        ("■ Layer 02 & 01: OS, 가상화 & 하드웨어 인프라 (H/W & Platform)", COLOR_NVIDIA_GREEN, [
            ("• Red Hat OpenShift: ", "엔터프라이즈 K8s & RHOAI SNO (Single Node OpenShift) — https://www.redhat.com/openshift"),
            ("• Nutanix: ", "NCI (무상 AHV 하이퍼바이저) & Nutanix NKP (Kubernetes Platform) — https://www.nutanix.com"),
            ("• Dell Technologies: ", "Dell PowerEdge AI 서버 & Dell PowerScale AI 스토리지 — https://www.dell.com/ai"),
            ("• NVIDIA Enterprise: ", "가속 컴퓨팅 GPU (B200/300, H100/200, L40S, RTX 6000 Ada/Blackwell) — https://www.nvidia.com/data-center/"),
            ("• Rebellions (리벨리온): ", "고성능 국산 AI 가속 NPU (ATOM, REBEL) — https://rebellions.ai"),
            ("• FuriosaAI (퓨리오사AI): ", "차세대 엔터프라이즈 LLM 가속 NPU (RNGD) — https://furiosa.ai")
        ])
    ]

    for g_title, g_color, g_items in ref_groups:
        add_sub_header(doc, g_title)
        for i_title, i_desc in g_items:
            p_i = doc.add_paragraph()
            p_i.paragraph_format.space_before = Pt(1)
            p_i.paragraph_format.space_after = Pt(3)
            p_i.paragraph_format.left_indent = Cm(0.4)
            r_it = p_i.add_run(i_title)
            r_it.font.name = FONT_MAIN
            r_it.font.size = Pt(8.5)
            r_it.bold = True
            r_it.font.color.rgb = g_color

            r_id = p_i.add_run(i_desc)
            r_id.font.name = FONT_MAIN
            r_id.font.size = Pt(8.5)
            r_id.font.color.rgb = COLOR_TEXT_MAIN

    # 저장 실행
    doc.save(str(output_path))
    print(f"✅ AI Full Stack 전략 서비스 보고서 DOCX 생성 완료: {output_path}")
    print(f"   파일 크기: {output_path.stat().st_size / 1024:.1f} KB")


def main():
    base_dir = Path(__file__).resolve().parent
    docx_dir = base_dir / "docx"
    docx_dir.mkdir(exist_ok=True)
    
    out_dated = docx_dir / "2026-08-20_MZC_AI_Fullstack_Strategy_Service_Report.docx"
    out_docs = base_dir.parent / "docs" / "2026-08-20_MZC_AI_Fullstack_Strategy_Service_Report.docx"

    build_coe_report(out_dated)
    build_coe_report(out_docs)


if __name__ == "__main__":
    main()
