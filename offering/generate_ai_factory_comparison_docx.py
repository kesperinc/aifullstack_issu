#!/usr/bin/env python3
"""
NVIDIA AI Factory vs Dell AI Factory vs MZC 모듈형 AI Fullstack 비교 구축 문서 전용 DOCX 생성기
- 2026.08 메가존클라우드 공식 전략 오퍼링 기준
- 특정 단일 솔루션 종속 탈피: '모듈형 4-Layer 엔터프라이즈 AI 아키텍처(Vendor-Agnostic Modular Fullstack)' 전면 적용
- Nutanix, Articul8, Cohere, Red Hat, MinIO, Redis Enterprise 등은 고객 요구에 따른 선택형 생태계 옵션(Pluggable Options)으로 포지셔닝
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# ── 디자인 색상 팔레트 ──────────────────────────────────────────────
COLOR_NVIDIA_GREEN = RGBColor(0x76, 0xB9, 0x00)   # #76B900 (NVIDIA Green)
COLOR_DELL_BLUE     = RGBColor(0x00, 0x71, 0xC5)   # #0071C5 (Dell Blue)
COLOR_MZC_BLUE      = RGBColor(0x00, 0xAB, 0xF0)   # #00ABF0 (MZC Blue)
COLOR_NAVY_DARK     = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A (Deep Navy)
COLOR_PRIMARY_DARK  = RGBColor(0x1E, 0x29, 0x3B)   # #1E293B (Dark Slate)
COLOR_GOLD          = RGBColor(0xD4, 0xAF, 0x37)   # #D4AF37 (Gold)
COLOR_RED           = RGBColor(0xEE, 0x00, 0x00)   # #EE0000 (Red)
COLOR_TEXT_MAIN     = RGBColor(0x1F, 0x24, 0x2E)   # #1F242E (Main Body Text)
COLOR_TEXT_MUTED    = RGBColor(0x64, 0x74, 0x8B)   # #64748B (Muted)
COLOR_WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

HEX_BG_LIGHT_BLUE   = "F0F7FF"
HEX_BG_LIGHT_GREEN  = "F4FBF0"
HEX_BG_LIGHT_DELL   = "F0F6FC"
HEX_BG_LIGHT_GRAY   = "F8FAFC"
HEX_BG_DARK_NAVY    = "0F172A"
HEX_BORDER_LIGHT    = "CBD5E1"
HEX_BORDER_MZC      = "00ABF0"
HEX_BORDER_NVIDIA   = "76B900"
HEX_BORDER_DELL     = "0071C5"

FONT_MAIN = "맑은 고딕"
FONT_CODE = "Consolas"


def set_cell_shading(cell, color_hex):
    """테이블 셀 배경색 설정"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_margins(cell, top=140, bottom=140, left=160, right=160):
    """테이블 셀 안쪽 여백 설정 (dxa 단위)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)


def set_table_borders(table, color_hex=HEX_BORDER_LIGHT, sz="4", val="single"):
    """테이블 전체 테두리 설정"""
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)


def add_section_header(doc, number_str, title_str, subtitle_str=None):
    """대섹션 헤더 추가 (H1 스타일)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    r_num = p.add_run(f"[{number_str}] ")
    r_num.font.name = FONT_MAIN
    r_num.font.size = Pt(15)
    r_num.bold = True
    r_num.font.color.rgb = COLOR_MZC_BLUE
    
    r_title = p.add_run(title_str)
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(15)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    if subtitle_str:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(10)
        p_sub.paragraph_format.keep_with_next = True
        r_sub = p_sub.add_run(subtitle_str)
        r_sub.font.name = FONT_MAIN
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = COLOR_TEXT_MUTED
        r_sub.italic = True


def add_sub_header(doc, title_str, color=COLOR_PRIMARY_DARK):
    """중섹션 헤더 추가 (H2 스타일)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title_str)
    r.font.name = FONT_MAIN
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = color


def add_callout_box(doc, title, text_lines, border_color_hex="00ABF0", bg_color_hex="F0F7FF"):
    """강조 콜아웃 박스"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_shading(cell, bg_color_hex)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:color="{border_color_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    
    r_title = p.add_run(f"💡 {title}\n")
    r_title.bold = True
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(10.5)
    if border_color_hex == "00ABF0":
        r_title.font.color.rgb = COLOR_MZC_BLUE
    elif border_color_hex == "76B900":
        r_title.font.color.rgb = COLOR_NVIDIA_GREEN
    elif border_color_hex == "0071C5":
        r_title.font.color.rgb = COLOR_DELL_BLUE
    else:
        r_title.font.color.rgb = COLOR_NAVY_DARK

    for line in text_lines:
        p_line = cell.add_paragraph()
        p_line.paragraph_format.space_before = Pt(2)
        p_line.paragraph_format.space_after = Pt(2)
        r_line = p_line.add_run(line)
        r_line.font.name = FONT_MAIN
        r_line.font.size = Pt(9.5)
        r_line.font.color.rgb = COLOR_TEXT_MAIN

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)


def build_comparison_document(output_path):
    """모듈형·개방형 아키텍처 원칙이 적용된 비교 구축 DOCX 문서 생성"""
    doc = Document()
    
    # ── 기본 페이지 설정 (A4 여백 2.2cm) ──────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        
        # 헤더 / 푸터 설정
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("MegazoneCloud Modular Enterprise AI Fullstack Offering | Confidential")
        hrun.font.name = FONT_MAIN
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = COLOR_TEXT_MUTED
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("© 2026 MEGAZONECLOUD Corp. All Rights Reserved. | Vendor-Agnostic Modular AI Architecture")
        frun.font.name = FONT_MAIN
        frun.font.size = Pt(8)
        frun.font.color.rgb = COLOR_TEXT_MUTED

    # ── [표지 / 커버 헤더 영역] ──────────────────────────────────────────
    p_badge = doc.add_paragraph()
    p_badge.paragraph_format.space_before = Pt(10)
    p_badge.paragraph_format.space_after = Pt(6)
    r_b1 = p_badge.add_run("⚡ NVIDIA AI Factory (수직 통합)")
    r_b1.font.name = FONT_MAIN
    r_b1.font.size = Pt(9)
    r_b1.bold = True
    r_b1.font.color.rgb = COLOR_NVIDIA_GREEN
    
    r_sep1 = p_badge.add_run("   |   ")
    r_sep1.font.color.rgb = COLOR_TEXT_MUTED
    
    r_b2 = p_badge.add_run("💻 Dell AI Factory (턴키 패키지)")
    r_b2.font.name = FONT_MAIN
    r_b2.font.size = Pt(9)
    r_b2.bold = True
    r_b2.font.color.rgb = COLOR_DELL_BLUE
    
    r_sep2 = p_badge.add_run("   |   ")
    r_sep2.font.color.rgb = COLOR_TEXT_MUTED
    
    r_b3 = p_badge.add_run("🌐 MZC Modular 7-Layer (개방형·소버린)")
    r_b3.font.name = FONT_MAIN
    r_b3.font.size = Pt(9)
    r_b3.bold = True
    r_b3.font.color.rgb = COLOR_MZC_BLUE

    # 메인 타이틀
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("NVIDIA AI Factory vs Dell AI Factory vs MZC 모듈형 풀스택 비교 및 엔터프라이즈 구축 가이드")
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(20)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    # 부제목
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("단일 벤더 종속을 탈피한 개방형·모듈형(Vendor-Agnostic) 7-Layer 아키텍처 및 4대 Phase 구축 로드맵")
    r_sub.font.name = FONT_MAIN
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    # 메타 정보 박스
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(0)
    p_meta.paragraph_format.space_after = Pt(16)
    r_m = p_meta.add_run(
        "• 문서 버전: v3.0 (모듈형 개방형 아키텍처 전면 개정판)\n"
        "• 핵심 원칙: 특정 솔루션(Nutanix, Articul8, Cohere, Acryl 등)에 종속되지 않고, 고객 요구에 최적화된 Best-of-Breed 옵션을 결합하는 개방형 프레임워크\n"
        "• 발행 조직: 메가존클라우드 Integrated Solution Sales Unit (ISSU) / AI Solution Architect Team\n"
        "• 적용 범위: 공공·금융·제조·의료 엔터프라이즈 온프레미스 AI 인프라, 망분리 소버린 AI, 하이브리드 클라우드 AI"
    )
    r_m.font.name = FONT_MAIN
    r_m.font.size = Pt(8.5)
    r_m.font.color.rgb = COLOR_TEXT_MUTED

    # 구분선
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(0)
    p_hr.paragraph_format.space_after = Pt(12)
    r_hr = p_hr.add_run("━" * 58)
    r_hr.font.color.rgb = COLOR_MZC_BLUE
    r_hr.font.size = Pt(8)

    # ── [1. Executive Summary & 패러다임 전환] ──────────────────────────
    add_section_header(doc, "Section 01", "2026 AI 패러다임의 전환: AI Factory vs 개방형 모듈형 AI Fullstack", 
                       "단일 벤더 종속(Lock-in)을 극복하고 고객 맞춤형 Best-of-Breed 인프라를 실현하는 패러다임")

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(8)
    r_i = p_intro.add_run(
        "2026년 기업 AI 환경은 단순한 거대 모델 학습(LLM Training)을 넘어, 기업의 원천 데이터가 24/7 지속적인 인텔리전스(토큰, 에이전틱 태스크)로 생산되는 'AI Factory(AI 공장)' 시대로 진입했습니다. "
        "그러나 특정 단일 벤더의 수직 통합 스택이나 고정된 하드웨어 턴키 패키지만으로는 기업마다 상이한 레거시 인프라(서버/스토리지), 가상화 환경(VMware/OpenShift/Nutanix 등), "
        "사내 보안 규제(Air-Gap 망분리), 다양한 파운데이션 모델(오픈소스/Cohere/Acryl/Upstage 등)에 유연하게 대응하기 어렵습니다.\n\n"
        "메가존클라우드의 AI 풀스택은 특정 단일 상용 소프트웨어나 벤더에 종속되지 않는 **'개방형·모듈형 7-Layer 엔터프라이즈 AI 아키텍처(Vendor-Agnostic & Modular Architecture)'**를 지향합니다. "
        "각 계층(H/W, 가상화, 데이터, 모델, 오케스트레이션, UI, 보안)마다 검증된 다양한 솔루션 옵션들을 고객의 예산과 환경에 맞춰 자유롭게 조합(Pluggable)할 수 있도록 지원합니다."
    )
    r_i.font.name = FONT_MAIN
    r_i.font.size = Pt(9.5)
    r_i.font.color.rgb = COLOR_TEXT_MAIN

    # 3대 접근 방식 비교 표
    t_summary = doc.add_table(rows=4, cols=3)
    t_summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_summary.autofit = False
    set_table_borders(t_summary)

    headers = [
        ("⚡ 순수 NVIDIA AI Factory", HEX_BG_LIGHT_GREEN, COLOR_NVIDIA_GREEN),
        ("💻 Dell AI Factory with NVIDIA", HEX_BG_LIGHT_DELL, COLOR_DELL_BLUE),
        ("🌐 MZC Modular AI Fullstack", HEX_BG_LIGHT_BLUE, COLOR_MZC_BLUE)
    ]
    for col_idx, (h_text, bg_hex, font_clr) in enumerate(headers):
        cell = t_summary.cell(0, col_idx)
        cell.width = Cm(5.5)
        set_cell_shading(cell, bg_hex)
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.font.name = FONT_MAIN
        r.font.size = Pt(10)
        r.bold = True
        r.font.color.rgb = font_clr

    summary_rows = [
        [
            "【아키텍처 성격】\n• 단일 벤더 수직 통합(Vertical Lock-in)\n• 실리콘부터 최신 모델 가속(NIM)까지 초고속 GPU 파이프라인 제공",
            "【아키텍처 성격】\n• 엔터프라이즈 H/W 완제품형 턴키\n• PowerEdge 서버 + PowerScale 스토리지 + APEX 과금 일체화 패키지",
            "【아키텍처 성격】\n• 개방형·모듈형 소버린 AI 플랫폼\n• 벤더 종속 없이 각 계층별 최적 솔루션(Best-of-Breed) 자유 결합"
        ],
        [
            "【주요 구성 요소】\n• DGX B200 / GB200 NVL72\n• Quantum InfiniBand / Spectrum-X\n• DGX OS, Run:ai, NVAIE, NIM, NeMo",
            "【주요 구성 요소】\n• PowerEdge XE9680 / XE9640\n• PowerScale (DGX 인증 스토리지)\n• Dell Enterprise Hub on Hugging Face",
            "【주요 구성 요소】\n• H/W: Dell, HPE, NVIDIA, Intel, NPU\n• 가상화: OpenShift, Nutanix, VMware\n• 모델: 오픈소스, Cohere, Acryl, Solar\n• 오케스트레이션: Articul8, Dify, LangChain"
        ],
        [
            "【주요 한계 및 특성】\n• 고비용 NVAIE 연간 라이선스 부담\n• 자체 스토리지 부재로 서드파티 필수\n• 대규모 초기 일시불 CaPex 중심",
            "【주요 한계 및 특성】\n• 인프라 H/W 및 스토리지 패키지 중심\n• 사내 레거시 ERP/데이터 파이프라인 및 맞춤형 SI는 전문 파트너 협업 필요",
            "【차별화 강점】\n• 기존 인프라(서버/스토리지) 100% 재활용\n• 완전 폐쇄망(Air-Gap) 소버린 보안 충족\n• 도입 규모별 4대 Phase 점진적 확장"
        ]
    ]

    for row_idx, row_data in enumerate(summary_rows, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = t_summary.cell(row_idx, col_idx)
            cell.width = Cm(5.5)
            set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_TEXT_MAIN

    p_after_sum = doc.add_paragraph()
    p_after_sum.paragraph_format.space_before = Pt(8)

    add_callout_box(
        doc,
        "메가존클라우드의 핵심 철학: 'Zero Lock-in, Open Modular AI Governance'",
        [
            "1. 메가존클라우드는 특정 단일 솔루션(Nutanix, Articul8, Cohere 등)만을 전제로 인프라를 제한하지 않습니다. 이들은 고객의 워크로드와 예산에 맞춘 '선택 가능한 제안 옵션(Pluggable Offerings)' 중 하나입니다.",
            "2. 고객의 기존 투자 자산(Dell, HPE, VMware, OpenShift, 자체 데이터레이크 등)을 최대한 보호하면서, 필요한 계층마다 최적의 컴포넌트(NVIDIA NIM, MinIO, Redis Enterprise, 오픈소스 LLM, 상용 구축형 AI 등)를 유연하게 결합합니다.",
            "3. 메가존클라우드는 독립적인 최상위 시스템 인티그레이터(SI)이자 클라우드 매니지드 서비스 사업자(MSP)로서, 전체 계층의 아키텍처 통합, 보안 가드레일, 운영 관제를 완벽히 책임집니다."
        ],
        border_color_hex="00ABF0",
        bg_color_hex="F0F7FF"
    )

    # ── [2. 스택 구조 심층 해부: 엔비디아 5계층 vs MZC 모듈형 7-Layer] ──
    add_section_header(doc, "Section 02", "스택 구조 심층 해부: 엔비디아 5계층 vs MZC 모듈형 7-Layer",
                       "수직 통합 스택과 계층별 모듈형 개방 아키텍처의 상세 컴포넌트 매핑")

    add_sub_header(doc, "1. NVIDIA AI Factory 5-Tier Hierarchy (단일 벤더 가속 계층)")
    
    t_nv = doc.add_table(rows=6, cols=3)
    t_nv.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_nv.autofit = False
    set_table_borders(t_nv)

    nv_headers = [("계층 구분 (Tier)", Cm(3.8)), ("핵심 기술 및 컴포넌트", Cm(6.5)), ("역할 및 비즈니스 가치", Cm(6.2))]
    for col_idx, (h_title, w_val) in enumerate(nv_headers):
        cell = t_nv.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, "EBF5D8")
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(h_title)
        r.font.name = FONT_MAIN
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = COLOR_NVIDIA_GREEN

    nv_data = [
        ("Tier 5\nApplications & Agents", "• NVIDIA AI Blueprints\n• Omniverse (제조/디지털트윈)\n• Isaac (로보틱스) / Clara (의료)", "엔터프라이즈 RAG, 시각화, 도메인별 사전 훈련된 산업 청사진 제공"),
        ("Tier 4\nAI Engine & Microservices", "• NVIDIA NIM (추론 마이크로서비스)\n• NeMo Guardrails / NeMo Retriever\n• NeMo Curator / Customizer", "TensorRT-LLM 기반 컨테이너화 고속 추론, 가드레일 및 검색 증강 최적화"),
        ("Tier 3\nPlatform S/W & MLOps", "• NVIDIA AI Enterprise (NVAIE)\n• NVIDIA Run:ai (GPU 스케줄링)\n• DGX OS, Base Command, K8s Operator", "엔터프라이즈 런타임 라이선스, GPU 동적 파티셔닝(MIG) 및 워크로드 관리"),
        ("Tier 2\nAI Fabric & Network", "• Quantum-2 / X800 InfiniBand\n• Spectrum-X Ethernet (RoCEv2)\n• BlueField-3 DPU, NVLink (1.8TB/s)", "초저지연 무손실 패브릭, In-Network Computing(SHARP), GPU-to-GPU 고속 통신"),
        ("Tier 1\nAccelerated Compute", "• DGX B200 / GB200 NVL72 (수랭식)\n• HGX H100 / H200, Grace Hopper\n• GPUDirect Storage (GDS)", "Blackwell 아키텍처 기반 초대용량 고밀도 연산 처리 및 초고속 NVMe I/O")
    ]
    for row_idx, (t_tier, t_tech, t_val) in enumerate(nv_data, start=1):
        row_cells = t_nv.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([t_tier, t_tech, t_val], [Cm(3.8), Cm(6.5), Cm(6.2)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True

    p_mid_space = doc.add_paragraph()
    p_mid_space.paragraph_format.space_before = Pt(8)

    add_sub_header(doc, "2. MZC Modular 7-Layer Architecture (계층별 선택형 생태계 옵션)")

    t_mzc = doc.add_table(rows=9, cols=3)
    t_mzc.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_mzc.autofit = False
    set_table_borders(t_mzc)

    mzc_headers = [("계층 구분 (Layer)", Cm(3.8)), ("표준 기능 및 모듈 역할", Cm(5.8)), ("선택 가능한 주요 솔루션 옵션 (Pluggable Options)", Cm(6.9))]
    for col_idx, (h_title, w_val) in enumerate(mzc_headers):
        cell = t_mzc.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, "E0F2FE")
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        r = p.add_run(h_title)
        r.font.name = FONT_MAIN
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = COLOR_MZC_BLUE

    mzc_data = [
        (
            "Layer 7\nApplication & UI",
            "현업 업무용 UI 및 기간계 연동\n• 엔터프라이즈 AI 챗봇 / 검색 포털\n• 사내 코딩 에이전트 (IDE/CLI)\n• ERP / CRM / 그룹웨어 / AICC 결합",
            "• 메가존 사내 챗봇 UI 템플릿\n• 코딩 에이전트 (MCP / Continue.dev)\n• 엔터프라이즈 기간계(SAP/Oracle) 실시간 연동\n• 고객사 사내 레거시 포털 커스텀 SI"
        ),
        (
            "Layer 6\nOrchestration & Agentic",
            "지능형 모델 오케스트레이션 및 API 게이트웨이\n• 멀티 모델 지능형 라우팅\n• 엔터프라이즈 RAG 파이프라인\n• 에이전틱 워크플로우 제어",
            "• [옵션 1] Articul8 Model Mesh Orchestrator\n• [옵션 2] 오픈소스 LangChain / LlamaIndex / Dify\n• [옵션 3] 사내 보안 API Gateway & 커스텀 라우터"
        ),
        (
            "Layer 5\nFoundation Models & SLM",
            "사내 구축형 소버린 파운데이션 모델\n• 보안 데이터 외부 유출 없는 온프레미스 구동\n• 산업 도메인 특화 파인튜닝 (LoRA/QLoRA)",
            "• [오픈소스] Llama 3.3, Mistral, Qwen 2.5, DeepSeek\n• [구축형 상용] Cohere (Command R+ / Embed / Rerank)\n• [하이브리드] 사내 sLLM + 외부 클라우드 LLM 연계"
        ),
        (
            "Layer 4\nAI Serving & MLOps",
            "고속 추론 서빙 및 MLOps 라이프사이클\n• 실시간 토큰 생성 가속\n• 모델 모니터링, 버전 관리 및 리소스 제어",
            "• [추론 엔진] vLLM, NVIDIA NIM, Triton, TGI\n• [프레임워크] PyTorch, Ray, KServe\n• [MLOps 관제] MLflow, Kubeflow, Prometheus, Grafana"
        ),
        (
            "Layer 3\nData & Knowledge Platform",
            "정형/비정형 데이터 레이크 및 거버넌스\n• 사내 지식 베이스 벡터화\n• 전사 데이터 카탈로그 및 접근 제어",
            "• [데이터 플랫폼] MinIO (AI 오브젝트 스토리지), Redis Enterprise (시맨틱 캐시), Confluent (Kafka)\n• [Vector DB] Milvus, Qdrant, pgvector\n• [사내 RDBMS] Oracle, PostgreSQL, SAP HANA"
        ),
        (
            "Layer 2\nOS & Virtualization",
            "하이브리드 가상화 및 컨테이너 플랫폼\n• GPU 리소스 파티셔닝(MIG/vGPU)\n• VM 및 K8s 컨테이너 통합 운영",
            "• [옵션 1] Red Hat OpenShift / RHOAI (SNO)\n• [옵션 2] Nutanix NCI(무상 AHV) + Nutanix NKP\n• [옵션 3] 기존 VMware vSphere 환경 유지 연동\n• [옵션 4] Baremetal K8s / SUSE Rancher"
        ),
        (
            "Layer 1\nCompute, Storage & Fabric",
            "가속 컴퓨팅 서버 및 고성능 스토리지\n• 워크로드(학습 vs 추론)별 하드웨어 최적화\n• 고속 I/O 패브릭 구성",
            "• [서버] Dell PowerEdge, HPE ProLiant, Supermicro\n• [가속기] NVIDIA H100/L40S, 리벨리온, 퓨리오사AI\n• [스토리지] Dell PowerScale, NetApp"
        ),
        (
            "Cross-Cutting\nSecurity & Governance",
            "전 계층 관통 Zero-Trust 보안 가드레일\n• 망분리 및 데이터 주권(Data Sovereignty)\n• 개인정보 보호 및 규제 감사 준수",
            "• Air-Gapped 완전 폐쇄망 오프라인 패키징\n• 부서/직무별 세분화된 RBAC 권한 통제\n• PII 개인정보 자동 비식별화 필터링\n• 전 과정 감사 로그 및 응답 출처 추적성(Audit)"
        )
    ]
    for row_idx, (t_layer, t_func, t_opt) in enumerate(mzc_data, start=1):
        row_cells = t_mzc.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([t_layer, t_func, t_opt], [Cm(3.8), Cm(5.8), Cm(6.9)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            if row_idx == 8:  # 보안 강조
                set_cell_shading(cell, "FEF2F2")
            elif row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                if row_idx == 8:
                    r.font.color.rgb = COLOR_RED

    # ── [3. Dell AI Factory vs NVIDIA AI Factory vs MZC 개방형 역할] ──
    add_section_header(doc, "Section 03", "Dell AI Factory with NVIDIA vs 순수 NVIDIA AI Factory 정밀 비교",
                       "엔터프라이즈 H/W 패키지와 순수 GPU 슈퍼컴의 차이, 그리고 MZC의 개방형 통합 가치")

    p_dell_intro = doc.add_paragraph()
    p_dell_intro.paragraph_format.space_after = Pt(8)
    r_di = p_dell_intro.add_run(
        "순수 NVIDIA AI Factory와 Dell AI Factory with NVIDIA는 상호 대립하는 개념이 아닌, 엔비디아의 최신 가속 엔진을 델(Dell)의 완제품형 엔터프라이즈 섀시 및 "
        "전용 스토리지(PowerScale)와 결합한 '공동 턴키(Turnkey) 생태계'입니다. 메가존클라우드는 이 인프라 위에 특정 제품에 구애받지 않고 고객의 레거시 및 소버린 요구를 충족하는 "
        "독립적인 '최상위 시스템 인티그레이터(SI)이자 아키텍처 오케스트레이터' 역할을 수행합니다."
    )
    r_di.font.name = FONT_MAIN
    r_di.font.size = Pt(9.5)
    r_di.font.color.rgb = COLOR_TEXT_MAIN

    t_dell_comp = doc.add_table(rows=7, cols=4)
    t_dell_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_dell_comp.autofit = False
    set_table_borders(t_dell_comp)

    d_headers = [("비교 항목", Cm(2.8)), ("⚡ 순수 NVIDIA AI Factory", Cm(4.5)), ("💻 Dell AI Factory with NVIDIA", Cm(4.6)), ("🌐 MZC 개방형 역할 및 통합 가치", Cm(4.6))]
    for col_idx, (h_title, w_val) in enumerate(d_headers):
        cell = t_dell_comp.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, "E2E8F0")
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_title)
        r.font.name = FONT_MAIN
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = COLOR_NAVY_DARK

    d_data = [
        ("서버 폼팩터", "• 순수 DGX 시스템\n(DGX B200, GB200 NVL72)\n• 데이터센터 랙 일체형", "• Dell PowerEdge XE9680 / XE9640 (수랭식)\n• iDRAC 통합 서버 관리 체계", "• Dell, HPE 등 고객 선호 H/W 위에 최적 가상화 배포\n• NVIDIA GPU, 리벨리온, 퓨리오사AI 복합 구성"),
        ("스토리지 계층", "• 자체 스토리지 없음\n• 서드파티 파트너(GDS 인증) 별도 소싱 필요", "• Dell PowerScale (세계 최초 DGX 이더넷 인증)\n• PowerStore 일체형 제공", "• 사내 기존 스토리지 데이터 원본 Air-Gap 보호\n• MinIO / 다양한 데이터 파이프라인과 무중단 연동"),
        ("엣지 및 클라이언트", "• 데이터센터 및 클라우드 대규모 클러스터 중심", "• Dell Precision AI 워크스테이션\n• Dell Latitude AI PC & NativeEdge 엣지 관리", "• 개발자 로컬 PC/워크스테이션부터 프라이빗 데이터센터까지 엔드투엔드 연결 구축"),
        ("AI 소프트웨어", "• NVAIE, NIM, NeMo, Run:ai", "• NVAIE / NIM 탑재\n• Dell Enterprise Hub on Hugging Face", "• NIM 및 오픈소스 추론 엔진(vLLM)을 사내 폐쇄망 오프라인 환경에 맞게 커스텀 패키징"),
        ("과금 및 조달", "• 대규모 초기 CaPex 중심 일시불", "• Dell APEX (As-a-Service 사용량 과금) 지원", "• 4대 Phase 점진적 도입 로드맵 + 클라우드 펀딩 바우처 연계 + TCO 최적화"),
        ("레거시 SI 연동", "• 하드웨어 및 레퍼런스 청사진 제공 (SI 미수행)", "• Dell Professional Services for GenAI 지원", "• 엔터프라이즈 ERP/기간계 연동 + 사내 시스템 결합\n• VMware 유지/대체(OpenShift/AHV) 맞춤형 SI")
    ]

    for row_idx, (c_item, c_nv, c_dell, c_mzc) in enumerate(d_data, start=1):
        row_cells = t_dell_comp.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([c_item, c_nv, c_dell, c_mzc], [Cm(2.8), Cm(4.5), Cm(4.6), Cm(4.6)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True

    # ── [4. 7대 핵심 레이어 1:1 정밀 비교 매트릭스] ──────────────────────
    add_section_header(doc, "Section 04", "7대 핵심 레이어 1:1 정밀 비교 매트릭스",
                       "성능, 유연성, 보안, 거버넌스, TCO 관점에서의 종합 평가")

    t_matrix = doc.add_table(rows=9, cols=4)
    t_matrix.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_matrix.autofit = False
    set_table_borders(t_matrix)

    m_headers = [("비교 차원 (Dimension)", Cm(3.2)), ("⚡ NVIDIA AI Factory Stack", Cm(4.8)), ("🌐 MZC Modular AI Fullstack", Cm(5.5)), ("전략적 권장 방안", Cm(3.0))]
    for col_idx, (h_title, w_val) in enumerate(m_headers):
        cell = t_matrix.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_title)
        r.font.name = FONT_MAIN
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    matrix_data = [
        ("1. 하드웨어 & 가속기", "• NVIDIA 전용 실리콘 (Blackwell, H100)\n• 독점 생태계로 최고 성능 보장\n• 타사 가속기 믹스 불가", "• 멀티 벤더 가속기 지원 (NVIDIA + 리벨리온, 퓨리오사AI)\n• Dell / HPE / Supermicro 유연 구성\n• 워크로드(학습 vs 추론)별 가성비 믹스", "MZC 우위\n(비용·유연성)"),
        ("2. 네트워크 & 패브릭", "• Quantum-2 InfiniBand & Spectrum-X\n• NVLink 1.8TB/s, BlueField-3 DPU\n• 초거대 분산학습 최적화", "• 표준 400G RoCEv2 Ethernet & InfiniBand\n• DWDM 원격 DR 복제 및 사내망 인터페이스 연동\n• 기존 엔터프라이즈망 재활용", "NVIDIA 우위\n(분산학습 성능)"),
        ("3. 가상화 & 오케스트레이션", "• DGX OS + Run:ai + Base Command\n• GPU 파티셔닝(MIG) 및 동적 스케줄링\n• 순수 AI 컨테이너 집중", "• Red Hat OpenShift / Nutanix AHV / VMware 선택\n• 고객 라이선스 정책에 맞춘 유연한 가상화 전략\n• VM과 K8s 컨테이너 단일 통합 운영", "MZC 우위\n(환경 유연성)"),
        ("4. 데이터 플랫폼 & 레이크", "• GPUDirect Storage (GDS) & NeMo Curator\n• RAPIDS 가속 라이브러리\n• GPU I/O 가속 중심", "• MinIO AI 스토리지, Redis Enterprise 캐시, Confluent Kafka\n• Milvus, Qdrant 등 유연한 Vector DB 선택\n• Dell PowerScale 스토리지 직접 연동", "융합 권장\n(Storage+GDS)"),
        ("5. AI 모델 & 서빙 런타임", "• NVIDIA NIM (TensorRT-LLM, Triton)\n• 컨테이너 기반 1-Click 자동 최적화\n• 독자 NeMo 생태계 중심", "• A8, Cohere, 오픈소스 LLM 등 다양\n• vLLM, NVIDIA NIM, Triton 하이브리드 탑재\n• 모델 교체 및 파인튜닝 자유도 확보", "MZC 우위\n(모델 선택권)"),
        ("6. RAG, 에이전트 & UI", "• NeMo Framework & AI Blueprints\n• 기술 레퍼런스 및 청사진 제공", "• Articul8, Dify, LangChain/LlamaIndex 선택형 결합\n• 사내 코딩 에이전트(MCP) + Enterprise UI\n• 엔터프라이즈 ERP 및 기간계 직접 커스텀 SI", "MZC 우위\n(SI·비즈니스 연동)"),
        ("7. 보안 & 컴플라이언스", "• NVAIE 엔터프라이즈 패치\n• 클라우드/하이브리드 인증 요구\n• 오프라인 폐쇄망 구성 시 추가 공수", "• Air-Gapped 완전 폐쇄망 아키텍처\n• 부서별 RBAC 접근제어 & PII 자동 마스킹\n• 감사 로그 및 답변 출처 추적성 100%", "MZC 우위\n(금융·공공 규제)"),
        ("8. 도입 방식 & TCO", "• 대규모 초기 CaPex (수십억 원 대)\n• GPU당 고가 NVAIE 연간 구독료\n• 소규모 단계적 도입 시 단위비용 상승", "• 4대 Phase 점진적 로드맵 (PoC부터 단계적 확장)\n• 기존 H/W 보호 및 하이브리드 버스팅\n• 클라우드 송환으로 TCO 40~60% 절감", "MZC 압도적 우위\n(도입 장벽 완화)")
    ]

    for row_idx, (m_dim, m_nv, m_mzc, m_eval) in enumerate(matrix_data, start=1):
        row_cells = t_matrix.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([m_dim, m_nv, m_mzc, m_eval], [Cm(3.2), Cm(4.8), Cm(5.5), Cm(3.0)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
            elif col_idx == 3:
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if "MZC" in text:
                    r.font.color.rgb = COLOR_MZC_BLUE
                elif "NVIDIA" in text:
                    r.font.color.rgb = COLOR_NVIDIA_GREEN
                elif "융합" in text:
                    r.font.color.rgb = COLOR_GOLD

    # ── [5. MZC 모듈형 융합 레퍼런스 아키텍처] ──────────────────────────
    add_section_header(doc, "Section 05", "MZC 모듈형 융합 시너지 아키텍처 및 3대 결합 가치",
                       "고객 인프라에 맞춘 최적의 Best-of-Breed 조합과 통합 거버넌스")

    add_callout_box(
        doc,
        "하이브리드 모듈형 융합 아키텍처 원칙 (Modular Synergy Blueprint)",
        [
            "1. 멀티 추론 엔진 하이브리드 배포: NVIDIA NIM OCI 컨테이너와 오픈소스 vLLM/Triton을 고객의 컨테이너 플랫폼(OpenShift / Nutanix NKP / Vanilla K8s) 상에 병행 배포하여, 고성능 가속과 비용 효율성을 동시에 확보합니다.",
            "2. 데이터 플랫폼 & 초고속 패브릭 연계: 사내 Dell PowerScale/MinIO 데이터 스토리지와 NVIDIA GPUDirect Storage(GDS)를 직접 연계하여, 대용량 임베딩/파인튜닝 시 GPU 메모리 직접 전송을 지원합니다.",
            "3. 멀티 모델 & 지능형 라우팅: 오케스트레이션 계층에서 고객 요구에 따라 Articul8, Dify, LangChain 등을 적용하여, 단순 질의는 로컬 경량 sLLM/NPU로, 복잡한 고난도 추론은 최신 고성능 모델(NIM, Cohere 등)로 최적 라우팅합니다."
        ],
        border_color_hex="D4AF37",
        bg_color_hex="FFFDF0"
    )

    p_bene_title = doc.add_paragraph()
    p_bene_title.paragraph_format.space_before = Pt(8)
    p_bene_title.paragraph_format.space_after = Pt(4)
    r_bt = p_bene_title.add_run("■ 고객이 누리는 3대 핵심 결합 가치")
    r_bt.font.name = FONT_MAIN
    r_bt.font.size = Pt(11)
    r_bt.bold = True
    r_bt.font.color.rgb = COLOR_NAVY_DARK

    benefits = [
        ("1. 성능 극대화 (High Performance)", "엔비디아 NIM 및 최적화 엔진을 결합하여 추론 레이턴시 50% 단축 및 초당 토큰 생성량(TPS) 2배 향상"),
        ("2. 완벽한 주권 및 보안 (Sovereignty & Security)", "Air-Gapped 완전 폐쇄망과 PII 마스킹, 모델 가중치 사내 보관으로 금융·공공 보안 규정 100% 충족"),
        ("3. TCO 최적화 및 벤더 독립성 (Zero Lock-in)", "기존 H/W 자산 활용, 유연한 가상화 옵션 및 단계적 확장을 통해 불필요한 라이선스 낭비를 없애고 TCO 40~60% 절감")
    ]
    for b_title, b_desc in benefits:
        p_b = doc.add_paragraph()
        p_b.paragraph_format.space_before = Pt(2)
        p_b.paragraph_format.space_after = Pt(2)
        p_b.paragraph_format.left_indent = Cm(0.6)
        r1 = p_b.add_run(f"• {b_title}: ")
        r1.bold = True
        r1.font.name = FONT_MAIN
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_MZC_BLUE
        r2 = p_b.add_run(b_desc)
        r2.font.name = FONT_MAIN
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_TEXT_MAIN

    # ── [6. 4대 Phase 압축 로드맵 & 도입 패키지] ────────────────────────
    add_section_header(doc, "Section 06", "4대 Phase 압축 로드맵 & 단계별 도입 가이드",
                       "Small Start 파일럿 검증부터 전사 AI Factory 확장까지의 안전한 전환 여정")

    p_phase_intro = doc.add_paragraph()
    p_phase_intro.paragraph_format.space_after = Pt(8)
    r_pi = p_phase_intro.add_run(
        "엔터프라이즈 AI 전환의 위험을 최소화하기 위해, 고객의 의사결정 단계에 맞춘 4개의 핵심 Phase로 로드맵을 구성하였습니다. "
        "사전 진단 및 파일럿 검증(Phase 1)부터 시작하여 코어 인프라 현대화(Phase 2), AI 플랫폼 구축(Phase 3), 전사 운영 및 확장(Phase 4)으로 단계별 확장됩니다."
    )
    r_pi.font.name = FONT_MAIN
    r_pi.font.size = Pt(9.5)
    r_pi.font.color.rgb = COLOR_TEXT_MAIN

    t_phases = doc.add_table(rows=5, cols=4)
    t_phases.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_phases.autofit = False
    set_table_borders(t_phases)

    p_headers = [("Phase 구분", Cm(3.2)), ("주요 목표 및 기간", Cm(3.8)), ("핵심 활동 및 모듈 옵션", Cm(4.8)), ("주요 산출물 (Deliverables)", Cm(4.7))]
    for col_idx, (h_title, w_val) in enumerate(p_headers):
        cell = t_phases.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=140, bottom=140, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_title)
        r.font.name = FONT_MAIN
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = COLOR_WHITE

    phases_data = [
        (
            "PHASE 01\n진단 & 파일럿 검증",
            "• 소요 기간: 2~4주\n• 목표: 리스크 0% 기회 검증 및 파트너 펀딩 수혜",
            "• 인벤토리·종속성 분석 및 TCO 사전 진단\n• 소규모 기술 검증(PoC / Pilot Landing Zone)\n• 클라우드/파트너 펀딩 프로그램 지원금 신청",
            "• 전환 우선순위 진단 보고서\n• Pilot 검증 결과서 & Go/No-Go 판정\n• TCO 절감 및 지원금 적용 제안서"
        ),
        (
            "PHASE 02\n코어 인프라 현대화",
            "• 소요 기간: 1~3개월\n• 목표: 가상화 최적화, 무중단 본 이관 및 DR 체계 구축",
            "• 고객 맞춤 가상화(OpenShift / Nutanix / VMware) 구성\n• 자동화 도구 기반 무중단 데이터/VM 이관\n• 하이브리드 클라우드 연계 및 망분리 DR 수립",
            "• 무중단 마이그레이션 Runbook\n• 하이브리드 인프라 아키텍처 설계서\n• 재해복구(DR) 및 망분리 보안 정책서"
        ),
        (
            "PHASE 03\nAI 플랫폼 구축",
            "• 소요 기간: 1~2개월\n• 목표: 사내 소버린 AI 플랫폼 및 비즈니스 연동 가동",
            "• GPU/NPU 가속 인프라(Dell/HPE 등) 증설\n• 선택형 모델(오픈소스/Cohere/Acryl 등) 및 NIM 배포\n• RAG 파이프라인 및 사내 코딩 에이전트 가동",
            "• AI 전력/인프라 상세 설계서\n• 사내 RAG 및 에이전틱 파이프라인 구성서\n• 업무 시스템(ERP/CRM) 연동 명세서"
        ),
        (
            "PHASE 04\n운영 관리 & 전사 확장",
            "• 소요 기간: 지속 운영\n• 목표: 24x7 SLA 무장애 관제 및 전사 AI Factory 확장",
            "• 통합 모니터링 및 24x7 SLA 지원\n• 부서별 맞춤 Use Case 발굴 및 파인튜닝 지원\n• 전사 AI 클러스터 Scale-out 확장",
            "• 전사 통합 운영 매뉴얼 & 정기 리포트\n• 부서별 AI 확장 제안서\n• 전사 AI Factory 마스터 로드맵"
        )
    ]

    for row_idx, (ph_name, ph_goal, ph_pkgs, ph_deliv) in enumerate(phases_data, start=1):
        row_cells = t_phases.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip([ph_name, ph_goal, ph_pkgs, ph_deliv], [Cm(3.2), Cm(3.8), Cm(4.8), Cm(4.7)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=110, bottom=110, left=110, right=110)
            if row_idx % 2 == 1:
                set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = FONT_MAIN
            r.font.size = Pt(8.5)
            r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0:
                r.bold = True
                if row_idx == 1:
                    r.font.color.rgb = COLOR_MZC_BLUE
                elif row_idx == 2:
                    r.font.color.rgb = COLOR_PRIMARY_DARK
                elif row_idx == 3:
                    r.font.color.rgb = COLOR_NVIDIA_GREEN
                elif row_idx == 4:
                    r.font.color.rgb = COLOR_GOLD

    # ── [7. 산업군별 타깃 시나리오 및 모듈형 구축 방안] ──────────────────
    add_section_header(doc, "Section 07", "산업군별 타깃 시나리오 및 최적 모듈형 구축 방안",
                       "금융, 제조/ERP, 의료, 공공/국방, HPC, 미디어 등 6대 핵심 레퍼런스")

    usecases = [
        ("1. 백업 & 데이터 주권 재해복구 (DR & Air-Gap Vault)", "Veeam, Dell PowerProtect, Nutanix Sync, Air-Gap Vault", 
         "단일 클라우드 백업을 넘어 온프레미스와 다중 클라우드 환경에서 실시간 스냅샷 및 비동기 복제를 수행하여 랜섬웨어 감염 및 재난 발생 시 데이터 손실 없이 무중단 복구를 보장합니다."),
        ("2. 엔터프라이즈 ERP & 사내 기간계 결합", "엔터프라이즈 ERP, HPE GreenLake, Dell APEX, 프라이빗 클라우드",
         "해외 퍼블릭 클라우드로 핵심 비즈니스 데이터 반출 없이 사내 프라이빗 클라우드 형태로 구축하여 데이터 주권과 컴플라이언스를 완벽히 준수합니다."),
        ("3. 생명과학 & 헬스케어 (PACS & 바이오 연구)", "PACS 데이터 분석, 의료법 10년 보관, 사내 특화 LLM/NIM, NVIDIA Clara",
         "임상실험, 게놈연구, 암 연구를 위한 대외비 민감 정보를 사내에 격리 보관하고 의료정보 법적 10년 이상 장기보관 의무를 충족하며 고속 AI 분석을 수행합니다."),
        ("4. 금융 이상거래 탐지 (FDS) & 규제 준수 보고서", "MinIO/Redis, 사내 구축형 LLM(Cohere/Articul8), PII 자동 마스킹",
         "페타바이트(PB) 단위의 금융 거래 데이터에 대한 초고속 액세스와 망분리 환경 내 실시간 FDS 탐지 및 개인정보 자동 비식별화를 통해 규제 감사 보고서를 안전하게 자동 생성합니다."),
        ("5. HPC & 온사이트 AI 추론 토큰 비용 절감", "NVIDIA NIM, vLLM, 오픈소스 LLM, 리벨리온, 퓨리오사AI",
         "지속적인 외부 상용 LLM API 호출에 따른 토큰 과금 폭증을 방어하기 위해 사내 온사이트 GPU/NPU 클러스터를 구성하여 장기 운영 TCO를 60% 이상 절감합니다."),
        ("6. 미디어 & 엔터테인먼트 렌더링 / 생성형 협업", "Cloud Bursting, Dell PowerScale, AWS 하이브리드 연계",
         "고화질 3D 영상 렌더링 및 생성형 AI 비디오 제작 시 로컬 스토리지의 초고속 입출력과 퍼블릭 클라우드의 탄력적인 버스팅(Cloud Bursting)을 결합합니다.")
    ]

    for u_title, u_techs, u_desc in usecases:
        p_u = doc.add_paragraph()
        p_u.paragraph_format.space_before = Pt(8)
        p_u.paragraph_format.space_after = Pt(2)
        r_ut = p_u.add_run(f"■ {u_title}")
        r_ut.font.name = FONT_MAIN
        r_ut.font.size = Pt(10.5)
        r_ut.bold = True
        r_ut.font.color.rgb = COLOR_NAVY_DARK

        p_utag = doc.add_paragraph()
        p_utag.paragraph_format.space_before = Pt(0)
        p_utag.paragraph_format.space_after = Pt(2)
        p_utag.paragraph_format.left_indent = Cm(0.4)
        r_tag = p_utag.add_run(f"【추천 모듈 조합】 {u_techs}")
        r_tag.font.name = FONT_MAIN
        r_tag.font.size = Pt(8.5)
        r_tag.font.color.rgb = COLOR_DELL_BLUE
        r_tag.bold = True

        p_ud = doc.add_paragraph()
        p_ud.paragraph_format.space_before = Pt(0)
        p_ud.paragraph_format.space_after = Pt(6)
        p_ud.paragraph_format.left_indent = Cm(0.4)
        r_ud = p_ud.add_run(u_desc)
        r_ud.font.name = FONT_MAIN
        r_ud.font.size = Pt(9.5)
        r_ud.font.color.rgb = COLOR_TEXT_MAIN

    # ── [8. 의사결정 및 권장 액션 플랜] ─────────────────────────────────
    add_section_header(doc, "Section 08", "엔터프라이즈 도입 의사결정 가이드 & Next Steps",
                       "고객 유형별 최적 추진 경로 및 제안 문의")

    add_callout_box(
        doc,
        "고객 상황별 맞춤 추천 아키텍처 경로 (Decision Framework)",
        [
            "• 시나리오 A (초기 비용 최소화 & 빠른 실증): Phase 1 사전 진단 및 클라우드/파트너 펀딩을 활용하여 오픈소스 LLM + vLLM 기반 2~4주 경량 PoC를 수행합니다.",
            "• 시나리오 B (사내 보안 망분리 & 최신 상용 모델 도입): Dell/HPE 온프레미스 서버 상에 Cohere (Command R+ / Embed) 또는 사내 오픈소스 sLLM 및 NVIDIA NIM을 Air-Gapped 환경으로 배포합니다.",
            "• 시나리오 C (인프라 현대화 및 점진적 전환): 기존 H/W 인프라를 보호하면서 Nutanix AHV 또는 OpenShift AI를 얹어 AI 워크로드를 점진적으로 증설합니다."
        ],
        border_color_hex="76B900",
        bg_color_hex="F4FBF0"
    )

    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(14)
    p_contact.paragraph_format.space_after = Pt(4)
    r_c_title = p_contact.add_run("■ 세일즈 오퍼링 및 아키텍처 컨설팅 문의")
    r_c_title.font.name = FONT_MAIN
    r_c_title.font.size = Pt(11)
    r_c_title.bold = True
    r_c_title.font.color.rgb = COLOR_NAVY_DARK

    p_ci = doc.add_paragraph()
    p_ci.paragraph_format.space_before = Pt(2)
    p_ci.paragraph_format.space_after = Pt(14)
    p_ci.paragraph_format.left_indent = Cm(0.4)
    r_ci = p_ci.add_run(
        "• 담당 부서: 메가존클라우드 Integrated Solution Sales Unit (ISSU) / AI 하이브리드 솔루션팀\n"
        "• 웹 포털: offering/index.html 및 offering/nvidia_ai_factory_vs_mzc_fullstack_comparison.html\n"
        "• 주요 서비스: 1:1 고객 맞춤형 아키텍처 컨설팅, 벤더 중립적 TCO 시뮬레이션, 온프레미스 AI PoC 랩 지원"
    )
    r_ci.font.name = FONT_MAIN
    r_ci.font.size = Pt(9)
    r_ci.font.color.rgb = COLOR_TEXT_MUTED

    # 저장 실행
    doc.save(str(output_path))
    print(f"✅ 비교 구축 DOCX 문서 생성 완료: {output_path}")
    print(f"   파일 크기: {output_path.stat().st_size / 1024:.1f} KB")


def main():
    base_dir = Path(__file__).resolve().parent
    docx_dir = base_dir / "docx"
    docx_dir.mkdir(exist_ok=True)
    
    output_dated = docx_dir / "2026-08-20_NVIDIA_Dell_MZC_AI_Factory_Comparison_and_Implementation_Guide.docx"
    output_std = docx_dir / "nvidia_ai_factory_vs_mzc_fullstack_comparison.docx"
    
    docs_dir = base_dir.parent / "docs"
    docs_dir.mkdir(exist_ok=True)
    output_docs = docs_dir / "2026-08-20_NVIDIA_Dell_MZC_AI_Factory_Comparison_and_Implementation_Guide.docx"

    build_comparison_document(output_dated)
    build_comparison_document(output_std)
    build_comparison_document(output_docs)


if __name__ == "__main__":
    main()
