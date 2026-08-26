# -*- coding: utf-8 -*-
"""
SambaNova RDU AI Full Stack Feasibility Analysis Report - Word (DOCX) Generator
MEGAZONECLOUD Integrated Solution Sales Unit (ISSU)
"""

import os
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

FONT_MAIN = "맑은 고딕"
FONT_CODE = "Consolas"

# Color Palette
COLOR_NAVY_DARK = RGBColor(0x07, 0x09, 0x0E)
COLOR_TEXT_MAIN = RGBColor(0x22, 0x22, 0x22)
COLOR_TEXT_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_MZC_BLUE = RGBColor(0x00, 0xAB, 0xF0)
COLOR_SAMBA_ORANGE = RGBColor(0xFF, 0x5E, 0x14)
COLOR_CYAN = RGBColor(0x00, 0x99, 0xAA)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_RED = RGBColor(0xCC, 0x00, 0x00)

HEX_BG_DARK_NAVY = "0D111D"
HEX_BG_LIGHT_GRAY = "F4F6F9"
HEX_BG_CALLOUT_RED = "FDE8E8"
HEX_BG_CALLOUT_GOLD = "FEF9E7"
HEX_BG_CALLOUT_BLUE = "E8F4FD"
HEX_BORDER_GRAY = "D1D5DB"
HEX_BORDER_ORANGE = "FF5E14"
HEX_BORDER_RED = "E02424"

def set_cell_shading(cell, color_hex):
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
      <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
      <w:left w:val="none"/>
      <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
      <w:right w:val="none"/>
      <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
      <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def add_cover_page(doc):
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_meta = p_meta.add_run("MEGAZONECLOUD ISSU | FEASIBILITY EVALUATION REPORT")
    r_meta.font.name = FONT_MAIN
    r_meta.font.size = Pt(8.5)
    r_meta.font.color.rgb = COLOR_TEXT_MUTED
    r_meta.bold = True

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(40)

    p_badge = doc.add_paragraph()
    r_badge = p_badge.add_run("【 Third-Party ISV Feasibility Evaluation 】")
    r_badge.font.name = FONT_MAIN
    r_badge.font.size = Pt(10.5)
    r_badge.font.color.rgb = COLOR_SAMBA_ORANGE
    r_badge.bold = True

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(14)
    r_title = p_title.add_run("SambaNova RDU 기술 분석 및\nMZC AI Full Stack 편입 가능성 종합 검토 보고서")
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(21)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(28)
    r_sub = p_sub.add_run("스탠퍼드 RDA SN50 분석, 비즈니스 가드레일(프레임워크 파편화 방지 / Dell·NVIDIA 총판 집중 / 1년 BEP 시뮬레이션) 중심 검토")
    r_sub.font.name = FONT_MAIN
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = COLOR_TEXT_MUTED

    t_box = doc.add_table(rows=4, cols=2)
    t_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_box.autofit = False
    set_table_borders(t_box, color="00ABF0", sz="6")

    meta_info = [
        ("문서 번호", "SPEC-20260820-SAMBANOVA-01"),
        ("발행 일자", "2026년 08월 20일"),
        ("작성 부서", "메가존클라우드(주) ISSU 아키텍처 및 솔루션 프리세일즈 팀"),
        ("분석 대상", "SambaNova Overview Deck (SN40/SN50 RDU, SambaRack, SambaStack)")
    ]

    for idx, (label, val) in enumerate(meta_info):
        row = t_box.rows[idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width, c1.width = Cm(4.0), Cm(12.5)
        set_cell_shading(c0, "EBF6FC")
        set_cell_shading(c1, "F8FBFE")
        set_cell_margins(c0, top=70, bottom=70, left=100, right=100)
        set_cell_margins(c1, top=70, bottom=70, left=100, right=100)

        p0 = c0.paragraphs[0]
        r0 = p0.add_run(label)
        r0.font.name = FONT_MAIN
        r0.font.size = Pt(9.5)
        r0.bold = True
        r0.font.color.rgb = COLOR_MZC_BLUE

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(val)
        r1.font.name = FONT_MAIN
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_TEXT_MAIN

    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(30)
    p_div.paragraph_format.space_after = Pt(20)

    p_exec = doc.add_paragraph()
    r_exec_h = p_exec.add_run("■ Executive Summary (핵심 검토 요약)\n")
    r_exec_h.font.name = FONT_MAIN
    r_exec_h.font.size = Pt(11)
    r_exec_h.bold = True
    r_exec_h.font.color.rgb = COLOR_SAMBA_ORANGE

    r_exec_b = p_exec.add_run(
        "1. 기술적 우수성: SambaNova SN50 RDU는 순차적 토큰 생성(Decode)에 극도로 최적화되어 GPU 대비 10배 빠른 토큰 속도(250+ Tok/s)와 30kW 공랭 전산실 수용성을 제공함.\n"
        "2. 비즈니스 가드레일 검토: MZC는 Dell/NVIDIA 공인 총판 지위를 보유하고 있어 검증된 메인스트림 H/W 중심 영업이 압도적으로 유리하며, 독자 칩셋(RDU) 편입 시 솔루션 프레임워크 파편화가 우려됨.\n"
        "3. 조직 셋업 및 BEP 시뮬레이션: 5인 전담 조직 셋업 시 연간 고정비 약 10~12억 원이 소요되나 국내 예상 수주는 연 0~1건에 불과하여 예상 BEP 도달 기간이 3~4년 소요됨 (1년 내 회수 불가).\n"
        "4. 최종 판정: 200여 개 ISV 재평가 기준(1년 BEP 초과 시 배제)에 따라 'MZC 공식 솔루션 편입 부정적 (No-Go / 단순 기술 검토 단계로 동결)' 판정을 내리며, H/W 직접 핸들링을 배제하고 특수 고객 요청 시에만 전문 수입 파트너를 통한 선별적 서드파티 연계로 제한함."
    )
    r_exec_b.font.name = FONT_MAIN
    r_exec_b.font.size = Pt(9.2)
    r_exec_b.font.color.rgb = COLOR_TEXT_MAIN

    doc.add_page_break()

def add_section_header(doc, num_str, title_str, subtitle_str=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(3)

    r_tag = p.add_run(f"제 {num_str} 장. ")
    r_tag.font.name = FONT_MAIN
    r_tag.font.size = Pt(13)
    r_tag.bold = True
    r_tag.font.color.rgb = COLOR_SAMBA_ORANGE

    r_title = p.add_run(title_str)
    r_title.font.name = FONT_MAIN
    r_title.font.size = Pt(13)
    r_title.bold = True
    r_title.font.color.rgb = COLOR_NAVY_DARK

    if subtitle_str:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(10)
        r_sub = p_sub.add_run(f"— {subtitle_str}")
        r_sub.font.name = FONT_MAIN
        r_sub.font.size = Pt(9.0)
        r_sub.font.color.rgb = COLOR_TEXT_MUTED

def add_sub_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = FONT_MAIN
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = COLOR_NAVY_DARK

def add_narrative_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    r.font.name = FONT_MAIN
    r.font.size = Pt(9.2)
    r.font.color.rgb = COLOR_TEXT_MAIN

def build_docx():
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0)
        sec.right_margin = Cm(2.0)

    # 1. Cover
    add_cover_page(doc)

    # ── [1. ISSU 조직 배경 및 분석 기본 전제] ───────────────────────────
    add_section_header(doc, "1", "ISSU 조직의 본질 및 분석 기본 전제",
                       "솔루션 세일즈/프리세일즈 중심 조직, 파트너 H/W 판매 구조 및 자체 인프라 제약")
    
    add_narrative_paragraph(
        doc,
        "1. ISSU 조직의 본질: 메가존클라우드의 솔루션 세일즈(Solution Sales), 기술 영업 및 프리세일즈(Presales) 중심 조직입니다.\n"
        "2. 하드웨어 특화 솔루션의 성격: 하드웨어를 직접 제조하거나 사내에 실물 장비를 보유하는 구조가 아니며, 당사는 '세일즈 오퍼링(Offering)' 및 아키텍처 컨설팅 패키지 형태로 포트폴리오를 보유합니다 (현재 AIDC 사업에 대하여 검토 중이나, 구체화된 내용은 없음).\n"
        "3. 판매 및 공급 체계: 실제 하드웨어의 조달, 납품, 유지보수 및 구축은 델(Dell)과 엔비디아(NVIDIA) 파트너 에코시스템을 통하여 이루어집니다.\n"
        "4. 현재의 인프라 제약: 현재 ISSU 내부에는 SambaRack이나 RDU 칩셋 등 독자 하드웨어를 직접 설치하여 테스트할 수 있는 자체 실증 Infra가 부재한 상태입니다."
    )

    # ── [2. SambaNova 핵심 기술 분석] ──────────────────────────────────
    add_section_header(doc, "2", "SambaNova 핵심 기술 분석 (RDU & Dataflow Architecture)",
                       "스탠퍼드대 출신 유니콘, TSMC 5nm SN50 칩셋 및 Fast Decode 특화")

    add_narrative_paragraph(
        doc,
        "SambaNova Systems는 2017년 스탠퍼드대 Kunle Olukotun, Christopher Ré 교수 등이 창업한 AI 반도체 전문 기업으로, "
        "최근 10억 달러 규모의 Series F 투자를 유치(기업가치 110억 달러, 약 15조 원)하고 JPMorgan Chase 등의 엔터프라이즈 레퍼런스를 확보하였습니다."
    )

    t_tech = doc.add_table(rows=4, cols=2)
    t_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_tech.autofit = False
    set_table_borders(t_tech)

    tech_data = [
        ("SN50 RDU 프로세서\n(5nm Dataflow Chip)", "• TSMC 5nm 공정 기반 5세대 Reconfigurable Dataflow Unit (RDU)\n• SRAM + HBM2e + DDR 3-Tier 메모리 계층 통합으로 온칩 메모리 병목 원천 해결\n• GPU 대비 최대 10배 빠른 토큰 생성 및 4~5배 향상된 추론 Throughput"),
        ("Fast Decode 특화\n(250+ Tokens/s/user)", "• LLM 추론 중 순차적 토큰 생성(Decode) 단계에 극도로 최적화\n• 250+ Tokens/s/user의 초고속 인터랙티브 속도로 다단계 Agentic Tool Call 지연 극복\n• Llama 3.3 70B, DeepSeek-R1, Qwen3 등 최신 오픈 모델 완벽 지원"),
        ("SambaRack & SambaStack\n(Enterprise Appliance)", "• SambaRack: 256개 RDU 집적 어플라이언스 (랙당 30kW로 일반 공랭 전산실 수용 가능)\n• SambaStack: 온프레미스 및 완전 폐쇄망(Air-Gapped) 프라이빗 추론 플랫폼")
    ]

    for row_idx, (cat, desc) in enumerate(tech_data, start=1):
        cell_c, cell_d = t_tech.rows[row_idx].cells[0], t_tech.rows[row_idx].cells[1]
        cell_c.width, cell_d.width = Cm(4.5), Cm(12.5)
        set_cell_margins(cell_c, top=80, bottom=80, left=80, right=80)
        set_cell_margins(cell_d, top=80, bottom=80, left=80, right=80)
        if row_idx % 2 == 1: set_cell_shading(cell_c, HEX_BG_LIGHT_GRAY); set_cell_shading(cell_d, HEX_BG_LIGHT_GRAY)

        p_c = cell_c.paragraphs[0]
        r_c = p_c.add_run(cat)
        r_c.font.name = FONT_MAIN; r_c.font.size = Pt(8.5); r_c.bold = True; r_c.font.color.rgb = COLOR_SAMBA_ORANGE

        p_d = cell_d.paragraphs[0]
        r_d = p_d.add_run(desc)
        r_d.font.name = FONT_MAIN; r_d.font.size = Pt(8.2); r_d.font.color.rgb = COLOR_TEXT_MAIN

    # Table Header
    c0, c1 = t_tech.rows[0].cells[0], t_tech.rows[0].cells[1]
    c0.width, c1.width = Cm(4.5), Cm(12.5)
    set_cell_shading(c0, HEX_BG_DARK_NAVY); set_cell_shading(c1, HEX_BG_DARK_NAVY)
    set_cell_margins(c0, top=100, bottom=100, left=100, right=100); set_cell_margins(c1, top=100, bottom=100, left=100, right=100)
    p0 = c0.paragraphs[0]; r0 = p0.add_run("구분"); r0.font.name = FONT_MAIN; r0.font.size = Pt(8.5); r0.bold = True; r0.font.color.rgb = COLOR_WHITE
    p1 = c1.paragraphs[0]; r1 = p1.add_run("SambaNova 핵심 아키텍처 및 상세 스펙"); r1.font.name = FONT_MAIN; r1.font.size = Pt(8.5); r1.bold = True; r1.font.color.rgb = COLOR_WHITE

    # ── [3. NVIDIA GPU vs SambaNova RDU 비교 분석] ─────────────────────
    add_section_header(doc, "3", "NVIDIA GPU vs SambaNova RDU 심층 기술 비교",
                       "아키텍처, Decode 속도, 소프트웨어 생태계 및 전력 소모 비교")

    t_cmp = doc.add_table(rows=6, cols=3)
    t_cmp.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_cmp.autofit = False
    set_table_borders(t_cmp)

    cmp_headers = [("비교 항목", Cm(3.5)), ("NVIDIA GPU 클러스터\n(B200 / H100 / L40S)", Cm(6.5)), ("SambaNova RDU 어플라이언스\n(SN50 SambaRack)", Cm(7.0))]
    for col_idx, (h_text, w_val) in enumerate(cmp_headers):
        cell = t_cmp.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]; r = p.add_run(h_text)
        r.font.name = FONT_MAIN; r.font.size = Pt(8.2); r.bold = True; r.font.color.rgb = COLOR_WHITE

    cmp_data = [
        ("칩셋 아키텍처", "범용 SIMD/SIMT 병렬 GPU", "재구성 가능 데이터플로우(RDA) RDU"),
        ("추론 속도 (Decode)", "사용자당 30 ~ 80 Tokens/sec", "사용자당 200 ~ 250+ Tokens/sec (3~5배 빠름)"),
        ("소프트웨어 생태계", "NVIDIA CUDA, TensorRT-LLM (글로벌 표준)", "SambaFlow 독자 컴파일러 (이식성 장벽 존재)"),
        ("데이터센터 전력", "NVL-72 기준 랙당 150kW (수랭식 필수)", "랙당 30kW (일반 엔터프라이즈 공랭 수용 가능)"),
        ("소프트웨어 호환성", "모든 오픈소스/상용 LLM 100% 즉시 지원", "PyTorch 모델 컴파일 필요 (커스텀 오퍼레이터 제약)")
    ]

    for row_idx, r_data in enumerate(cmp_data, start=1):
        row_cells = t_cmp.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(3.5), Cm(6.5), Cm(7.0)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]; r = p.add_run(text)
            r.font.name = FONT_MAIN; r.font.size = Pt(7.8); r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0: r.bold = True; r.font.color.rgb = COLOR_NAVY_DARK

    # ── [4. 신규 조직 셋업 비용 및 BEP 시뮬레이션] ─────────────────────
    add_section_header(doc, "4", "신규 비즈니스 인력/조직 셋업 비용 및 BEP 시뮬레이션",
                       "5인 전담 조직 셋업 고정비, 국내 시장 수주 현실 및 1년 BEP 초과 평가")

    t_bep = doc.add_table(rows=5, cols=3)
    t_bep.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_bep.autofit = False
    set_table_borders(t_bep)

    bep_headers = [("비즈니스 셋업 & 수익 항목", Cm(4.5)), ("비용 및 시장 현실 산정 근거", Cm(8.0)), ("연간 재무 영향", Cm(4.5))]
    for col_idx, (h_text, w_val) in enumerate(bep_headers):
        cell = t_bep.cell(0, col_idx)
        cell.width = w_val
        set_cell_shading(cell, HEX_BG_DARK_NAVY)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]; r = p.add_run(h_text)
        r.font.name = FONT_MAIN; r.font.size = Pt(8.2); r.bold = True; r.font.color.rgb = COLOR_WHITE

    bep_data = [
        ("전담 조직 셋업 비용\n(5인 전담팀 구성)", "• RDU 전문 솔루션 아키텍트 2명 + 세일즈 2명 + PM 1명\n• 미국 본사 공인 엔지니어링 인증 및 데모 랩 운영비", "연간 고정비 약 10억 ~ 12억 원"),
        ("건당 사업 순마진\n(Unit Margin)", "• SambaRack 어플라이언스 및 플랫폼 구축 공급 마진\n• 대규모 구축 건당 평균 순마진 약 1.5억 ~ 2.5억 원", "건당 순마진 약 2.0억 원"),
        ("국내 시장 수주 현실\n(Annual Pipeline)", "• 10억 원 이상 RDU 전용 랙 도입 가능한 국내 엔터프라이즈 극소수\n• 1차년도 보수적 예상 수주: 연간 0 ~ 1건", "1차년도 매출 이익 약 0 ~ 2억 원\n(연간 8~10억 적자 발생)"),
        ("예상 BEP 도달 기간", "연간 고정비(11억) 회수를 위해 연간 최소 5~6대 이상 지속 수주 필요\n➔ 단기간 내 달성 불가능", "최소 3년 ~ 4년 이상 소요\n(1년 내 달성 불가 ➔ 부정적)")
    ]

    for row_idx, r_data in enumerate(bep_data, start=1):
        row_cells = t_bep.rows[row_idx].cells
        for col_idx, (text, w_val) in enumerate(zip(r_data, [Cm(4.5), Cm(8.0), Cm(4.5)])):
            cell = row_cells[col_idx]
            cell.width = w_val
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            if row_idx % 2 == 1: set_cell_shading(cell, HEX_BG_LIGHT_GRAY)
            p = cell.paragraphs[0]; r = p.add_run(text)
            r.font.name = FONT_MAIN; r.font.size = Pt(7.8); r.font.color.rgb = COLOR_TEXT_MAIN
            if col_idx == 0: r.bold = True; r.font.color.rgb = COLOR_NAVY_DARK
            elif col_idx == 2 and row_idx == 4: r.bold = True; r.font.color.rgb = COLOR_RED

    add_narrative_paragraph(
        doc,
        "• 200여 가지 ISV 파트너 재평가 기준 적용: 경쟁력, 시장 성장성, 매출 기여도를 기준으로 200여 개 파트너십을 엄격히 정예화했듯이, "
        "기술적으로 문제가 없더라도 BEP가 1년 단위 이상 소요(3~4년)되는 신규 솔루션은 비즈니스 편입에 대해 부정적으로 평가합니다.\n"
        "• 현재 상태 규정: 현재 SambaNova는 '단순 기술적 검토(Technical Exploration)' 단계로 동결하며, 신규 전담 조직 셋업이나 공식 솔루션 추가는 전면 보류합니다."
    )

    # ── [5. AI Full Stack 합류 및 총판 사업 타당성 판단] ───────────────
    add_section_header(doc, "5", "MZC AI Full Stack 합류 및 총판 사업 타당성 판단",
                       "프레임워크 파편화 방지, 분기 100억 총판 허들 및 H/W 직접 핸들링 배제")

    add_sub_header(doc, "■ 1. 솔루션 프레임워크 파편화 방지 및 공식 편입 부정적")
    add_narrative_paragraph(
        doc,
        "• 프레임워크 파편화 리스크: 독자 RDU 칩셋 및 전용 SambaRack을 표준 스택에 편입할 경우, 기존 NVIDIA CUDA 중심 표준 아키텍처와 이원화되어 엔지니어링 리소스가 파편화됩니다.\n"
        "• 분기 100억 원 총판 허들 미달: MZC 총판 비즈니스 편입 기준(분기 100억 원 이상)에 크게 미달하며, BEP 도달 기간이 1년 이상 소요되므로 총판 사업 편입은 불가합니다.\n"
        "• 하드웨어 직접 핸들링 배제: MZC가 실물 SambaRack 수입·재고·AS를 직접 담당하는 위험은 원천 차단합니다."
    )

    add_sub_header(doc, "■ 2. Dell / NVIDIA 공인 총판 지위 활용에 집중")
    add_narrative_paragraph(
        doc,
        "• 기존 핵심 총판 역량 레버리지: MZC가 이미 공인 총판 지위를 확보하고 있는 Dell PowerEdge 서버 및 NVIDIA GPU 표준 라인업에 모든 영업 및 엔지니어링 역량을 집중합니다.\n"
        "• 특수 고객 제한적 연계: JPMorgan 사례처럼 금융사 등에서 SambaNova를 콕 집어 요구하는 특수 사이트에 한하여, 물리 H/W는 서드파티 파트너(Unitrontech 등)가 납품하고 MZC는 S/W SI로만 선별 대응합니다."
    )

    # ── [6. 결론 및 실행 제언] ─────────────────────────────────────────
    add_section_header(doc, "6", "종합 결론 및 전략적 의사결정 가이드라인",
                       "Dell/NVIDIA 총판 비즈니스 집중 및 H/W 직접 핸들링 배제")

    conclusions = [
        ("• 1. 현재 상태 규정: '단순 기술적 검토' 단계로 유지: ",
         "250+ Tok/s Fast Decode의 기술적 우수성은 확인하였으나, BEP 1년 초과(3~4년) 및 전담 조직 셋업 부담으로 인해 MZC 공식 솔루션 편입은 부정적(보류)으로 결론"),
        ("• 2. Dell / NVIDIA 공인 총판 사업에 전사 역량 집중 (Core Business First): ",
         "솔루션 프레임워크 파편화를 방지하고, 메가존클라우드의 핵심 강점인 Dell PowerEdge + NVIDIA 표준 엔터프라이즈 인프라 비즈니스에 전사 역량 집중"),
        ("• 3. H/W 직접 핸들링 배제 및 선별적 서드파티 파트너 연계: ",
         "MZC의 하드웨어 직접 총판/유통은 추진하지 않으며, 고객사 특수 요청 시에만 전문 H/W 수입 파트너를 통한 서드파티 납품 연계로 위험 통제")
    ]

    for c_title, c_desc in conclusions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r_t = p.add_run(c_title)
        r_t.font.name = FONT_MAIN; r_t.font.size = Pt(8.8); r_t.bold = True; r_t.font.color.rgb = COLOR_NAVY_DARK
        r_d = p.add_run(c_desc)
        r_d.font.name = FONT_MAIN; r_d.font.size = Pt(8.8); r_d.font.color.rgb = COLOR_TEXT_MAIN

    # Output paths
    out_dir_offering = r"C:\dev\antigravity-workspace\aifullstack\offering\docx"
    out_dir_docs = r"C:\dev\antigravity-workspace\aifullstack\docs"
    os.makedirs(out_dir_offering, exist_ok=True)
    os.makedirs(out_dir_docs, exist_ok=True)

    file_name = "2026-08-20_SambaNova_AI_Fullstack_Analysis_Report.docx"
    path_offering = os.path.join(out_dir_offering, file_name)
    path_docs = os.path.join(out_dir_docs, file_name)

    doc.save(path_offering)
    shutil.copy2(path_offering, path_docs)

    sz_kb = os.path.getsize(path_offering) / 1024
    print(f"✅ SambaNova 분석 보고서 DOCX 생성 완료: {path_offering}")
    print(f"   파일 크기: {sz_kb:.1f} KB")
    print(f"✅ SambaNova 분석 보고서 DOCX 생성 완료: {path_docs}")
    print(f"   파일 크기: {sz_kb:.1f} KB")

if __name__ == "__main__":
    build_docx()
